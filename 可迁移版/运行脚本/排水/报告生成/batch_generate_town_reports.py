import os
import re
import zipfile
from collections import defaultdict
from copy import deepcopy
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


PACKAGE_ROOT = Path(__file__).resolve().parents[3]
PROJECT_ROOT = PACKAGE_ROOT.parent if (PACKAGE_ROOT.parent / "资料收集").is_dir() else PACKAGE_ROOT
SOURCE = PACKAGE_ROOT / "skills" / "report" / "assets" / "正文底稿.docx"
OUTPUTS = PACKAGE_ROOT / "outputs"
FINAL_DIR = Path(os.environ.get("REPORT_OUTPUT_DIR", PROJECT_ROOT / "生成"))
DATA_DIR = Path(os.environ.get("REPORT_DATA_DIR", PROJECT_ROOT / "资料收集"))
AMOUNT_FALLBACK = PACKAGE_ROOT / "skills" / "report" / "assets" / "common_amount_basis.xlsx"

TOWNS = [
    "北陡镇",
    "汶村镇",
    "白沙镇",
    "三合镇",
    "深井镇",
    "四九镇",
    "大江镇",
    "都斛镇",
    "赤溪镇",
    "冲蒌镇",
    "端芬镇",
    "川岛镇",
    "广海镇",
    "海宴镇",
    "台城街道",
    "水步镇",
    "斗山镇",
]

def clean_text(text):
    return "".join(str(text or "").split())


def row_texts_xml(row):
    cells = []
    for tc in row._tr.tc_lst:
        cells.append(clean_text("".join(t.text or "" for t in tc.iter() if t.tag.endswith("}t"))))
    return cells


def set_tc_text(tc, text):
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs = tc.findall(f".//{ns}p")
    if not paragraphs:
        p = OxmlElement("w:p")
        tc.append(p)
        paragraphs = [p]
    first_p = paragraphs[0]
    for child in list(first_p):
        first_p.remove(child)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    first_p.append(r)


def set_para_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def is_title_paragraph(paragraph):
    text = paragraph.text.strip()
    if not text:
        return False
    style = paragraph.style.name if paragraph.style is not None else ""
    if style.startswith("Heading") or "标题" in style:
        return True
    if re.match(r"^第[一二三四五六七八九十]+章", text):
        return True
    if re.match(r"^\d+(\.\d+)*\s+", text):
        return True
    if text.endswith("考核情况") or text.endswith("评分情况") or text.endswith("汇总表") or text.endswith("评价表"):
        return True
    if "绩效考核服务项目" in text and ("报告" in text or "正文" in text):
        return True
    return False


def apply_font_rules(doc):
    for paragraph in doc.paragraphs:
        title = is_title_paragraph(paragraph)
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run._element.rPr.rFonts.set(qn("w:ascii"), "宋体")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "宋体")
            run.bold = True if title else False
    for table in doc.tables:
        for row in table._tbl.tr_lst:
            for tc in row.tc_lst:
                for r in tc.iter():
                    if not r.tag.endswith("}r"):
                        continue
                    rpr = r.find(qn("w:rPr"))
                    if rpr is None:
                        rpr = OxmlElement("w:rPr")
                        r.insert(0, rpr)
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is None:
                        rfonts = OxmlElement("w:rFonts")
                        rpr.append(rfonts)
                    rfonts.set(qn("w:eastAsia"), "宋体")
                    rfonts.set(qn("w:ascii"), "宋体")
                    rfonts.set(qn("w:hAnsi"), "宋体")
                    for b in rpr.findall(qn("w:b")) + rpr.findall(qn("w:bCs")):
                        rpr.remove(b)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run._element.rPr.rFonts.set(qn("w:ascii"), "宋体")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "宋体")
                    run.bold = False


def decimal2(value):
    return Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def display_number(value, places=2):
    if value in (None, ""):
        return ""
    quantized = Decimal(str(value)).quantize(
        Decimal("1") if places == 0 else Decimal("1." + "0" * places),
        rounding=ROUND_HALF_UP,
    )
    text = f"{quantized:.{places}f}"
    return text.rstrip("0").rstrip(".") if "." in text else text


def read_xlsx_rows(path):
    """Read the first worksheet with only the standard library."""
    ns = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        shared = []
        if "xl/sharedStrings.xml" in archive.namelist():
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            for item in root.findall("x:si", ns):
                shared.append("".join(node.text or "" for node in item.findall(".//x:t", ns)))

        workbook = ET.fromstring(archive.read("xl/workbook.xml"))
        rels = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        rel_map = {
            rel.attrib["Id"]: rel.attrib["Target"]
            for rel in rels
        }
        sheet = workbook.find("x:sheets/x:sheet", ns)
        rel_id = sheet.attrib["{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"]
        target = rel_map[rel_id].lstrip("/")
        sheet_path = target if target.startswith("xl/") else f"xl/{target}"
        root = ET.fromstring(archive.read(sheet_path))

        rows = []
        for row in root.findall(".//x:sheetData/x:row", ns):
            values = {}
            max_col = -1
            for cell in row.findall("x:c", ns):
                ref = cell.attrib.get("r", "A1")
                letters = re.match(r"[A-Z]+", ref).group(0)
                col = 0
                for letter in letters:
                    col = col * 26 + ord(letter) - 64
                col -= 1
                max_col = max(max_col, col)
                value_node = cell.find("x:v", ns)
                inline = cell.find("x:is", ns)
                if inline is not None:
                    value = "".join(node.text or "" for node in inline.findall(".//x:t", ns))
                elif value_node is None:
                    value = ""
                elif cell.attrib.get("t") == "s":
                    value = shared[int(value_node.text)]
                else:
                    value = value_node.text or ""
                values[col] = value
            rows.append([values.get(i, "") for i in range(max_col + 1)])
    return rows


def load_amount_basis():
    path = DATA_DIR / "金额基础数据.xlsx"
    if not path.exists():
        path = AMOUNT_FALLBACK
    if not path.exists():
        raise FileNotFoundError("缺少金额基础数据.xlsx，且未找到通用金额基数。")

    rows = read_xlsx_rows(path)
    if not rows:
        raise ValueError(f"金额基础表为空：{path}")
    headers = [clean_text(value) for value in rows[0]]
    required = [
        "镇街", "可用性付费基数Pk3（万元/年）", "运维服务费基数Py3（万元/年）",
        "第一批运维服务费基数（元/月）", "第九批运维服务费基数（元/月）",
        "建设期考核系数E1",
    ]
    missing = [name for name in required if name not in headers]
    if missing:
        raise ValueError("金额基础表缺少列：" + "、".join(missing))

    result = {}
    for row in rows[1:]:
        record = {headers[i]: row[i] if i < len(row) else "" for i in range(len(headers))}
        town = clean_text(record.get("镇街"))
        if town not in TOWNS:
            continue
        batch_names = ["第一", "第二", "第三", "第四", "第五", "第六", "第七", "第八", "第九"]
        batches = [
            Decimal(record.get(f"{name}批运维服务费基数（元/月）") or "0")
            for name in batch_names
        ]
        result[town] = {
            "facility_count": record.get("农村污水处理设施点数", ""),
            "design_scale": record.get("设计处理规模", ""),
            "pk3": Decimal(record["可用性付费基数Pk3（万元/年）"] or "0"),
            "py3": Decimal(record["运维服务费基数Py3（万元/年）"] or "0"),
            "batches": batches,
            "e1": Decimal(record.get("建设期考核系数E1") or "1"),
        }
    missing_towns = [town for town in TOWNS if town not in result]
    if missing_towns:
        raise ValueError("金额基础表缺少镇街：" + "、".join(missing_towns))
    return result


def validate_common_material():
    path = DATA_DIR / "公共资料" / "公共资料.docx"
    if not path.exists():
        raise FileNotFoundError(f"缺少公共资料：{path}")
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "考核标准" not in text or len(doc.tables) < 2:
        raise ValueError(f"公共资料缺少考核标准或标准表：{path}")
    return path


def append_row_from_template(table, values):
    row = deepcopy(table.rows[-1]._tr)
    table._tbl.append(row)
    cells = row.tc_lst
    for index, cell in enumerate(cells):
        set_tc_text(cell, str(values[index]) if index < len(values) else "")
    return row


def set_table_font(table, size=8):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(size)
                    run.bold = False


def collect_metrics():
    metrics = {town: {"sample": [], "pay": {}} for town in TOWNS}

    for town, data in metrics.items():
        path = DATA_DIR / town / f"{town}附件资料.docx"
        if not path.exists():
            raise FileNotFoundError(f"缺少资料收集文件：{path}")
        doc = Document(path)
        titles = []
        town_summary_title = f"{town}农村污水处理设施绩效评价表"
        for paragraph in doc.paragraphs:
            title = paragraph.text.strip()
            if title.endswith("绩效评价表") and title != town_summary_title:
                titles.append(title)
        score_tables = []
        for table in doc.tables:
            if len(table.rows) < 3 or len(table.columns) < 7:
                continue
            last = [clean_text(c.text) for c in table.rows[-1].cells]
            if "评分" in last[:5]:
                try:
                    score = Decimal(last[6])
                except Exception:
                    continue
                score_tables.append((table, score))
        for idx, (table, score) in enumerate(score_tables):
            if score >= 90:
                ec1, ec2 = Decimal("1"), Decimal("1")
            elif score >= 80:
                ec1, ec2 = (Decimal("100") - (Decimal("90") - score) * Decimal("0.5")) / Decimal("100"), Decimal("1")
            elif score >= 70:
                ec1 = (Decimal("95") - (Decimal("80") - score) * Decimal("1")) / Decimal("100")
                ec2 = (Decimal("100") - (Decimal("80") - score) * Decimal("0.5")) / Decimal("100")
            else:
                ec1 = (Decimal("85") - (Decimal("70") - score) * Decimal("1.5")) / Decimal("100")
                ec2 = (Decimal("95") - (Decimal("70") - score) * Decimal("1")) / Decimal("100")
            title = titles[idx] if idx < len(titles) else f"第{idx + 1}个设施"
            match = re.search(r"（(.+)）污水(?:运营|处理设施)绩效评价表$", title)
            name = match.group(1) if match else re.sub(r"污水(?:运营|处理设施)绩效评价表$", "", title)
            data["sample"].append({
                "name": name,
                "title": title,
                "score": score,
                "ec1": ec1,
                "ec2": ec2,
                "deductions": [row_texts_xml(row)[6] if len(row_texts_xml(row)) > 6 else "" for row in table.rows],
            })

        data["criteria"] = []
        if score_tables:
            for row in score_tables[0][0].rows:
                values = row_texts_xml(row)
                data["criteria"].append((values + [""] * 4)[:4])

        for table in doc.tables:
            header = "|".join(clean_text(c.text) for r in table.rows[:2] for c in r.cells)
            if "进水" in header and "出水是否达标" in header and "COD" in header:
                water = {"count": 0, "out_ok": 0, "cod_low": 0, "cod_high": 0, "nh3_high": 0}
                water_rows = []
                for row in table.rows[2:]:
                    vals = [clean_text(c.text) for c in row.cells]
                    if len(vals) < 13 or vals[1] != town:
                        continue
                    water_rows.append(vals)
                    water["count"] += 1
                    water["out_ok"] += 1 if vals[12] == "是" else 0
                    try:
                        cod = Decimal(vals[8].replace("L", ""))
                        nh3 = Decimal(vals[9].replace("L", ""))
                    except Exception:
                        continue
                    water["cod_low"] += 1 if cod < 100 else 0
                    water["cod_high"] += 1 if cod > 170 else 0
                    water["nh3_high"] += 1 if nh3 > 28 else 0
                data["water"] = water
                data["water_rows"] = water_rows
            if "村民满意度得分" in header and "实施机构满意度评分" in header:
                last = [clean_text(c.text) for c in table.rows[-1].cells]
                if len(last) >= 9:
                    data["satisfaction"] = {"public": last[-3], "town": last[-2], "agency": last[-1]}

        scores = [item["score"] for item in data["sample"]]
        if scores:
            data["sample_count"] = len(scores)
            data["score_ge90"] = sum(1 for s in scores if s >= 90)
            data["score_80_90"] = sum(1 for s in scores if 80 <= s < 90)
            data["score_lt80"] = sum(1 for s in scores if s < 80)
            ec1_values = [Decimal(item["ec1"]) for item in data["sample"] if item["ec1"]]
            ec2_values = [Decimal(item["ec2"]) for item in data["sample"] if item["ec2"]]
            data["ec1"] = (sum(ec1_values) / Decimal(len(ec1_values))).quantize(Decimal("0.001"), rounding=ROUND_HALF_UP)
            data["ec2"] = (sum(ec2_values) / Decimal(len(ec2_values))).quantize(Decimal("0.001"), rounding=ROUND_HALF_UP)
        if len(data.get("water_rows", [])) != len(data["sample"]):
            raise ValueError(
                f"{town}水质抽检行数({len(data.get('water_rows', []))})与评分表数({len(data['sample'])})不一致"
            )
    return metrics


def attach_payments(metrics, amount_basis):
    for town, data in metrics.items():
        basis = amount_basis[town]
        ec1 = data["ec1"]
        ec2 = data["ec2"]
        pk_month = decimal2(basis["pk3"] * Decimal("10000") / Decimal("12"))
        availability = decimal2(pk_month * basis["e1"] * ec2)
        om_base = sum(basis["batches"][:7])
        om_payment = decimal2(om_base * ec1)
        batch8_payment = decimal2(basis["batches"][7] * ec1)
        data["basis"] = basis
        data["pay"] = {
            "availability_base": pk_month,
            "availability": availability,
            "availability_deduct": decimal2(pk_month - availability),
            "om_base": decimal2(om_base),
            "om": om_payment,
            "om_deduct": decimal2(om_base - om_payment),
            "batch8": batch8_payment,
            "batch8_deduct": decimal2(basis["batches"][7] - batch8_payment),
            "batch9": basis["batches"][8],
        }
        data["pay"]["total_deduct"] = decimal2(
            data["pay"]["availability_deduct"]
            + data["pay"]["om_deduct"]
            + data["pay"]["batch8_deduct"]
        )


def update_toc(doc, targets):
    if isinstance(targets, str):
        targets = [targets]
    body = doc._body._element
    anchor = None
    for p in list(doc.paragraphs):
        if p.text.strip() == "目录":
            anchor = p
        if p.style.name.lower().startswith("toc"):
            body.remove(p._element)
    if anchor is None:
        return
    lines = [
        "第一章  考核工作概述",
        "1.1  考核目的",
        "1.2  考核要求",
        "1.3  考核依据",
        "1.4  考核成员及分工责任",
        "1.5  考核频次",
        "1.6  考核方法",
        "1.7  考核安排",
        "第二章  农村污水处理设施考核结果",
        "2.1  农村污水处理设施考核情况",
        *[
            f"2.{index + 2}  {target}农村污水处理设施考核情况"
            for index, target in enumerate(targets)
        ],
        "第三章  绩效付费计算",
        "3.1  付费范围",
        "3.2  付费依据",
        "3.3  农村污水处理服务费支付计算方法",
        "3.4  农村污水处理设施项目服务费",
        "3.5  计费系数",
        "3.6  农村污水处理设施服务费计算",
        "3.7  付费建议",
        "3.8  考核付费结果分析",
        "第四章  主要问题及整改工作建议",
        "4.1  主要问题",
        "4.2  项目公司整改建议",
    ]
    after = anchor._element
    for line in lines:
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = line
        r.append(t)
        p.append(r)
        after.addnext(p)
        after = p


def replace_overview_paragraphs(doc, target, metrics):
    data = metrics.get(target, {})
    count = data.get("sample_count", 0)
    ge90 = data.get("score_ge90", 0)
    mid = data.get("score_80_90", 0)
    low = data.get("score_lt80", 0)
    ec1 = data.get("ec1", "")
    ec2 = data.get("ec2", "")
    pay = data.get("pay", {})
    availability = pay.get("availability", "")
    om = pay.get("om", "")
    availability_deduct = pay.get("availability_deduct", "0")
    om_deduct = pay.get("om_deduct", "0")
    batch8_deduct = pay.get("batch8_deduct", "0")
    total_deduct = pay.get("total_deduct", "0")

    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("各镇村级污水处理设施陆续于2021") and "共抽检220个农村污水处理设施" in text:
            set_para_text(
                p,
                f"各镇村级污水处理设施陆续于2021 ~2023 年正式投入商业运营，并满足考核条件。"
                f"2022年10~12 月，我司按照《PPP项目合同》及相关绩效考核要求开展了本项目第一次村级污水处理设施绩效考核工作。"
                f"2023年11月20日，台山市水利局与江门路航环保科技有限公司签订《PPP项目合同补充协议》。"
                f"我司根据《PPP项目合同补充协议》要求开展绩效考核工作，并给出绩效考核结果与服务费挂钩结论。"
                f"本次为2023年下半年度考核，是本项目村级设施的第三次考核，本报告涉及{target}共抽检{count}个农村污水处理设施，"
                f"考核结果作为该镇已转运营农村污水处理设施下一周期的付费依据。",
            )
        elif text.startswith("根据台山市水利局工作安排及合同相关约定，2023年12月~2024年1月") and "17个镇街220个农村污水处理设施" in text:
            set_para_text(
                p,
                f"根据台山市水利局工作安排及合同相关约定，2023年12月~2024年1月，由台山市水利局、"
                f"广东省建筑设计研究院集团股份有限公司、镇级行政主管单位等各派代表组成考核小组，"
                f"对{target}{count}个农村污水处理设施进行了2023年下半年度现场考核，同步开展问卷调查，"
                f"并让各考核组成员及被考核单位现场签字确认。结合水质检测结果，查阅项目公司和实施机构提供的考核资料，"
                f"对农村污水处理设施运维情况进行考核评分。",
            )
        elif text.startswith("各镇农村污水处理设施绩效考核中，220个农村污水处理设施中"):
            low_part = f"，{low}个设施评分低于80分" if low else ""
            set_para_text(
                p,
                f"{target}农村污水处理设施绩效考核中，{count}个农村污水处理设施中有{ge90}个设施考核评分在90分及以上，"
                f"{mid}个设施评分在80~90分之间{low_part}。全镇运维期绩效考核运维服务付费系数Ec1为{ec1}，"
                f"运维期绩效考核可用性服务付费系数Ec2为{ec2}。",
            )
        elif text.startswith("根据台山市水利局工作安排及合同相关约定，本次集中对台山市第二轮农村生活污水处理设施建设PPP项目") and "17个镇街220个农村污水处理设施" in text:
            set_para_text(
                p,
                f"根据台山市水利局工作安排及合同相关约定，本次集中对台山市第二轮农村生活污水处理设施建设PPP项目中"
                f"{target}农村生活污水处理设施开展2023年下半年度绩效考核，考核对象包括{target}{count}个抽检农村污水处理设施。",
            )
        elif text.startswith("2023年12月~2024年1月，考核小组对台山市17个镇共220个农村污水处理设施进行了考核"):
            low_part = f"，{low}个设施评分低于80分" if low else ""
            set_para_text(
                p,
                f"2023年12月~2024年1月，考核小组对{target}{count}个农村污水处理设施进行了考核，"
                f"其中有{ge90}个设施的现场运维绩效考核评分在90分及以上，{mid}个设施评分在80~90分之间{low_part}。",
            )
        elif text.startswith("根据《PPP项目合同》、《PPP项目合同补充协议》和项目完工转运营情况") and "共1256个" in text:
            set_para_text(
                p,
                f"根据《PPP项目合同》、《PPP项目合同补充协议》和项目完工转运营情况，2023年下半年度付费范围包括"
                f"第二轮农村生活污水处理设施建设PPP项目中{target}已转入运营的农村污水处理设施，"
                f"详细情况见《PPP项目合同补充协议》附件二表3。",
            )
        elif availability and text.startswith("根据台山市第二轮农村污水处理设施建设PPP项目农村污水处理设施2023年下半年度绩效考核结果") and "可用性付费每月付费建议" in text:
            set_para_text(
                p,
                f"根据台山市第二轮农村污水处理设施建设PPP项目农村污水处理设施2023年下半年度绩效考核结果，"
                f"2024年1~6月{target}转运营农村污水处理设施可用性付费每月付费建议为{availability}元（详见表3-4）。",
            )
        elif om and text.startswith("1.转运营的第一~七批子项目运营2024年1月~6月每月的农村污水处理设施运维服务费每月付费建议为"):
            set_para_text(
                p,
                f"1.转运营的第一~七批子项目运营2024年1月~6月每月的农村污水处理设施运维服务费每月付费建议为{om}元（详见表3-5）。",
            )
        elif text.startswith("2.第八批子项目转运营未满6个月"):
            set_para_text(p, f"2.{target}第八批子项目运维服务费按表3-6列示情况计算。")
        elif text.startswith("3.第九批子项目转运营未满6个月"):
            set_para_text(p, f"3.{target}第九批子项目运维服务费按表3-6列示情况计算。")
        elif total_deduct and text.startswith("根据第3.6、3.7章节的付费情况，相比于合同约定的可用性付费") and "合计扣减费用" in text:
            set_para_text(
                p,
                f"根据第3.6、3.7章节的付费情况，相比于合同约定的可用性付费和第一~七批农村运维服务费基数，"
                f"{target}农村污水处理设施每月扣减可用性付费为{availability_deduct}元，"
                f"第一~七批农村运维服务费每月扣减费用为{om_deduct}元，第八批农村运维服务费每月扣减费用为{batch8_deduct}元，"
                f"合计扣减费用为{total_deduct}元/月，详细见下表。",
            )


def scrub_cross_town_examples(doc, target):
    other_towns = [town for town in TOWNS if town != target]
    replacements = [
        ("六是社会服务评价方面有待提升", f"六是社会服务评价方面有待提升。{target}农村污水设施点中，镇级主管部门及公众评价存在少量扣分，后续仍需结合村民反馈意见持续提升服务质量。"),
        ("注：因在计算过程中小数点四舍五入影响", "注：因在计算过程中小数点四舍五入影响，各镇农村污水处理服务费中可用性付费之和与总可用性付费存在差异，总可用性付费保持不变。"),
        ("n1—某镇第n批一体化设施总数", "n1—某镇第n批一体化设施总数，单位为个。"),
        ("村内仅采取末端截污方式收集污水", f"1.村内污水收集效果仍需提升。{target}部分设施服务范围内仍存在末端截污、明渠收集、局部污水未完全收集等情况，影响村内水体环境和污水处理设施运行效能。"),
        ("截污口设置不合理", f"2.截污口及检查井设置维护仍需加强。{target}部分设施存在截污口设置、管网衔接或检查井维护不到位等问题，需结合现场情况持续整改。"),
        ("检查井或集水井存在溢流直排", f"3.检查井或集水井运行维护仍需加强。{target}部分设施需进一步排查高水位、溢流及排水不畅等情况，保障污水有效收集处理。"),
        ("部分设施进水CODCr或NH3-N浓度较高", f"2.部分设施进水水质指标波动较大。{target}抽检设施中存在进水CODCr或NH3-N浓度偏高情形，需持续排查进水来源并强化运行调控。"),
        ("管网运维有待进一步提升", f"1.管网运维有待进一步提升。{target}部分管网存在水位偏高、排水不畅、井内杂物清理不及时等问题，需进一步加强巡查维护。"),
        ("检查井或管网建设不规范", f"2.检查井或管网建设维护需进一步规范。{target}部分设施需结合现场排水路径、井体设置和管网运行情况开展针对性整改。"),
        ("设施运维有待进一步提升", f"1.设施运维有待进一步提升。{target}部分预处理设施、提升泵、曝气设备及控制系统需加强巡查维护，确保设施稳定运行。"),
        ("采用VFS简易设施运行不稳定", f"2.采用VFS等简易设施的站点需加强运行维护。{target}相关设施应结合供电、曝气、抽水及周边遮挡等情况及时维护，提升运行稳定性。"),
        ("部分设施点存在安全隐患", f"部分设施点安全管理仍需持续加强。{target}应继续排查防护栏、警示标识、信息公示牌及周边环境等安全管理事项，及时消除隐患。"),
        ("台城街道的农村设施村民满意度有待进一步提升", f"{target}农村设施村民满意度仍有提升空间。下半年度无政府部门处罚、社会有效投诉或公众媒体有效负面报道，但仍需结合镇级主管部门及村民反馈意见，持续提升污水收集效果、运行维护质量和群众满意度。"),
    ]
    for p in doc.paragraphs:
        text = p.text.strip()
        if not any(town in text for town in other_towns):
            continue
        for prefix, new_text in replacements:
            if text.startswith(prefix) or prefix in text:
                set_para_text(p, new_text)
                break


def replace_global_metrics(doc, targets, metrics):
    sample_count = sum(metrics[town]["sample_count"] for town in targets)
    ge90 = sum(metrics[town]["score_ge90"] for town in targets)
    mid = sum(metrics[town]["score_80_90"] for town in targets)
    low = sum(metrics[town]["score_lt80"] for town in targets)
    operated_count = sum(Decimal(str(metrics[town]["basis"]["facility_count"] or "0")) for town in targets)
    water_count = sum(metrics[town].get("water", {}).get("count", 0) for town in targets)
    out_ok = sum(metrics[town].get("water", {}).get("out_ok", 0) for town in targets)
    cod_low = sum(metrics[town].get("water", {}).get("cod_low", 0) for town in targets)
    cod_high = sum(metrics[town].get("water", {}).get("cod_high", 0) for town in targets)
    nh3_high = sum(metrics[town].get("water", {}).get("nh3_high", 0) for town in targets)
    scope = "、".join(targets)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("各镇村级污水处理设施陆续于2021"):
            set_para_text(
                paragraph,
                f"各镇村级污水处理设施陆续于2021~2023年正式投入商业运营，并满足考核条件。"
                f"本次为2023年下半年度考核，是本项目村级设施的第三次考核，考核范围为{scope}，"
                f"共抽检{sample_count}个农村污水处理设施，考核结果作为已转运营{display_number(operated_count, 0)}个"
                "农村污水处理设施下一周期的付费依据。",
            )
        elif text.startswith("各镇农村污水处理设施绩效考核中"):
            set_para_text(
                paragraph,
                f"本次农村污水处理设施绩效考核共抽检{sample_count}个设施，其中{ge90}个设施考核评分在90分及以上，"
                f"{mid}个设施评分在80~90分之间，{low}个设施评分低于80分。",
            )
        elif text.startswith("二是设施进水CODCr浓度普遍偏低"):
            set_para_text(
                paragraph,
                f"二是部分设施进水水质指标存在异常。本次共核查{water_count}个水质抽检点，其中进水CODCr浓度低于"
                f"100mg/L的设施有{cod_low}个，进水CODCr浓度高于170mg/L的设施有{cod_high}个，进水NH3-N浓度高于"
                f"28mg/L的设施有{nh3_high}个，需持续排查进水来源并强化运行调控。",
            )
        elif text.startswith("2023年12月~2024年1月，考核小组对台山市17个镇共220个"):
            set_para_text(
                paragraph,
                f"2023年12月~2024年1月，考核小组对{scope}共{sample_count}个农村污水处理设施进行了考核，"
                f"其中{ge90}个设施评分在90分及以上，{mid}个设施评分在80~90分之间，{low}个设施评分低于80分。",
            )
        elif text.startswith("本次考核的220个农村污水处理设施中"):
            rate = decimal2(Decimal(out_ok) * Decimal("100") / Decimal(water_count)) if water_count else Decimal("0")
            set_para_text(
                paragraph,
                f"本次共采集{water_count}个农村污水处理设施点进、出水水样，其中{out_ok}个设施点出水水质达标，"
                f"达标率为{display_number(rate, 2)}%；进水CODCr浓度低于100mg/L的设施有{cod_low}个，"
                f"进水CODCr浓度高于170mg/L的设施有{cod_high}个，进水NH3-N浓度高于28mg/L的设施有{nh3_high}个。",
            )


def find_table(doc, required_headers):
    for table in doc.tables:
        if not table.rows:
            continue
        headers = row_texts_xml(table.rows[0])
        joined = "|".join(headers)
        if all(header in joined for header in required_headers):
            return table
    raise ValueError("底稿缺少表格：" + "、".join(required_headers))


def clear_table_data(table):
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)


def populate_schedule_table(doc, targets, metrics):
    table = find_table(doc, ["时间", "第一组", "第二组"])
    clear_table_data(table)
    by_date = defaultdict(list)
    for town in targets:
        for row in metrics[town].get("water_rows", []):
            by_date[row[7]].append((town, row[2], row[3]))
    for date, records in by_date.items():
        descriptions = []
        for town in targets:
            town_rows = [record for record in records if record[0] == town]
            if not town_rows:
                continue
            villages = "，".join(f"{admin}（{natural}）" for _, admin, natural in town_rows)
            descriptions.append(f"{town}（{len(town_rows)}）：{villages}")
        midpoint = (len(descriptions) + 1) // 2
        append_row_from_template(table, [date, "；".join(descriptions[:midpoint]), "；".join(descriptions[midpoint:])])
    total_count = sum(metrics[town]["sample_count"] for town in targets)
    label = "、".join(f"{town}（{metrics[town]['sample_count']}）" for town in targets)
    append_row_from_template(table, ["共计", label, f"共{total_count}个设施"])


def populate_sample_table(doc, targets, metrics):
    table = find_table(doc, ["考核评分", "运维服务费系数Ec1", "可用性服务付费系数Ec2"])
    clear_table_data(table)
    sequence = 1
    for town in targets:
        data = metrics[town]
        for index, (water_row, sample) in enumerate(zip(data["water_rows"], data["sample"])):
            append_row_from_template(table, [
                sequence,
                town,
                water_row[2],
                water_row[3],
                water_row[4],
                water_row[5],
                water_row[6],
                water_row[12],
                display_number(sample["score"], 1),
                display_number(sample["ec1"], 4),
                display_number(data["ec1"], 3) if index == len(data["sample"]) - 1 else "",
                display_number(sample["ec2"], 4),
                display_number(data["ec2"], 3) if index == len(data["sample"]) - 1 else "",
            ])
            sequence += 1


def make_summary_table(doc, data, samples):
    rows = max(len(data["criteria"]), 31)
    table = doc.add_table(rows=rows, cols=4 + len(samples))
    table.style = "Table Grid"
    for row_index in range(rows):
        criteria = data["criteria"][row_index] if row_index < len(data["criteria"]) else ["", "", "", ""]
        for col_index in range(4):
            table.cell(row_index, col_index).text = criteria[col_index]
        for sample_index, sample in enumerate(samples):
            if row_index == 0:
                value = sample["name"]
            elif row_index == rows - 1:
                value = display_number(sample["score"], 1)
            else:
                value = sample["deductions"][row_index] if row_index < len(sample["deductions"]) else ""
            table.cell(row_index, 4 + sample_index).text = value
    if rows:
        score_cell = table.cell(rows - 1, 0)
        for col_index in range(1, 4):
            score_cell = score_cell.merge(table.cell(rows - 1, col_index))
        score_cell.text = "评分"
    set_table_font(table, 7)
    return table


def add_paragraph_element(doc, text, style=None, bold=False):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    return paragraph._element


def build_town_section(doc, town, data):
    water = data.get("water", {})
    satisfaction = data.get("satisfaction", {})
    names = "、".join(sample["name"] for sample in data["sample"])
    issue_count = sum(
        1
        for sample in data["sample"]
        for value in sample["deductions"][1:-1]
        if value not in ("", "/", "0", "0.0")
    )
    elements = [
        add_paragraph_element(doc, f"{town}农村污水处理设施考核情况", "Heading 2"),
        add_paragraph_element(doc, "基本情况", "Heading 3"),
        add_paragraph_element(
            doc,
            f"根据农村污水处理设施清单及本次考核资料，{town}本次共抽检{data['sample_count']}个农村污水处理设施，"
            f"抽检设施包括{names}。考核内容涵盖设施运行、水质达标、维护管理、安全保障、满意度及档案管理等方面。",
        ),
        add_paragraph_element(doc, "农村污水处理设施运维考核情况", "Heading 3"),
        add_paragraph_element(
            doc,
            f"本次考核小组分别对上述{data['sample_count']}个农村污水处理设施进行绩效评价，"
            f"其中{data['score_ge90']}个设施点得分在90分及以上，{data['score_80_90']}个设施点得分在80~90分之间，"
            f"{data['score_lt80']}个设施点得分低于80分。全镇运维期绩效考核运维服务付费系数Ec1为"
            f"{display_number(data['ec1'], 3)}，运维期绩效考核可用性服务付费系数Ec2为{display_number(data['ec2'], 3)}。",
        ),
        add_paragraph_element(doc, "1.产出", bold=True),
        add_paragraph_element(
            doc,
            f"抽检设施基本能正常运行，且设施规模与实际处理水量基本相适应。本次评分表共记录{issue_count}项扣分事项，"
            "项目公司应结合现场检查意见持续做好设施、管网、检查井及机电设备的巡查维护。",
        ),
        add_paragraph_element(doc, "2.效果", bold=True),
        add_paragraph_element(
            doc,
            f"本次共核查{water.get('count', 0)}个水质抽检点，其中{water.get('out_ok', 0)}个出水检测结果达标；"
            f"进水CODCr浓度低于100mg/L的设施有{water.get('cod_low', 0)}个，进水NH3-N浓度高于28mg/L的设施有"
            f"{water.get('nh3_high', 0)}个。村民满意度、镇街满意度和实施机构满意度评分分别为"
            f"{satisfaction.get('public', '未提供')}、{satisfaction.get('town', '未提供')}和{satisfaction.get('agency', '未提供')}。",
        ),
        add_paragraph_element(doc, "3.管理", bold=True),
        add_paragraph_element(
            doc,
            "项目公司提供了运营维护、生产安全、财务及档案管理等资料，建立了相应管理制度。后续应继续完善巡查、维修、"
            "水质检测和群众意见反馈闭环，确保农村污水处理设施稳定运行。",
        ),
        add_paragraph_element(doc, f"{town}农村污水处理设施运行维护绩效评价见下表。"),
    ]
    for start in range(0, len(data["sample"]), 11):
        elements.append(make_summary_table(doc, data, data["sample"][start:start + 11])._element)
    return elements


def insert_town_sections(doc, targets, metrics):
    anchor = find_table(doc, ["考核评分", "运维服务费系数Ec1", "可用性服务付费系数Ec2"])._element
    for town in targets:
        for element in build_town_section(doc, town, metrics[town]):
            anchor.addnext(element)
            anchor = element


def populate_payment_tables(doc, targets, metrics):
    table_specs = [
        (["农村污水处理设施点数", "设计处理规模"], "basis"),
        (["可用性付费基数Pk3", "第九批"], "amount_basis"),
        (["建设期考核系数E1", "运维服务绩效考核系数Ec1"], "coefficients"),
        (["可用性付费基数（元/月）Pk3", "可用性付费(元/月)"], "availability"),
        (["第一批运维服务费(元/月)", "运维服务费合计"], "om"),
        (["第八批运维服务费(元/月)", "第九批运维服务费(元/月)"], "batch89"),
        (["每月合计扣减费用", "第一~七批运维服务费"], "deductions"),
    ]
    tables = {}
    for headers, key in table_specs:
        table = find_table(doc, headers)
        clear_table_data(table)
        tables[key] = table

    for sequence, town in enumerate(targets, 1):
        data = metrics[town]
        basis = data["basis"]
        pay = data["pay"]
        batches = basis["batches"]
        append_row_from_template(tables["basis"], [
            sequence, town, basis["facility_count"], basis["design_scale"],
            display_number(basis["pk3"]), display_number(basis["py3"]),
        ])
        append_row_from_template(tables["amount_basis"], [
            sequence, town, display_number(basis["pk3"]), display_number(basis["py3"]),
            *[display_number(value) for value in batches],
        ])
        append_row_from_template(tables["coefficients"], [
            sequence, town, display_number(basis["e1"], 3), display_number(data["ec1"], 3), display_number(data["ec2"], 3),
        ])
        append_row_from_template(tables["availability"], [
            sequence, town, display_number(basis["pk3"]), display_number(pay["availability_base"]),
            display_number(basis["e1"], 3), display_number(data["ec2"], 3), display_number(pay["availability"]),
        ])
        calculated_batches = [decimal2(value * data["ec1"]) for value in batches[:7]]
        append_row_from_template(tables["om"], [
            sequence, town, display_number(data["ec1"], 3),
            *[display_number(value) for value in calculated_batches], display_number(pay["om"]),
        ])
        append_row_from_template(tables["batch89"], [
            sequence, town, "1", display_number(batches[7]), display_number(data["ec1"], 3),
            display_number(pay["batch8"]), display_number(batches[8]), "",
        ])
        append_row_from_template(tables["deductions"], [
            sequence, town,
            display_number(pay["availability_base"]), display_number(pay["availability"]), display_number(pay["availability_deduct"]),
            display_number(pay["om_base"]), display_number(pay["om"]), display_number(pay["om_deduct"]),
            display_number(pay["batch8_deduct"]), display_number(pay["total_deduct"]),
        ])


def replace_problem_sections(doc, targets, metrics):
    section_texts = {
        "环境整治效果方面": (
            f"本次对{'、'.join(targets)}共{sum(metrics[t]['sample_count'] for t in targets)}个抽检设施开展现场考核。"
            "部分设施点仍需加强周边环境、格栅集水井、构筑物及标识标牌的日常维护，及时完成现场问题整改。"
        ),
        "进、出水水质方面": (
            f"本次共核查{sum(metrics[t].get('water', {}).get('count', 0) for t in targets)}个水质抽检点，"
            f"其中{sum(metrics[t].get('water', {}).get('out_ok', 0) for t in targets)}个出水检测结果达标；"
            f"进水CODCr浓度低于100mg/L的设施有{sum(metrics[t].get('water', {}).get('cod_low', 0) for t in targets)}个，"
            f"进水NH3-N浓度高于28mg/L的设施有{sum(metrics[t].get('water', {}).get('nh3_high', 0) for t in targets)}个。"
            "项目公司应持续分析异常进水原因，强化工艺运行调控和水质检测。"
        ),
        "污水收集系统方面": (
            "部分抽检设施仍存在污水收集不充分、管渠排水不畅、检查井或截流井维护不到位等问题。"
            "建议结合评分表和现场记录逐项排查，提升污水有效收集率。"
        ),
        "污水处理系统方面": (
            "部分设施的预处理单元、提升泵、曝气设备及控制系统仍需加强巡查保养。"
            "项目公司应根据设施工艺和运行负荷落实针对性维护，确保处理系统稳定运行。"
        ),
        "安全管理措施方面": (
            "应持续检查防护栏、安全警示标识、信息公示牌、电气设备和有限空间作业管理情况，"
            "对发现的安全隐患建立清单并及时闭环整改。"
        ),
        "组织管理方面": (
            "项目公司已建立相应运维管理制度，后续仍需规范巡查、维修、水质检测和档案记录，"
            "提高问题响应和整改闭环效率。"
        ),
        "社会服务评价方面": (
            "各镇街应结合村民、镇级主管部门及实施机构满意度评价结果，持续改进污水收集、设施维护和信息公开工作，"
            "及时回应群众意见。"
        ),
    }
    paragraphs = list(doc.paragraphs)
    heading_indexes = [index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() in section_texts]
    for index in reversed(heading_indexes):
        heading = paragraphs[index]
        end = len(paragraphs)
        for candidate in range(index + 1, len(paragraphs)):
            text = paragraphs[candidate].text.strip()
            style = paragraphs[candidate].style.name if paragraphs[candidate].style else ""
            if text in section_texts or text == "项目公司整改建议" or style in {"Heading 1", "Heading 2", "Heading 3"}:
                end = candidate
                break
        for paragraph in paragraphs[index + 1:end]:
            parent = paragraph._element.getparent()
            if parent is not None:
                parent.remove(paragraph._element)
        replacement = doc.add_paragraph(section_texts[heading.text.strip()])
        heading._element.addnext(replacement._element)


def scrub_template_placeholders(doc, targets):
    placeholders = ("本段根据后续提供", "不保留既有镇街数据", "目标镇街材料填写")
    previous_text = ""
    scope = "、".join(targets)
    for paragraph in doc.paragraphs:
        if any(marker in paragraph.text for marker in placeholders):
            if previous_text == "考核安排":
                replacement = f"本次对{scope}开展现场考核，具体时间及分组安排见下表。"
            elif previous_text == "付费范围":
                replacement = f"2023年下半年度付费范围包括{scope}已转入运营并满足合同考核条件的农村污水处理设施。"
            elif "表3-1" in previous_text:
                replacement = "农村污水处理设施项目服务费基数见下表。"
            elif previous_text == "二、考核工作开展情况":
                replacement = f"考核小组按照合同约定对{scope}开展现场检查、资料查阅、问卷调查和水质检测，并形成考核评分结果。"
            else:
                replacement = ""
            set_para_text(paragraph, replacement)
        if paragraph.text.strip():
            previous_text = paragraph.text.strip()


def generate_report(targets, metrics, filename):
    out = FINAL_DIR / filename
    doc = Document(SOURCE)
    update_toc(doc, targets)
    populate_schedule_table(doc, targets, metrics)
    populate_sample_table(doc, targets, metrics)
    insert_town_sections(doc, targets, metrics)
    populate_payment_tables(doc, targets, metrics)
    if len(targets) == 1:
        target = targets[0]
        replace_overview_paragraphs(doc, target, metrics)
        scrub_cross_town_examples(doc, target)
    replace_global_metrics(doc, targets, metrics)
    replace_problem_sections(doc, targets, metrics)
    scrub_template_placeholders(doc, targets)
    apply_font_rules(doc)
    doc.save(out)
    return out


def generate_town_report(target, metrics):
    return generate_report([target], metrics, f"{target}2023年下半年度村级设施考核报告（正文）.docx")


def delete_old_baisha_chixi_outputs():
    for name in [
        "白沙镇第二章高仿原文正式版.docx",
        "赤溪镇第二章高仿原文正式版.docx",
    ]:
        path = os.path.join(OUTPUTS, name)
        if os.path.exists(path):
            os.remove(path)


def copy_existing_or_generate(town, metrics):
    return generate_town_report(town, metrics)


def structural_check(paths):
    rows = []
    for path in paths:
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        town = os.path.basename(path).split("2023年")[0]
        if town == "台山市":
            ok = all(t in text for t in TOWNS)
        else:
            unrelated = [t for t in TOWNS if t != town and f"{t}农村污水处理设施考核情况" in text]
            ok = (f"{town}农村污水处理设施考核情况" in text) and not unrelated
        rows.append((os.path.basename(path), ok, len(doc.tables)))
    return rows


def main():
    os.makedirs(FINAL_DIR, exist_ok=True)
    delete_old_baisha_chixi_outputs()
    common_material = validate_common_material()
    metrics = collect_metrics()
    attach_payments(metrics, load_amount_basis())

    generated = []
    for town in TOWNS:
        generated.append(copy_existing_or_generate(town, metrics))

    generated.append(generate_report(
        TOWNS,
        metrics,
        "台山市2023年下半年度村级设施考核报告（正文）.docx",
    ))

    checks = structural_check(generated)
    print(f"final_dir={FINAL_DIR}")
    print(f"data_dir={DATA_DIR}")
    print(f"common_material={common_material}")
    print(f"generated_count={len(generated)}")
    for name, ok, table_count in checks:
        print(f"{name}\tOK={ok}\ttables={table_count}")


if __name__ == "__main__":
    main()
