import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


FORBIDDEN_LABELS = ["套话", "框架", "合并版", "技能", "高仿", "专项", "保留版", "生成版"]
PLACEHOLDER_TEXTS = [
    "本段根据后续提供", "不保留既有镇街数据", "目标镇街材料填写",
    "约73%（159个）", "本次考核的220个农村污水处理设施中",
]
REPORT_NAME_RE = re.compile(r"^(.+?)2023年下半年度村级设施考核报告（正文）\.docx$")
SECTION_RE = re.compile(r"([^，。、；\s]+(?:镇|街道|乡|村|区|县|市))农村污水处理设施考核情况")


def clean(text):
    return "".join(str(text or "").split())


def element_text(element):
    return "".join(t.text or "" for t in element.iter() if t.tag.endswith("}t"))


def doc_text(doc):
    return "\n".join(element_text(element) for element in doc._body._element)


def row_values(row):
    return [clean(element_text(tc)) for tc in row._tr.tc_lst]


def font_issues(doc):
    issues = []
    for run in doc._element.iter():
        if not run.tag.endswith("}r"):
            continue
        text = element_text(run).strip()
        if not text:
            continue
        rpr = run.find(qn("w:rPr"))
        rfonts = None if rpr is None else rpr.find(qn("w:rFonts"))
        east = None if rfonts is None else rfonts.get(qn("w:eastAsia"))
        if east not in (None, "宋体"):
            issues.append(f"非宋体文字：{text[:20]}")
            if len(issues) >= 10:
                break
    return issues


def sequence_issues(doc):
    issues = []
    for table_index, table in enumerate(doc.tables):
        header_index = None
        for i, row in enumerate(table.rows[:3]):
            vals = row_values(row)
            if vals and vals[0] == "序号":
                header_index = i
                break
        if header_index is None:
            continue
        expected = 1
        for row_index, row in enumerate(table.rows[header_index + 1 :], start=header_index + 1):
            vals = row_values(row)
            if not vals:
                continue
            first = vals[0]
            if first in {"", "合计", "共计", "总计"} or not re.fullmatch(r"\d+", first):
                continue
            if int(first) != expected:
                issues.append(f"表{table_index + 1}第{row_index + 1}行序号为{first}，应为{expected}")
                break
            expected += 1
    return issues


def report_target(path):
    match = REPORT_NAME_RE.match(path.name)
    return match.group(1) if match else path.name.split("2023年")[0]


def unique_sorted(names):
    return sorted(dict.fromkeys(name for name in names if name))


def parse_towns(value):
    if not value:
        return []
    towns = [item.strip() for item in re.split(r"[、,，\s]+", value) if item.strip()]
    return unique_sorted(towns)


def infer_towns_from_data_dir(data_dir):
    if not data_dir or not data_dir.is_dir():
        return []
    towns = []
    for child in data_dir.iterdir():
        if not child.is_dir():
            continue
        town = child.name
        if (child / f"{town}附件资料.docx").is_file():
            towns.append(town)
    return unique_sorted(towns)


def infer_uploaded_towns(files, explicit_towns, data_dir=None):
    if explicit_towns:
        return explicit_towns
    data_towns = infer_towns_from_data_dir(data_dir)
    if data_towns:
        return data_towns
    return unique_sorted(report_target(file) for file in files if REPORT_NAME_RE.match(file.name))


def section_towns(text):
    return unique_sorted(match.group(1) for match in SECTION_RE.finditer(text))


def validate_file(path, expected_towns=None, total_report=False):
    result = []
    name = path.name
    if any(label in name for label in FORBIDDEN_LABELS):
        result.append("文件名含制作痕迹")
    if not REPORT_NAME_RE.match(name):
        result.append("文件名不符合标准格式")

    doc = Document(str(path))
    text = doc_text(doc)
    target = report_target(path)
    expected_towns = expected_towns or []
    if total_report:
        missing_towns = [town for town in expected_towns if f"{town}农村污水处理设施考核情况" not in text]
        if missing_towns:
            result.append("总报告缺少镇街章节：" + "、".join(missing_towns))
        extra_towns = [town for town in section_towns(text) if town not in expected_towns]
        if extra_towns:
            result.append("总报告含非本次上传镇街章节：" + "、".join(extra_towns))
    elif target in expected_towns:
        others = [town for town in expected_towns if town != target and f"{town}农村污水处理设施考核情况" in text]
        if others:
            result.append("单镇报告含其他镇名：" + "、".join(others))
        if f"{target}农村污水处理设施考核情况" not in text:
            result.append("缺少本镇考核章节")

    body_labels = [label for label in FORBIDDEN_LABELS if label in text]
    if body_labels:
        result.append("正文含制作痕迹：" + "、".join(body_labels))
    placeholders = [label for label in PLACEHOLDER_TEXTS if label in text]
    if placeholders:
        result.append("正文含底稿占位文字：" + "、".join(placeholders))

    sample_tables = []
    summary_tables = []
    payment_tables = []
    for table in doc.tables:
        if not table.rows:
            continue
        header = "|".join(row_values(table.rows[0]))
        if "考核评分" in header and "运维服务费系数Ec1" in header:
            sample_tables.append(table)
        if "一级指标" in header and "评价方法" in header and len(table.columns) > 4:
            summary_tables.append(table)
        if "每月合计扣减费用" in header:
            payment_tables.append(table)
    if not sample_tables or len(sample_tables[0].rows) <= 1:
        result.append("考核评分汇总表未填充")
    if not summary_tables:
        result.append("缺少设施绩效评价汇总表")
    if not payment_tables or len(payment_tables[0].rows) <= 1:
        result.append("绩效付费扣减表未填充")

    result.extend(sequence_issues(doc))
    result.extend(font_issues(doc))
    return result


def main():
    parser = argparse.ArgumentParser()
    package_root = Path(__file__).resolve().parents[3]
    project_root = package_root.parent if (package_root.parent / "资料收集").is_dir() else package_root
    parser.add_argument("path", nargs="?", default=str(project_root / "生成"))
    parser.add_argument(
        "--towns",
        default="",
        help="本次上传/生成的镇街，用逗号或顿号分隔；不填时从输出目录里的镇街报告文件名自动识别。",
    )
    parser.add_argument(
        "--data-dir",
        default=str(project_root / "资料收集"),
        help="本次上传资料目录；不显式传 --towns 时，用其中的镇街附件资料识别本次上传镇街。",
    )
    args = parser.parse_args()

    root = Path(args.path)
    files = [root] if root.is_file() else sorted(root.glob("*.docx"))
    expected_towns = infer_uploaded_towns(files, parse_towns(args.towns), Path(args.data_dir))
    bad = []
    if root.is_dir():
        expected = {f"{town}2023年下半年度村级设施考核报告（正文）.docx" for town in expected_towns}
        actual = {file.name for file in files}
        missing = sorted(expected - actual)
        total_files = [
            file for file in files
            if REPORT_NAME_RE.match(file.name) and report_target(file) not in expected_towns
        ]
        extra = sorted(
            name for name in actual - expected
            if name not in {file.name for file in total_files}
        )
        if len(expected_towns) > 1 and not total_files:
            missing.append("总报告（文件名前缀按本项目名称确定）")
        if len(expected_towns) <= 1 and total_files:
            extra.extend(file.name for file in total_files)
        if len(expected_towns) > 1 and len(total_files) > 1:
            extra.extend(file.name for file in total_files[1:])
        if missing:
            bad.append(("输出目录", ["缺少报告：" + "、".join(missing)]))
        if extra:
            bad.append(("输出目录", ["存在非标准报告：" + "、".join(extra)]))
    for file in files:
        is_total = root.is_dir() and len(expected_towns) > 1 and report_target(file) not in expected_towns
        issues = validate_file(file, expected_towns, is_total)
        if issues:
            bad.append((file.name, issues))

    print(f"checked={len(files)}")
    print("towns=" + ("、".join(expected_towns) if expected_towns else "未识别"))
    print(f"bad={len(bad)}")
    for name, issues in bad:
        print(name)
        for issue in issues:
            print(f"  - {issue}")
    if bad:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
