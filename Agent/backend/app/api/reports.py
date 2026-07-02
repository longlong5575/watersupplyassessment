from pathlib import Path
import os
import subprocess
import sys

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from docx import Document
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.database import get_session
from app.core.security import admin_user
from app.models import AssessmentCycle, AssessmentRecord, City, Report, ReportTask, Town
from app.schemas import ReportTaskRequest
from app.services.report_dataset import build_report_dataset, validate_report_dataset
from app.services.report_tasks import run_report_task


router = APIRouter(tags=["reports"])


def serialize_report(item: Report) -> dict:
    town_name = None
    try:
        town_name = item.town.name if getattr(item, "town", None) else None
    except Exception:
        town_name = None
    return {
        "id": item.id,
        "taskId": item.task_id,
        "name": item.name,
        "status": item.status,
        "size": item.size,
        "createdAt": item.created_at.isoformat(),
        "town": town_name or item.town_id,
        "cycleId": item.cycle_id,
        "version": item.version,
        "format": item.format,
        "datasetHash": item.dataset_hash,
        "recordIds": item.data_snapshot.get("recordIds", []),
        "indicatorVersionIds": item.data_snapshot.get("indicatorVersionIds", []),
    }


def serialize_task(item: ReportTask, reports: list[Report] | None = None) -> dict:
    return {
        "id": item.id,
        "status": item.status,
        "progress": item.progress,
        "error": item.error,
        "createdAt": item.created_at.isoformat(),
        "updatedAt": item.updated_at.isoformat(),
        "startedAt": item.started_at.isoformat() if item.started_at else None,
        "completedAt": item.completed_at.isoformat() if item.completed_at else None,
        "payload": item.payload,
        "datasetHash": item.dataset_hash,
        "dataSnapshot": {
            "cycleId": item.data_snapshot.get("cycleId"),
            "cycleName": item.data_snapshot.get("cycleName"),
            "requestedTowns": item.data_snapshot.get("requestedTowns", []),
            "towns": item.data_snapshot.get("towns", []),
            "recordIds": item.data_snapshot.get("recordIds", []),
            "indicatorVersionIds": item.data_snapshot.get("indicatorVersionIds", []),
        },
        "reports": [serialize_report(report) for report in (reports or [])],
    }


def docx_preview(path: Path) -> dict:
    try:
        document = Document(path)
    except Exception as exc:
        raise HTTPException(status_code=422, detail="Report preview failed") from exc

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables[:8]:
        rows = []
        for row in table.rows[:12]:
            rows.append([cell.text.strip() for cell in row.cells[:6]])
        if rows:
            tables.append(rows)

    return {
        "paragraphs": paragraphs[:120],
        "tables": tables,
        "paragraphCount": len(paragraphs),
        "tableCount": len(document.tables),
    }


def resolve_report_request(payload: ReportTaskRequest, session: Session) -> tuple[City | None, AssessmentCycle | None, dict]:
    project = session.get(City, payload.projectId) if payload.projectId else None
    cycle = None
    if payload.source == "dashboard":
        cycle_query = (
            select(AssessmentCycle)
            .join(AssessmentRecord, AssessmentRecord.cycle_id == AssessmentCycle.id)
            .where(AssessmentRecord.status.in_(["submitted", "reviewed", "locked"]))
            .order_by(AssessmentRecord.updated_at.desc())
        )
        if project is not None:
            cycle_query = cycle_query.where(AssessmentRecord.city_id == project.id)
        if payload.townNames:
            town_ids = select(Town.id).where(Town.name.in_(payload.townNames))
            if project is not None:
                town_ids = town_ids.where(Town.city_id == project.id)
            cycle_query = cycle_query.where(AssessmentRecord.town_id.in_(town_ids))
        cycle = session.scalar(cycle_query)
    elif payload.period.strip():
        cycle_query = select(AssessmentCycle).where(AssessmentCycle.name == payload.period.strip())
        if project is not None:
            cycle_query = cycle_query.where(AssessmentCycle.city_id == project.id)
        cycle = session.scalar(cycle_query)
    if cycle is None:
        fallback_query = select(AssessmentCycle).where(AssessmentCycle.status == "active")
        if project is not None:
            fallback_query = fallback_query.where(AssessmentCycle.city_id == project.id)
        cycle = session.scalar(fallback_query)
    if payload.projectId and project is None:
        raise HTTPException(status_code=422, detail="Project not found")
    task_payload = payload.model_dump()
    task_payload["period"] = cycle.name if cycle is not None else payload.period.strip()
    if payload.townIds:
        towns = session.scalars(select(Town).where(Town.id.in_(payload.townIds))).all()
        if project and any(town.city_id != project.id for town in towns):
            raise HTTPException(status_code=422, detail="Selected towns do not belong to the project")
        task_payload["townNames"] = sorted({*payload.townNames, *(town.name for town in towns)})
    elif project and payload.townNames:
        valid_names = set(session.scalars(select(Town.name).where(Town.city_id == project.id, Town.name.in_(payload.townNames))).all())
        if valid_names != set(payload.townNames):
            raise HTTPException(status_code=422, detail="Selected towns do not belong to the project")
    return project, cycle, task_payload


def build_task_snapshot(session: Session, project: City | None, cycle: AssessmentCycle | None, task_payload: dict) -> dict:
    town_names = set(task_payload.get("townNames", []) or []) or None
    return build_report_dataset(session, cycle=cycle, town_names=town_names, city_id=project.id if project else None)


@router.post("/api/report-tasks/precheck")
def precheck_task(payload: ReportTaskRequest, session: Session = Depends(get_session), user=Depends(admin_user)):
    project, cycle, task_payload = resolve_report_request(payload, session)
    snapshot = build_task_snapshot(session, project, cycle, task_payload)
    errors: list[str] = []
    warnings: list[str] = []
    try:
        validate_report_dataset(snapshot)
    except RuntimeError as exc:
        errors.append(str(exc))
    requested = set(task_payload.get("townNames", []) or [])
    available = {item.get("town") for item in snapshot.get("towns", [])}
    missing = sorted(name for name in requested if name not in available)
    if missing:
        errors.append("No reviewed or locked data for: " + ", ".join(missing))
    for town in snapshot.get("towns", []):
        if town.get("waterQualityCount", 0) <= 0:
            warnings.append(f"{town.get('town')} has no water quality record.")
        if town.get("surveyCount", 0) <= 0:
            warnings.append(f"{town.get('town')} has no survey record.")
        if town.get("attachmentCount", 0) <= 0:
            warnings.append(f"{town.get('town')} has no attachment.")
    return {
        "ok": not errors,
        "errors": errors,
        "warnings": warnings,
        "summary": {
            "projectName": snapshot.get("projectName"),
            "cycleName": snapshot.get("cycleName"),
            "requestedTownCount": len(requested),
            "availableTownCount": len(snapshot.get("towns", [])),
            "recordCount": len(snapshot.get("records", [])),
            "recordIds": snapshot.get("recordIds", []),
            "indicatorVersionIds": snapshot.get("indicatorVersionIds", []),
            "datasetHash": snapshot.get("hash"),
        },
        "towns": snapshot.get("towns", []),
    }


@router.post("/api/report-tasks")
def create_task(payload: ReportTaskRequest, session: Session = Depends(get_session), user=Depends(admin_user)):
    project, cycle, task_payload = resolve_report_request(payload, session)
    snapshot = build_task_snapshot(session, project, cycle, task_payload)
    if task_payload.get("source") == "dashboard":
        try:
            validate_report_dataset(snapshot)
        except RuntimeError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
    task = ReportTask(
        cycle_id=cycle.id if cycle else None,
        created_by_id=user.id,
        payload=task_payload,
        data_snapshot=snapshot,
        dataset_hash=snapshot.get("hash"),
    )
    session.add(task)
    session.commit()
    # Redis/Celery is the production path; local development remains usable without a worker.
    if settings.celery_task_always_eager:
        run_report_task(task.id)
    else:
        from app.workers.tasks import generate_report

        try:
            generate_report.delay(task.id)
        except Exception:
            run_report_task(task.id)
    session.refresh(task)
    return {"id": task.id, "status": task.status, "progress": task.progress, "datasetHash": task.dataset_hash, "reports": []}


@router.get("/api/report-tasks/{task_id}")
def get_task(task_id: str, session: Session = Depends(get_session)):
    task = session.get(ReportTask, task_id)
    if task is None: raise HTTPException(status_code=404, detail="Report task not found")
    reports = session.scalars(select(Report).where(Report.task_id == task.id)).all()
    return serialize_task(task, list(reports))


@router.get("/api/report-tasks")
def list_tasks(session: Session = Depends(get_session), status: str | None = None):
    query = select(ReportTask).order_by(ReportTask.created_at.desc())
    if status:
        query = query.where(ReportTask.status == status)
    tasks = list(session.scalars(query).all())
    return {"items": [serialize_task(item) for item in tasks]}


@router.get("/api/reports")
def reports(session: Session = Depends(get_session)):
    return {"items": [serialize_report(item) for item in session.scalars(select(Report).order_by(Report.created_at.desc())).all()]}


@router.get("/api/reports/{report_id}/download")
def download(report_id: str, session: Session = Depends(get_session)):
    report = session.get(Report, report_id)
    if report is None: raise HTTPException(status_code=404, detail="Report not found")
    path = Path(report.storage_key)
    if not path.is_file(): raise HTTPException(status_code=404, detail="Report file is missing")
    return FileResponse(path, filename=path.name)


@router.get("/api/reports/{report_id}/preview")
def preview(report_id: str, session: Session = Depends(get_session)):
    report = session.get(Report, report_id)
    if report is None: raise HTTPException(status_code=404, detail="Report not found")
    path = Path(report.storage_key)
    if not path.is_file(): raise HTTPException(status_code=404, detail="Report file is missing")
    if path.suffix.lower() != ".docx": raise HTTPException(status_code=415, detail="Only DOCX reports can be previewed")
    return {"report": serialize_report(report), "content": docx_preview(path)}


@router.post("/api/reports/{report_id}/open-folder")
def open_report_folder(report_id: str, session: Session = Depends(get_session)):
    report = session.get(Report, report_id)
    if report is None: raise HTTPException(status_code=404, detail="Report not found")
    path = Path(report.storage_key)
    if not path.is_file(): raise HTTPException(status_code=404, detail="Report file is missing")
    folder = path.parent
    try:
        if sys.platform.startswith("win"):
            os.startfile(str(folder))  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(folder)])
        else:
            subprocess.Popen(["xdg-open", str(folder)])
    except Exception as exc:
        raise HTTPException(status_code=500, detail="Report folder could not be opened") from exc
    return {"ok": True, "folder": str(folder)}
