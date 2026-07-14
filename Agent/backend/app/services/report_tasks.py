import shutil
import traceback
import json
import re
from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Any

from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.database import SessionLocal
from app.core.config import settings
from app.models import AssessmentCycle, AssessmentRecord, Attachment, Report, ReportTask, SurveyRecord, Town, WaterQualityRecord
from app.models.entities import utcnow
from app.services.payment_basis import (
    maonan_payment_basis_for_point,
    maonan_payment_basis_rows,
    payment_source_summary,
    yunan_county_network_basis,
    yunan_rural_payment_basis_rows,
    yunan_town_payment_basis_for_point,
    yunan_town_payment_basis_rows,
)
from app.services.report_dataset import build_report_dataset, validate_report_dataset
from app.services.payment import (
    bounded_monthly_volume,
    maonan_annual_maximum_treatment_fee,
    maonan_network_monthly_fee,
    maonan_operation_coefficient,
    maonan_treatment_monthly_fee,
    maonan_water_quality_coefficient,
    town_average_coefficient,
    yunan_dry_season_quality_coefficient,
    yunan_network_monthly_fee,
    yunan_operation_coefficient,
    yunan_town_network_load_coefficient,
    yunan_town_treatment_monthly_fee,
    yunan_water_quality_coefficient,
)


REPORT_TYPE_LABELS = {
    "town_plant": "镇街污水处理厂",
    "town_network": "镇街污水收集管网",
    "rural_treatment": "农村污水处理设施",
}
STATUS_LABELS = {
    "reviewed": "已复核",
    "locked": "已锁定",
    "submitted": "已提交",
    "draft": "草稿",
}
SURVEY_TYPE_LABELS = {
    "satisfaction": "满意度调查",
    "sewage_collection": "污水收集效果调查",
    "overall_effect": "整体效果调查",
}
RESPONDENT_LABELS = {
    "villager1": "村民代表一",
    "villager2": "村民代表二",
    "gov_rep": "镇街代表",
    "assessment_team": "考核小组",
    "implementation_org": "实施机构",
}

REPORT_BODY_FONT = "宋体"
REPORT_HEADING_FONT = "宋体"
REPORT_LATIN_FONT = "宋体"
REPORT_BODY_SIZE_PT = 10.5
REPORT_TABLE_SIZE_PT = 9
REPORT_BODY_LINE_SPACING_PT = 24


def _normalize_report_text(value: Any) -> str:
    """Normalize legacy source units and punctuation at the report boundary."""
    text = "" if value is None else str(value)
    text = re.sub(r"(?i)m\s*(?:\^\s*)?3", "立方米", text)
    text = re.sub(r"(?i)立方米\s*/\s*d\b", "立方米/日", text)
    text = re.sub(r"(?i)(?<=\d)\s*km\b", "公里", text)
    text = re.sub(r"(?<=\d)\s*立方米", "立方米", text)
    text = text.replace("~", "至")
    return text


def _set_rfonts(r_pr, font_name: str = REPORT_BODY_FONT) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), font_name)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{attribute}"), None)


def _apply_run_font(run, font_name: str, size_pt: float, *, bold: bool | None = None) -> None:
    from docx.shared import Pt

    run.font.name = font_name
    _set_rfonts(run._element.get_or_add_rPr(), font_name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _apply_paragraph_format(paragraph, *, indent: bool = False, table: bool = False) -> None:
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(16 if table else REPORT_BODY_LINE_SPACING_PT)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if indent:
        fmt.first_line_indent = Pt(21)


def _set_style_font(style, font_name: str, size_pt: float, *, bold: bool = False) -> None:
    from docx.shared import Pt, RGBColor

    style.font.name = font_name
    _set_rfonts(style._element.get_or_add_rPr(), font_name)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def _enforce_document_fonts(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    styles = document.styles._element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.insert(0, r_pr_default)
    default_r_pr = r_pr_default.find(qn("w:rPr"))
    if default_r_pr is None:
        default_r_pr = OxmlElement("w:rPr")
        r_pr_default.append(default_r_pr)
    _set_rfonts(default_r_pr)

    for style in styles.findall(qn("w:style")):
        style_r_pr = style.find(qn("w:rPr"))
        if style_r_pr is None:
            style_r_pr = OxmlElement("w:rPr")
            style.append(style_r_pr)
        _set_rfonts(style_r_pr)

    for part in document.part.package.parts:
        root = getattr(part, "_element", None)
        if root is None:
            continue
        for text_node in root.iter(qn("w:t")):
            if text_node.text:
                text_node.text = _normalize_report_text(text_node.text)
        for run in root.iter(qn("w:r")):
            run_r_pr = run.find(qn("w:rPr"))
            if run_r_pr is None:
                run_r_pr = OxmlElement("w:rPr")
                run.insert(0, run_r_pr)
            _set_rfonts(run_r_pr)


def _fmt_score_value(value: Any) -> str:
    try:
        number = float(value or 0)
    except (TypeError, ValueError):
        return str(value or "0")
    if abs(number - round(number)) < 0.0001:
        return str(int(round(number)))
    return f"{number:.2f}".rstrip("0").rstrip(".")

PROJECT_REPORT_PROFILES = {
    "郁南项目": {
        "shortName": "郁南",
        "titleSuffix": "镇村污水处理设施绩效考核报告",
        "basis": "依据《PPP项目合同》、补充协议及郁南项目镇村考核报告确定的考核频次、考核对象、评分标准和水质限值执行。",
        "methods": ["现场检查", "查阅资料", "问卷调查（村级考核有）", "水质检测", "评分复核"],
        "waterStandard": "镇级污水处理设施出水执行城镇一级A及广东省地方标准较严值；农村生活污水处理设施执行广东省《农村生活污水处理排放标准》（DB44/2208-2019）。",
        "waterRows": [
            ["镇级污水厂及管网", "CODCr", "40", "mg/L"],
            ["镇级污水厂及管网", "NH3-N", "5（8）", "mg/L"],
            ["镇级污水厂及管网", "TP", "0.5", "mg/L"],
            ["农村污水处理设施", "CODCr", "60", "mg/L"],
            ["农村污水处理设施", "NH3-N", "8（15）", "mg/L"],
            ["农村污水处理设施", "TP", "1", "mg/L"],
        ],
        "hasSurvey": True,
    },
    "茂南项目": {
        "shortName": "茂南",
        "titleSuffix": "城镇设施绩效考核报告",
        "basis": "依据茂南区水质净化处理设施全区捆绑PPP项目城镇设施绩效考核报告口径，对水质净化厂及配套管网分别开展周期绩效考核。",
        "methods": ["现场检查", "查阅资料", "水质检测", "评分复核"],
        "waterStandard": "城镇污水处理设施出水执行《城镇污水处理厂污染物排放标准》一级A标准及广东省《水污染物排放限值》较严值。",
        "waterRows": [
            ["城镇污水处理设施", "CODCr", "40", "mg/L"],
            ["城镇污水处理设施", "BOD5", "10", "mg/L"],
            ["城镇污水处理设施", "SS", "10", "mg/L"],
            ["城镇污水处理设施", "NH3-N", "5（8）", "mg/L"],
            ["城镇污水处理设施", "TP", "0.5", "mg/L"],
        ],
        "hasSurvey": False,
    },
}


def _set_cell_text(cell, value: Any, *, bold: bool = False) -> None:
    text = _normalize_report_text(value)
    cell.text = text
    for paragraph in cell.paragraphs:
        _apply_paragraph_format(paragraph, table=True)
        for run in paragraph.runs:
            _apply_run_font(run, REPORT_BODY_FONT, REPORT_TABLE_SIZE_PT, bold=bold)


def _mark_table_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def _prepare_document(title: str):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    section.different_first_page_header_footer = True
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, cached, end])
    _apply_run_font(run, REPORT_BODY_FONT, 9)
    normal = document.styles["Normal"]
    _set_style_font(normal, REPORT_BODY_FONT, REPORT_BODY_SIZE_PT)
    normal.paragraph_format.line_spacing = Pt(REPORT_BODY_LINE_SPACING_PT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = document.styles[style_name]
        _set_style_font(style, REPORT_HEADING_FONT, size, bold=True)
        style.paragraph_format.line_spacing = Pt(REPORT_BODY_LINE_SPACING_PT)
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6)
    return document


def _add_assessment_object(document, town_data: dict[str, Any], records: list[dict[str, Any]]) -> None:
    section_code = (town_data.get("reportTemplate") or {}).get("assessmentObjectSection") or town_data.get("chapterCode") or "1"
    document.add_heading(f"{section_code} 考核对象", level=1)
    objects = town_data.get("assessmentObject") or {}
    for facility_type in town_data.get("assessmentTargets") or []:
        item = objects.get(facility_type) or {}
        document.add_heading(item.get("title") or REPORT_TYPE_LABELS.get(facility_type, facility_type), level=2)
        _add_paragraph(document, item.get("description") or "本次考核对象及其基本情况以项目资料和现场核查结果为准。")

    villages: dict[str, dict[str, Any]] = {}
    for record in records:
        if record.get("villageId"):
            villages[record["villageId"]] = record
    if villages:
        document.add_heading("农村设施点清单", level=2)
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for cell, value in zip(table.rows[0].cells, ["序号", "行政村", "设施点", "章节号", "考核对象"]):
            _set_cell_text(cell, value, bold=True)
        _mark_table_header(table.rows[0])
        for index, record in enumerate(sorted(villages.values(), key=lambda item: (item.get("villageChapterCode") or "", item.get("village") or "")), 1):
            row = table.add_row().cells
            obj = record.get("villageAssessmentObject") or {}
            for cell, value in zip(row, [index, record.get("administrativeVillage"), record.get("village"), record.get("villageChapterCode"), obj.get("title") or obj.get("description")]):
                _set_cell_text(cell, value)


def _add_paragraph(document, text: str, *, bold_prefix: str | None = None, indent: bool = True, size_pt: float | None = None) -> None:
    text = _normalize_report_text(text)
    bold_prefix = _normalize_report_text(bold_prefix) if bold_prefix else None
    paragraph = document.add_paragraph()
    _apply_paragraph_format(paragraph, indent=indent)
    font_size = size_pt or REPORT_BODY_SIZE_PT
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        _apply_run_font(run, REPORT_BODY_FONT, font_size, bold=True)
        text = text[len(bold_prefix):]
    run = paragraph.add_run(text)
    _apply_run_font(run, REPORT_BODY_FONT, font_size, bold=False)


def _add_simple_table(document, headers: list[str], rows: list[list[Any]]) -> None:
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    def keep_row_together(row) -> None:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, bold=True)
    _mark_table_header(table.rows[0])
    for row_values in rows:
        table_row = table.add_row()
        keep_row_together(table_row)
        for cell, value in zip(table_row.cells, row_values):
            _set_cell_text(cell, value)
    gap = document.add_paragraph()
    _apply_paragraph_format(gap, indent=False)
    gap.paragraph_format.line_spacing = Pt(4)


def _facility_types(records: list[dict[str, Any]]) -> list[str]:
    order = ["town_plant", "town_network", "rural_treatment"]
    present = {record.get("facilityType") or record.get("rawFacilityType") for record in records}
    return [item for item in order if item in present]


def _record_score_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for index, record in enumerate(records, 1):
        raw_type = record.get("facilityType") or record.get("rawFacilityType")
        rows.append([
            index,
            REPORT_TYPE_LABELS.get(raw_type, raw_type or "-"),
            record.get("administrativeVillage") or "-",
            record.get("village") or record.get("town") or "-",
            STATUS_LABELS.get(record.get("status"), record.get("status") or "-"),
            f"{float(record.get('totalScore') or 0):.2f}",
        ])
    return rows


def _record_point_name(record: dict[str, Any]) -> str:
    return record.get("village") or record.get("town") or "该项目点"


def _format_report_time(value: Any) -> str:
    if not value:
        return "-"
    text = str(value).strip()
    try:
        parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
        return parsed.strftime("%Y-%m-%d %H:%M")
    except ValueError:
        return text.replace("T", " ")


def _deduction_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for record in records:
        for score in record.get("scores") or []:
            if float(score.get("deduction") or 0) <= 0:
                continue
            rows.append([
                len(rows) + 1,
                record.get("village") or record.get("town") or "-",
                score.get("indicatorName") or "-",
                score.get("indicatorFullScore") or "-",
                score.get("deduction") or "0",
                score.get("reason") or score.get("deductionOptionName") or "-",
            ])
    return rows


def _add_deduction_narrative(document, records: list[dict[str, Any]], *, limit: int | None = None) -> None:
    rows = _deduction_rows(records)
    if limit is not None:
        rows = rows[:limit]
    if not rows:
        _add_paragraph(document, "经现场检查和资料核查，本期未发现需要按考核标准扣分的事项。项目公司仍应持续做好运行记录、设施巡查、水质检测和问题整改资料的整理归档。")
        return
    grouped: dict[str, list[list[Any]]] = {}
    for row in rows:
        grouped.setdefault(str(row[2]), []).append(row)
    for index, (indicator, items) in enumerate(grouped.items(), 1):
        total = sum(float(item[4] or 0) for item in items)
        points = "、".join(dict.fromkeys(str(item[1]) for item in items))
        reason_parts = []
        for item in items:
            if not item[5] or str(item[5]) == "-":
                continue
            reason_parts.extend(part.strip() for part in str(item[5]).replace(";", "；").split("；") if part.strip())
        reasons = "；".join(dict.fromkeys(reason_parts))
        if total >= 5:
            degree = "对本项得分影响较明显"
        elif len(items) > 1:
            degree = "属于同类问题重复出现"
        else:
            degree = "问题相对单一"
        text = f"{index}、{points}在“{indicator}”方面存在不符合考核要求的情况，合计扣{_fmt_score_value(total)}分，{degree}"
        if reasons:
            text += f"，主要表现为：{reasons}"
        _add_paragraph(document, text + "。")


def _all_score_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for record in records:
        point_name = record.get("village") or record.get("town") or "-"
        for score in record.get("scores") or []:
            full_score = float(score.get("indicatorFullScore") or 0)
            deduction = float(score.get("deduction") or 0)
            rows.append([
                len(rows) + 1,
                point_name,
                score.get("indicatorName") or "-",
                f"{full_score:.2f}",
                f"{max(full_score - deduction, 0):.2f}",
                f"{deduction:.2f}",
                score.get("reason") or score.get("deductionOptionName") or ("无扣分" if deduction == 0 else "-"),
            ])
    return rows


def _water_number(value: Any) -> float | None:
    match = re.search(r"\d+(?:\.\d+)?", str(value or ""))
    return float(match.group()) if match else None


def _water_result(value: Any, limit: Any) -> str:
    measured = _water_number(value)
    maximum = _water_number(limit)
    if measured is None or maximum is None:
        return "待判定"
    return "达标" if measured <= maximum else "不达标"


def _water_quality_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for record in records:
        for item in record.get("waterQuality") or []:
            payload = item.get("payload") or {}
            point_name = record.get("village") or record.get("town") or "-"
            sample_time = _format_report_time(payload.get("sampleTime") or item.get("sampledAt"))
            metric_rows = [
                ("CODCr", payload.get("codValue"), payload.get("codLimit")),
                ("BOD5", payload.get("bod5Value"), payload.get("bod5Limit")),
                ("SS", payload.get("ssValue"), payload.get("ssLimit")),
                ("NH3-N", payload.get("nh3nValue"), payload.get("nh3nLimit")),
                ("TP", payload.get("tpValue"), payload.get("tpLimit")),
            ]
            automatic_results: list[str] = []
            for metric, value, limit in metric_rows:
                if limit in (None, ""):
                    continue
                result = _water_result(value, limit)
                automatic_results.append(result)
                rows.append([
                    len(rows) + 1,
                    point_name,
                    sample_time,
                    metric,
                    value if value not in (None, "") else "-",
                    limit,
                    "mg/L",
                    result,
                    result,
                    "-",
                ])
            automatic_overall = "待判定" if "待判定" in automatic_results else "不达标" if "不达标" in automatic_results else "达标"
            if payload.get("monthlyMissingTest") or payload.get("monthlyRegulatorUnqualified"):
                automatic_overall = "不达标"
            conclusion = item.get("conclusion") or payload.get("conclusion")
            final_result = "达标" if conclusion == "qualified" else "不达标" if conclusion == "unqualified" else "待判定"
            overridden = bool(payload.get("conclusionOverridden") or payload.get("manualOverride"))
            note = payload.get("note") or payload.get("remark") or "-"
            rows.append([
                len(rows) + 1,
                point_name,
                sample_time,
                "综合判定",
                "-",
                "-",
                "-",
                automatic_overall,
                final_result,
                f"人工修改：{note}" if overridden else note,
            ])
    return rows


def _unqualified_water_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    return [row for row in _water_quality_rows(records) if row[7] == "不达标" or row[8] == "不达标"]


def _operation_volume_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for record in records:
        raw = record.get("rawPayload") or {}
        water = raw.get("waterQuality") or {}
        volume = (
            raw.get("actualTreatmentVolume")
            or raw.get("treatmentVolume")
            or raw.get("monthlyVolume")
            or water.get("actualTreatmentVolume")
            or water.get("monthlyVolume")
        )
        if volume in (None, ""):
            continue
        rows.append([
            len(rows) + 1,
            _record_point_name(record),
            water.get("designScale") or raw.get("designScale") or "-",
            volume,
            raw.get("volumePeriod") or water.get("volumePeriod") or "本考核周期",
        ])
    return rows


def _monthly_water_average_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for record in records:
        raw = record.get("rawPayload") or {}
        water = raw.get("waterQuality") or {}
        influent = raw.get("monthlyInfluentCod") or water.get("monthlyInfluentCod")
        effluent = raw.get("monthlyEffluentCod") or water.get("monthlyEffluentCod")
        if influent in (None, "") and effluent in (None, ""):
            continue
        rows.append([
            len(rows) + 1,
            _record_point_name(record),
            raw.get("waterMonth") or water.get("waterMonth") or "本考核周期",
            influent if influent not in (None, "") else "-",
            effluent if effluent not in (None, "") else "-",
            raw.get("monthlyWaterConclusion") or water.get("monthlyWaterConclusion") or "需结合月度资料核定",
        ])
    return rows


def _survey_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for record in records:
        for item in record.get("surveys") or []:
            rows.append([
                len(rows) + 1,
                record.get("village") or "-",
                SURVEY_TYPE_LABELS.get(item.get("surveyType"), item.get("surveyType") or "-"),
                RESPONDENT_LABELS.get(item.get("respondent"), item.get("respondent") or "-"),
                item.get("score") if item.get("score") is not None else "-",
            ])
    return rows


def _is_detection_attachment(item: dict[str, Any]) -> bool:
    filename = (item.get("filename") or "").lower()
    return Path(filename).suffix.lower() == ".pdf" or any(word in filename for word in ("检测", "化验", "水质报告", "检验"))


def _attachment_rows(records: list[dict[str, Any]], *, detection: bool | None = None) -> list[list[Any]]:
    rows = []
    for record in records:
        for item in record.get("attachments") or []:
            if detection is not None and _is_detection_attachment(item) != detection:
                continue
            rows.append([
                len(rows) + 1,
                record.get("village") or record.get("town") or "-",
                item.get("filename") or "-",
                item.get("scoreId") or "-",
                item.get("deductionOptionId") or "-",
                item.get("size") or 0,
            ])
    return rows


def _add_attachment_pictures(document, records: list[dict[str, Any]], *, detection: bool = False) -> tuple[int, int]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    inserted = 0
    skipped = 0
    image_content_types = {"image/jpeg", "image/png", "image/gif", "image/bmp", "image/tiff"}
    image_suffixes = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff"}
    for record in records:
        point_name = record.get("village") or record.get("town") or "项目点"
        for item in record.get("attachments") or []:
            if _is_detection_attachment(item) != detection:
                continue
            path = Path(item.get("storageKey") or "")
            content_type = (item.get("contentType") or "").lower()
            is_pdf = content_type == "application/pdf" or path.suffix.lower() == ".pdf"
            if content_type not in image_content_types and path.suffix.lower() not in image_suffixes and not is_pdf:
                continue
            if not path.is_file():
                skipped += 1
                continue
            try:
                pictures: list[tuple[Any, int, int]] = []
                if is_pdf:
                    import fitz

                    with fitz.open(path) as pdf:
                        for page in pdf:
                            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
                            pictures.append((BytesIO(pixmap.tobytes("png")), pixmap.width, pixmap.height))
                else:
                    import fitz

                    pixmap = fitz.Pixmap(str(path))
                    pictures.append((str(path), pixmap.width, pixmap.height))
                for page_index, (picture, pixel_width, pixel_height) in enumerate(pictures, 1):
                    max_width, max_height = 15.5, 20.0
                    aspect = pixel_width / max(pixel_height, 1)
                    width = max_width
                    height = width / max(aspect, 0.01)
                    if height > max_height:
                        height = max_height
                        width = height * aspect
                    if pixel_width < 64 or pixel_height < 64:
                        width = min(width, 5.0)
                        height = width / max(aspect, 0.01)
                    document.add_picture(picture, width=Cm(width), height=Cm(height))
                    picture_paragraph = document.paragraphs[-1]
                    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = document.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    page_suffix = f"（第{page_index}页）" if len(pictures) > 1 else ""
                    caption_text = (
                        f"检测资料 {inserted + 1}　项目点：{point_name}{page_suffix}"
                        if detection
                        else f"序号：{inserted + 1}　项目点：{point_name}"
                    )
                    run = caption.add_run(caption_text)
                    _apply_run_font(run, REPORT_BODY_FONT, REPORT_BODY_SIZE_PT)
                    _apply_paragraph_format(caption, indent=False)
                    inserted += 1
            except Exception:
                skipped += 1
    return inserted, skipped


def _accepted_agent_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    capability_labels = {
        "record_review_assist": "复核记录辅助校验",
        "report_semantic_check": "报告语义一致性校验",
    }
    rows = []
    for record in records:
        for item in record.get("agentRuns") or []:
            output = item.get("output") or {}
            capability = item.get("capability") or "-"
            rows.append([
                len(rows) + 1,
                record.get("village") or record.get("town") or "-",
                capability_labels.get(capability, capability),
                f"{float(item.get('confidence') or 0):.2f}",
                output.get("summary") or "-",
            ])
    return rows


def _data_file(*parts: str) -> Path:
    return settings.backend_dir / "app" / "data" / Path(*parts)


def _load_json_data(*parts: str) -> Any:
    path = _data_file(*parts)
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def _project_standard_key(project_name: str) -> str:
    return "maonan" if "茂南" in project_name else "yunan"


def _project_source_template(project_name: str) -> dict[str, Any]:
    name = "maonan-structure.json" if "茂南" in project_name else "yunan-legacy-structure.json"
    return _load_json_data("report_source_format", name) or {}


def _project_standards(project_name: str) -> dict[str, list[dict[str, Any]]]:
    data = _load_json_data("project_standards.json") or {}
    return data.get(_project_standard_key(project_name), {})


def _standard_rows_for_type(project_name: str, facility_type: str) -> list[list[Any]]:
    standards = _project_standards(project_name).get(facility_type) or []
    rows: list[list[Any]] = []
    for group in standards:
        for child in group.get("children") or []:
            for item in child.get("items") or []:
                rows.append([
                    len(rows) + 1,
                    group.get("name") or "-",
                    child.get("name") or "-",
                    item.get("name") or "-",
                    f"{float(item.get('maxScore') or 0):.2f}",
                    item.get("scoringMethod") or "-",
                    item.get("evaluationStandard") or item.get("standardText") or "-",
                ])
    return rows


def _standards_overview_rows(project_name: str, records: list[dict[str, Any]]) -> list[list[Any]]:
    standards = _project_standards(project_name)
    facility_types = list(standards.keys()) or _facility_types(records)
    rows = []
    for index, facility_type in enumerate(facility_types, 1):
        groups = standards.get(facility_type) or []
        item_count = sum(len(child.get("items") or []) for group in groups for child in group.get("children") or [])
        full_score = sum(
            float(item.get("maxScore") or 0)
            for group in groups
            for child in group.get("children") or []
            for item in child.get("items") or []
        )
        rows.append([index, REPORT_TYPE_LABELS.get(facility_type, facility_type), len(groups), item_count, f"{full_score:.2f}"])
    return rows


def _records_by_town(records: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    result: dict[str, list[dict[str, Any]]] = {}
    for record in records:
        result.setdefault(record.get("town") or "未命名镇街", []).append(record)
    return result


def _records_by_type(records: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    result: dict[str, list[dict[str, Any]]] = {}
    for record in records:
        facility_type = record.get("facilityType") or record.get("rawFacilityType") or "unknown"
        result.setdefault(facility_type, []).append(record)
    return result


def _project_full_name(project_name: str) -> str:
    if "茂南" in project_name:
        return "茂南区水质净化处理设施全区捆绑 PPP 项目"
    if "郁南" in project_name:
        return "郁南县整县生活污水处理设施捆绑 PPP 项目"
    return project_name


def _project_cover_unit(project_name: str) -> tuple[str, str]:
    if "茂南" in project_name:
        return "委托单位：茂名市茂南区住房和城乡建设局", "编制单位：广东省建筑设计研究院集团股份有限公司"
    if "郁南" in project_name:
        return "实施机构：郁南县住房和城乡建设局", "考核单位：广东省建筑设计研究院集团股份有限公司"
    return "委托单位：项目主管单位", "编制单位：绩效考核工作组"


def _add_project_personnel_page(document, project_name: str, cycle_name: str) -> None:
    document.add_heading("项目人员组成", level=1)
    if "茂南" in project_name:
        rows = [
            ["项目名称", f"茂南区水质净化处理设施全区捆绑PPP项目城镇水质净化设施运营期绩效考核服务项目（城镇设施{cycle_name}绩效考核报告）"],
            ["编制单位", "广东省建筑设计研究院集团股份有限公司"],
            ["工程咨询单位甲级资信证书编号", "甲232024011004（市政公用工程）"],
            ["法定代表人", "李巍"],
            ["审定人", "黄志聪"],
            ["审核人", "刘钰坤"],
            ["项目负责人", "曹雅娟"],
            ["校对", "陶艺婷"],
            ["现场考核", "陶艺婷、郭庆鑫"],
            ["报告编制", "陶艺婷、郭庆鑫、黄灿栩"],
            ["委托单位考核组", "组长：李庆龙；组员：梁忠智"],
            ["检测单位", "广东华蓝检测技术有限公司"],
            ["检测报告人员", "编制：罗桂珠；复核：潘浩贤；签发：孔令峰"],
        ]
    else:
        rows = [
            ["项目名称", f"郁南县整县生活污水处理设施捆绑PPP项目绩效考核（{cycle_name}镇级及农村设施考核报告）"],
            ["编制单位", "广东省建筑设计研究院集团股份有限公司"],
            ["工程咨询单位甲级资信证书编号", "甲232024011004（市政公用工程）"],
            ["法定代表人", "李巍"],
            ["技术总负责人", "罗赤宇"],
            ["审定人", "黄志聪"],
            ["审核人", "刘钰坤"],
            ["项目负责人", "陶艺婷"],
            ["校对", "曹雅娟、吴浩"],
            ["现场考核负责人", "陶艺婷、林敏仪"],
            ["报告编制", "陶艺婷、林敏仪"],
            ["实施机构考核组", "组长：彭树标；组员：伍界宇"],
            ["检测单位", "广东华蓝检测技术有限公司"],
            ["检测报告人员", "编制：罗桂珠；复核：潘浩贤；签发：孔令峰"],
        ]
    for label, value in rows:
        _add_paragraph(document, f"{label}：{value}", bold_prefix=f"{label}：", indent=False, size_pt=14)
    _add_paragraph(document, "本期项目人员组成以委托单位和编制单位最终确认的信息为准。", indent=False, size_pt=14)
    document.add_page_break()


def _add_source_cover(document, *, project_name: str, cycle_name: str, title: str, report_type: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    full_name = _project_full_name(project_name)
    owner_line, compiler_line = _project_cover_unit(project_name)
    is_maonan = "茂南" in project_name

    def add_centered(text: str, size: float, *, bold: bool = False) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        _apply_run_font(run, REPORT_HEADING_FONT, size, bold=bold)

    for _ in range(2):
        document.add_paragraph("")
    add_centered(full_name, 20 if is_maonan else 22, bold=True)
    if is_maonan:
        add_centered("城镇水质净化设施运营期绩效考核", 20, bold=True)
        add_centered("服务项目", 20, bold=True)
        for _ in range(3):
            document.add_paragraph("")
        add_centered(f"城镇设施{cycle_name}绩效考核报告", 22, bold=True)
        spacer_count = 7
    else:
        add_centered("绩效考核", 22, bold=True)
        for _ in range(2):
            document.add_paragraph("")
        add_centered(f"{cycle_name}镇级及农村设施考核报告", 22, bold=True)
        spacer_count = 8
    for _ in range(spacer_count):
        document.add_paragraph("")
    add_centered(owner_line, 16)
    add_centered(compiler_line, 16)
    now = datetime.now()
    digit_names = "〇一二三四五六七八九"
    month_names = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十", "十一", "十二"]
    cover_date = "".join(digit_names[int(value)] for value in str(now.year)) + "年" + month_names[now.month] + "月"
    add_centered(cover_date, 16)
    document.add_page_break()
    _add_project_personnel_page(document, project_name, cycle_name)


def _add_assessment_summary_text(document, *, project_name: str, cycle_name: str, records: list[dict[str, Any]], towns: list[dict[str, Any]], profile: dict[str, Any]) -> None:
    stats = _score_stats(records)
    deductions = _deduction_rows(records)
    town_count = len(towns) or len(_records_by_town(records))
    is_maonan = "茂南" in project_name
    document.add_heading("摘  要" if is_maonan else "摘要", level=1)
    if is_maonan:
        _add_paragraph(document, "根据《茂南区水质净化处理设施全区捆绑PPP项目城镇水质净化设施运营期绩效考核服务项目》约定，茂名市茂南区住房和城乡建设局委托广东省建筑设计研究院集团股份有限公司开展茂南区水质净化处理设施全区捆绑PPP项目绩效考核工作，绩效考核结果将作为支付项目服务费的依据。")
        document.add_heading("一、考核工作开展情况", level=2)
        _add_paragraph(document, f"根据《PPP项目合同》等文件规定的考核标准和原则，以及各城镇水质净化厂和污水收集管网批复转运营时间，本报告对{cycle_name}纳入考核范围的城镇水质净化厂、配套管网及相关运维资料开展绩效考核。现场考核完成后，考核组结合现场检查情况、水质检测结果和资料复核意见，对本期绩效情况作出客观、公正、全面的评价。")
        document.add_heading("二、考核评分情况", level=2)
        _add_paragraph(document, f"本期覆盖{town_count}个镇街、{len(records)}个已提交、已复核或已锁定考核对象，平均得分为{stats['average']:.2f}分，最高得分为{stats['max']:.2f}分，最低得分为{stats['min']:.2f}分，累计扣分{stats['deduction']:.2f}分。")
        document.add_heading("三、主要改进点", level=2)
        _add_paragraph(document, "本期重点评价设施运行、现场管理、在线监测、管网巡查和问题整改等工作，并结合现场检查记录、运维资料和复核意见，分别说明各镇街考核对象的得分情况。")
        document.add_heading("四、发现的主要问题", level=2)
    else:
        _add_paragraph(document, f"我方受郁南县住房和城乡建设局委托，依据《郁南县整县生活污水处理捆绑PPP项目合同》及补充协议约定，以及有关法律、法规、标准和规范要求，在审阅项目公司提交资料、开展现场检查并核实有关情况的基础上，对郁南县整县生活污水处理设施捆绑PPP项目所含镇级污水处理子项目开展{cycle_name}绩效考核工作。")
        document.add_heading("一、考核工作开展情况", level=2)
        _add_paragraph(document, f"根据镇级污水处理厂及污水收集管网建设完成情况，以及郁南县住房和城乡建设局工作安排及合同相关约定，考核组对{town_count}个镇街纳入本期考核范围的镇级污水处理厂、镇区污水收集管网和农村污水处理设施开展现场考核、资料核查、问卷调查、水质检测和评分复核。")
        document.add_heading("二、考核评分情况", level=2)
        _add_paragraph(document, f"本期纳入报告的已提交、已复核或已锁定考核对象共{len(records)}个，平均得分为{stats['average']:.2f}分，最高得分为{stats['max']:.2f}分，最低得分为{stats['min']:.2f}分，累计扣分{stats['deduction']:.2f}分。各镇级设施、镇区管网和农村设施的具体评分情况见正文及附件评分表。")
        document.add_heading("三、发现的主要问题", level=2)
    if deductions:
        leading = _deduction_rows(records)[:8]
        _add_paragraph(document, "根据现场检查记录和资料复核结果，本期主要问题如下。")
        _add_deduction_narrative(document, records, limit=8)
        _add_simple_table(document, ["序号", "设施点", "评分条目", "满分", "扣分", "依据或说明"], leading)
    else:
        _add_deduction_narrative(document, records)
    document.add_heading("五、建议" if is_maonan else "四、建议", level=2)
    if is_maonan:
        _add_paragraph(document, "建议项目公司加强各项运行管理制度落实，提升规范化管理水平，强化档案管理和数字化建设，规范设施运行，及时上报问题，增强对重点片区的巡查力度。")
        _add_paragraph(document, "针对负荷偏低、在线监测异常、管网巡查记录不完整、安全管理措施落实不到位等问题，应按考核标准逐项建立整改台账并落实复核材料。")
    else:
        _add_paragraph(document, "建议项目公司足额配备运维管理人员，加强培训力度，提升运维管理水平，有效落实巡检巡查、设施运行维护、污泥处置、安全管理和资料归档工作。")
        _add_paragraph(document, "针对设施负荷率偏低、进水COD浓度偏低、出水不达标、管网堵塞或整改不到位等情况，应系统研究提质增效对策，逐项闭环整改。")


def _town_basic_descriptions(town_data: dict[str, Any] | None, records: list[dict[str, Any]]) -> list[str]:
    descriptions: list[str] = []
    if town_data:
        for target in town_data.get("assessmentTargets") or []:
            item = (town_data.get("assessmentObject") or {}).get(target) or {}
            description = item.get("description")
            if description and description not in descriptions:
                descriptions.append(description)
    for record in records:
        village_object = record.get("villageAssessmentObject") or {}
        description = village_object.get("description")
        if description and description not in descriptions:
            descriptions.append(description)
    return descriptions


def _facility_basic_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for record in records:
        raw = record.get("rawPayload") or {}
        water = raw.get("waterQuality") or {}
        rows.append([
            len(rows) + 1,
            _record_point_name(record),
            REPORT_TYPE_LABELS.get(record.get("facilityType"), record.get("facilityType") or "-"),
            water.get("designScale") or raw.get("designScale") or "本期资料未记录",
            water.get("processType") or raw.get("processType") or "本期资料未记录",
            water.get("dischargeStandard") or raw.get("dischargeStandard") or "按项目适用标准执行",
        ])
    return rows


def _add_town_basic_intro(document, *, town_name: str, town_data: dict[str, Any] | None, records: list[dict[str, Any]], add_heading: bool = True) -> None:
    if add_heading:
        document.add_heading("基本情况", level=3)
    descriptions = _town_basic_descriptions(town_data, records)
    if descriptions:
        for description in descriptions:
            _add_paragraph(document, description)
    else:
        _add_paragraph(document, f"{town_name}本期考核对象及其基本情况以项目资料、现场检查记录和经核实的支撑材料为准。")
    if records:
        type_labels = [REPORT_TYPE_LABELS.get(item, item) for item in _facility_types(records)]
        _add_paragraph(document, f"本期共对{town_name}{len(records)}个考核对象完成资料复核，涉及{'、'.join(type_labels) or '污水处理设施'}。各考核对象的基本资料如下。")
        _add_simple_table(document, ["序号", "项目点", "考核对象", "设计规模（m³/d）", "处理工艺", "执行标准"], _facility_basic_rows(records))


def _add_town_type_sections(document, *, records: list[dict[str, Any]], project_name: str, towns: list[dict[str, Any]] | None = None) -> None:
    town_lookup = {town.get("town"): town for town in towns or []}
    by_town = _records_by_town(records)
    for town_index, (town_name, town_records) in enumerate(by_town.items(), 1):
        document.add_heading(f"{town_index}. {town_name}考核情况", level=2)
        _add_town_basic_intro(document, town_name=town_name, town_data=town_lookup.get(town_name), records=town_records)
        _add_paragraph(document, f"{town_name}本期纳入{len(town_records)}个考核对象。以下按正式报告“基本情况、运维考核情况、扣分情况、整改建议”的顺序展开。")
        town_types = _records_by_type(town_records)
        for type_index, (facility_type, items) in enumerate(town_types.items(), 1):
            label = REPORT_TYPE_LABELS.get(facility_type, facility_type)
            stats = _score_stats(items)
            document.add_heading(f"{town_index}.{type_index} {label}", level=3)
            _add_paragraph(document, f"本类别共{len(items)}个考核对象，平均得分{stats['average']:.2f}分，累计扣分{stats['deduction']:.2f}分。")
            _add_simple_table(document, ["序号", "考核类型", "行政村", "项目点", "状态", "得分"], _record_score_rows(items))
            deductions = _deduction_rows(items)
            if deductions:
                _add_simple_table(document, ["序号", "设施点", "评分条目", "满分", "扣分", "依据或说明"], deductions)
            else:
                _add_paragraph(document, "经核查，本类别未发现需扣分的问题。")


def _town_status_label(town: dict[str, Any], records: list[dict[str, Any]]) -> str:
    if not records:
        return "\u672a\u63d0\u4ea4/\u672a\u590d\u6838"
    if town.get("lockedCount", 0) and town.get("lockedCount") == len(records):
        return STATUS_LABELS.get("locked", "\u5df2\u9501\u5b9a")
    if town.get("reviewedCount", 0) + town.get("lockedCount", 0) == len(records):
        return STATUS_LABELS.get("reviewed", "\u5df2\u590d\u6838")
    return STATUS_LABELS.get("submitted", "\u5df2\u63d0\u4ea4")


def _towns_for_report(records: list[dict[str, Any]], towns: list[dict[str, Any]] | None) -> list[tuple[str, dict[str, Any] | None, list[dict[str, Any]]]]:
    by_town = _records_by_town(records)
    if towns:
        result = [(town.get("town") or "-", town, by_town.get(town.get("town") or "", [])) for town in _ordered_towns(towns)]
        known = {name for name, _, _ in result}
        for name, items in by_town.items():
            if name not in known:
                result.append((name, None, items))
        return result
    return [(name, None, items) for name, items in by_town.items()]


def _add_summary_town_type_sections(document, *, records: list[dict[str, Any]], project_name: str, towns: list[dict[str, Any]]) -> None:
    for town_index, (town_name, town_data, town_records) in enumerate(_towns_for_report(records, towns), 1):
        document.add_heading(f"{town_index}. {town_name}\u8003\u6838\u60c5\u51b5", level=2)
        if not town_records:
            _add_paragraph(document, f"{town_name}本期尚无经复核确认的考核资料，因此不纳入本期评分统计。待现场考核资料齐备并完成复核后，再补充该镇街的评分结果和问题分析。")
            _add_simple_table(
                document,
                ["\u5e8f\u53f7", "\u9547\u8857", "\u8003\u6838\u5bf9\u8c61\u6570", "\u72b6\u6001", "\u8bf4\u660e"],
                [[1, town_name, town_data.get("recordCount", 0) if town_data else 0, "\u672a\u63d0\u4ea4/\u672a\u590d\u6838", "\u672a\u7eb3\u5165\u672c\u671f\u8bc4\u5206\u7edf\u8ba1"]],
            )
            continue
        _add_town_type_sections(document, records=town_records, project_name=project_name, towns=[town_data] if town_data else None)


def _add_yunan_summary_score_overview(document, records: list[dict[str, Any]], towns: list[dict[str, Any]]) -> None:
    town_rows = []
    for index, (town_name, town_data, items) in enumerate(_towns_for_report(records, towns), 1):
        if items:
            stats = _score_stats(items)
            town_rows.append([index, town_name, len(items), f"{stats['average']:.2f}", f"{stats['max']:.2f}", f"{stats['min']:.2f}", f"{stats['deduction']:.2f}", _town_status_label(town_data or {}, items)])
        else:
            town_rows.append([index, town_name, 0, "-", "-", "-", "-", "\u672a\u63d0\u4ea4/\u672a\u590d\u6838"])
    _add_simple_table(document, ["\u5e8f\u53f7", "\u9547\u8857", "\u8003\u6838\u5bf9\u8c61\u6570", "\u5e73\u5747\u5f97\u5206", "\u6700\u9ad8\u5f97\u5206", "\u6700\u4f4e\u5f97\u5206", "\u7d2f\u8ba1\u6263\u5206", "\u72b6\u6001"], town_rows)

    type_rows = []
    by_type = _records_by_type(records)
    for index, (facility_type, items) in enumerate(by_type.items(), 1):
        stats = _score_stats(items)
        type_rows.append([index, REPORT_TYPE_LABELS.get(facility_type, facility_type), len(items), f"{stats['average']:.2f}", f"{stats['deduction']:.2f}"])
    _add_simple_table(document, ["\u5e8f\u53f7", "\u8003\u6838\u5bf9\u8c61\u7c7b\u522b", "\u8bb0\u5f55\u6570", "\u5e73\u5747\u5f97\u5206", "\u7d2f\u8ba1\u6263\u5206"], type_rows)

    object_rows = []
    for index, record in enumerate(records, 1):
        object_rows.append([
            index,
            record.get("town") or "-",
            REPORT_TYPE_LABELS.get(record.get("facilityType"), record.get("facilityType") or "-"),
            _record_point_name(record),
            f"{float(record.get('totalScore') or 0):.2f}",
            STATUS_LABELS.get(record.get("status"), record.get("status") or "-"),
        ])
    _add_simple_table(document, ["\u5e8f\u53f7", "\u9547\u8857", "\u8003\u6838\u5bf9\u8c61", "\u9879\u76ee\u70b9", "\u5f97\u5206", "\u72b6\u6001"], object_rows)

    deduction_rows = []
    for index, (facility_type, items) in enumerate(by_type.items(), 1):
        rows = _deduction_rows(items)
        stats = _score_stats(items)
        deduction_rows.append([
            index,
            REPORT_TYPE_LABELS.get(facility_type, facility_type),
            len(rows),
            f"{stats['deduction']:.2f}",
            "有扣分事项，详见附件2考核评分表" if rows else "本期未记录扣分事项",
        ])
    _add_simple_table(document, ["\u5e8f\u53f7", "\u8003\u6838\u5bf9\u8c61\u7c7b\u522b", "\u6263\u5206\u9879\u6570", "\u7d2f\u8ba1\u6263\u5206", "\u8bf4\u660e"], deduction_rows)

    support_rows = []
    for index, record in enumerate(records, 1):
        support_rows.append([
            index,
            record.get("town") or "-",
            _record_point_name(record),
            len(record.get("scores") or []),
            len(record.get("waterQuality") or []),
            len(record.get("attachments") or []),
        ])
    _add_simple_table(document, ["\u5e8f\u53f7", "\u9547\u8857", "\u9879\u76ee\u70b9", "\u8bc4\u5206\u6761\u76ee", "\u6c34\u8d28\u8bb0\u5f55", "\u73b0\u573a\u9644\u4ef6"], support_rows)

    _add_paragraph(document, "\u6c47\u603b\u62a5\u544a\u4ec5\u5217\u793a\u5f53\u524d\u9879\u76ee\u3001\u5f53\u524d\u5468\u671f\u5185\u5df2\u63d0\u4ea4\u3001\u5df2\u590d\u6838\u6216\u5df2\u9501\u5b9a\u5e76\u7eb3\u5165\u62a5\u544a\u7684\u9547\u8857\u6570\u636e\uff1b\u8349\u7a3f\u3001\u9000\u56de\u548c\u672a\u63d0\u4ea4\u7684\u9547\u8857\u4e0d\u8fdb\u5165\u6c47\u603b\u8868\u3001\u6b63\u6587\u7ae0\u8282\u548c\u7edf\u8ba1\u8ba1\u7b97\u3002")
    _add_paragraph(document, "\u5bf9\u5f97\u5206\u8f83\u4f4e\u6216\u6263\u5206\u8f83\u96c6\u4e2d\u7684\u9879\u76ee\u70b9\uff0c\u5e94\u7ed3\u5408\u73b0\u573a\u7167\u7247\u3001\u6c34\u8d28\u62bd\u68c0\u3001\u8fd0\u884c\u53f0\u8d26\u548c\u6574\u6539\u8bb0\u5f55\u9010\u9879\u590d\u6838\uff0c\u5f62\u6210\u4e0b\u4e00\u5468\u671f\u91cd\u70b9\u8ddf\u8e2a\u6e05\u5355\u3002")


def _add_yunan_work_section(document, *, project_name: str, cycle_name: str, scope_name: str, records: list[dict[str, Any]], profile: dict[str, Any], is_summary: bool) -> None:
    _add_paragraph(document, f"根据镇级污水处理厂、镇区污水收集管网及农村污水处理设施建设和运营情况，以及{project_name}绩效考核工作安排，考核组对{scope_name}开展{cycle_name}绩效考核。")
    _add_paragraph(document, f"本次考核依据项目合同、补充协议、考核标准及有关法律、法规、标准和规范要求，通过{'、'.join(profile['methods'])}等方式，对项目公司提供的运行维护资料、现场检查情况和水质检测结果进行核查。")
    _add_paragraph(document, "检查过程中，考核组分别对镇级污水处理厂、镇区污水收集管网和农村污水处理设施进行评价，现场形成检查记录，并经资料复核后确认扣分情况。")
    _add_simple_table(document, ["序号", "考核对象", "一级指标数", "评分项数", "满分"], _standards_overview_rows(project_name, records))
    _add_simple_table(document, ["序号", "考核方法", "资料或工作内容"], [[index, method, "按本次考核工作安排实施，并形成相应检查记录和支撑资料。"] for index, method in enumerate(profile["methods"], 1)])
    if is_summary:
        _add_paragraph(document, "汇总报告覆盖当前项目和当前周期内全部已提交、已复核或已锁定的镇街数据；草稿、退回和未提交数据不纳入本次评分和报告结论。")
    _add_paragraph(document, "报告先说明考核工作开展情况，再按考核对象汇总评分，随后分析发现的主要问题并提出整改建议；评分标准、评分明细、现场照片、水质资料和资料清单列入附件。")
    _add_paragraph(document, "现场考核和资料核查以项目公司提交资料、现场检查记录、复核意见及水质抽检资料为依据，涉及扣分的事项均在评分表中对应到具体评分条目。")


def _add_yunan_problem_section(document, records: list[dict[str, Any]]) -> None:
    deductions = _deduction_rows(records)
    if not deductions:
        _add_paragraph(document, "经核查，本期未发现需按考核标准扣分的问题。后续仍应持续关注巡检巡查、运行台账、水质检测、现场安全和问题整改等工作。")
        return
    grouped: dict[str, list[list[Any]]] = {}
    for row in deductions:
        grouped.setdefault(str(row[2]), []).append(row)
    leading = sorted(grouped.items(), key=lambda item: sum(float(row[4] or 0) for row in item[1]), reverse=True)
    _add_paragraph(document, "本期考核发现的主要问题如下。有关问题均经现场检查或资料核查确认，具体扣分情况详见附件评分表。")
    summary_rows = []
    for index, (name, rows) in enumerate(leading[:8], 1):
        total = sum(float(row[4] or 0) for row in rows)
        points = "、".join(str(row[1]) for row in rows[:6])
        summary_rows.append([index, name, len(rows), f"{total:.2f}", points])
        _add_paragraph(document, f"{index}、{name}方面存在问题，涉及{len(rows)}项扣分，累计扣{total:.2f}分，主要涉及{points}等项目点。")
    _add_simple_table(document, ["序号", "问题类别", "涉及项数", "累计扣分", "涉及项目点"], summary_rows)
    _add_simple_table(document, ["序号", "设施点", "评分条目", "满分", "扣分", "依据或说明"], deductions)


def _add_yunan_suggestion_section(document, records: list[dict[str, Any]]) -> None:
    _add_paragraph(document, "一是建议项目公司针对本期考核扣分项建立整改台账，明确责任部门、责任人员、整改措施和完成时限，整改完成后形成照片、记录、台账等佐证材料。")
    _add_paragraph(document, "二是建议持续加强运行维护管理，补齐生产运行、维护维修、巡查巡检、污泥处置、安全管理和人员持证等资料，确保资料内容真实、完整、可追溯。")
    _add_paragraph(document, "三是建议加强水质检测和工艺调控，对进出水水量、水质波动、设施负荷率和达标排放情况开展持续跟踪，发现异常及时分析原因并落实处理措施。")
    _add_paragraph(document, "四是建议落实公众调查、问题反馈和整改销号机制，对群众反映问题、现场检查问题和历次考核遗留问题逐项闭环，提升设施运行效果和服务质量。")
    if _deduction_rows(records):
        _add_paragraph(document, "对本期扣分较集中的评分点，应优先纳入下一周期复核重点，防止同类问题重复出现。")
    _add_simple_table(document, ["序号", "整改方向", "主要措施", "建议完成资料"], [
        [1, "运行维护", "完善生产运行、维护维修、巡查巡检记录，确保记录真实完整。", "运行台账、维修记录、巡查记录"],
        [2, "水质达标", "持续跟踪进出水水质和设施负荷，异常时及时调控工艺。", "检测报告、工艺调整记录"],
        [3, "现场管理", "清理现场杂物，维护构筑物、设备、管网和安全标识。", "整改前后照片、现场检查记录"],
        [4, "人员与安全", "补齐岗位人员、持证情况、安全培训和应急演练资料。", "人员证书、培训记录、演练记录"],
        [5, "整改闭环", "对考核发现问题逐项销号，并在下一周期复核。", "整改台账、复核意见"],
    ])


def _add_yunan_score_overview(document, records: list[dict[str, Any]]) -> None:
    by_town = _records_by_town(records)
    town_rows = []
    for index, (town, items) in enumerate(by_town.items(), 1):
        stats = _score_stats(items)
        town_rows.append([index, town, len(items), f"{stats['average']:.2f}", f"{stats['max']:.2f}", f"{stats['min']:.2f}", f"{stats['deduction']:.2f}"])
    _add_simple_table(document, ["序号", "镇街", "考核对象数", "平均得分", "最高得分", "最低得分", "累计扣分"], town_rows)

    type_rows = []
    by_type = _records_by_type(records)
    for index, (facility_type, items) in enumerate(by_type.items(), 1):
        stats = _score_stats(items)
        type_rows.append([index, REPORT_TYPE_LABELS.get(facility_type, facility_type), len(items), f"{stats['average']:.2f}", f"{stats['deduction']:.2f}"])
    _add_simple_table(document, ["序号", "考核对象类别", "记录数", "平均得分", "累计扣分"], type_rows)

    _add_paragraph(document, "从评分结果看，本期报告按镇级污水处理厂、镇区污水收集管网、农村污水处理设施分别统计。各类考核对象的具体评分依据和扣分条目详见附件2绩效考核评分表。")
    _add_paragraph(document, "对得分较低或扣分较集中的项目点，应结合现场照片、水质抽检、运行台账和整改记录逐项复核，形成下一周期重点跟踪清单。")


def _score_stats(records: list[dict[str, Any]]) -> dict[str, float]:
    scores = [float(item.get("totalScore") or 0) for item in records]
    deductions = [float(row[4] or 0) for row in _deduction_rows(records)]
    return {
        "count": float(len(scores)),
        "average": sum(scores) / len(scores) if scores else 0.0,
        "min": min(scores) if scores else 0.0,
        "max": max(scores) if scores else 0.0,
        "deduction": sum(deductions),
    }


def _score_level(score: float) -> str:
    if score >= 90:
        return "较好"
    if score >= 80:
        return "基本达标"
    if score >= 70:
        return "需持续整改"
    if score >= 60:
        return "问题较多"
    return "不达标"


def _project_report_type(project_name: str, is_summary: bool) -> str:
    if "茂南" in project_name:
        return "城镇设施全区捆绑PPP项目绩效考核报告" if is_summary else "城镇设施绩效考核报告"
    return "镇村污水处理设施绩效考核报告" if is_summary else "镇村污水处理设施绩效考核报告"


def _add_record_results(document, records: list[dict[str, Any]]) -> None:
    document.add_heading("考核结果", level=1)
    summary = document.add_table(rows=1, cols=6)
    summary.style = "Table Grid"
    for cell, value in zip(summary.rows[0].cells, ["序号", "考核类型", "行政村", "设施点", "状态", "得分"]):
        _set_cell_text(cell, value, bold=True)
    _mark_table_header(summary.rows[0])
    for index, record in enumerate(records, 1):
        raw_type = record.get("facilityType") or record.get("rawFacilityType")
        row = summary.add_row().cells
        for cell, value in zip(row, [index, REPORT_TYPE_LABELS.get(raw_type, raw_type or "-"), record.get("administrativeVillage") or "-", record.get("village") or "-", record.get("status"), f"{float(record.get('totalScore') or 0):.2f}"]):
            _set_cell_text(cell, value)

    deductions = []
    for record in records:
        for score in record.get("scores") or []:
            if float(score.get("deduction") or 0) > 0:
                deductions.append((record, score))
    document.add_heading("扣分明细", level=2)
    if not deductions:
        _add_paragraph(document, "经核查，本期未发现需扣分的问题。")
        return
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ["序号", "设施点", "评分条目", "满分", "扣分", "依据或说明"]):
        _set_cell_text(cell, value, bold=True)
    _mark_table_header(table.rows[0])
    for index, (record, score) in enumerate(deductions, 1):
        row = table.add_row().cells
        for cell, value in zip(row, [index, record.get("village") or record.get("town"), score.get("indicatorName"), score.get("indicatorFullScore"), score.get("deduction"), score.get("reason") or score.get("deductionOptionName")]):
            _set_cell_text(cell, value)


def _generate_project_reports(task: ReportTask, snapshot: dict[str, Any]) -> Path:
    output_dir = _storage_root() / "generated_reports" / "working" / task.id
    output_dir.mkdir(parents=True, exist_ok=True)
    project_name = snapshot.get("projectName") or "项目"
    cycle_name = snapshot.get("cycleName") or task.payload.get("period") or "本期"
    outputs = set(task.payload.get("outputs", []))
    requested_towns = set(task.payload.get("townNames", []) or [])
    if "separate" in outputs:
        for town_data in snapshot.get("towns") or []:
            town_name = town_data["town"]
            if requested_towns and town_name not in requested_towns:
                continue
            records = [item for item in snapshot.get("records") or [] if item.get("town") == town_name]
            if not records:
                continue
            profile = PROJECT_REPORT_PROFILES.get(project_name) or PROJECT_REPORT_PROFILES["\u90c1\u5357\u9879\u76ee"]
            document = _generate_town_document(project_name, cycle_name, town_data, records)
            document.save(output_dir / f"{town_name}-{cycle_name}-{profile['shortName']}{profile['titleSuffix']}\uff08\u6b63\u6587\uff09.docx")

    if "summary" in task.payload.get("outputs", []):
        profile = PROJECT_REPORT_PROFILES.get(project_name) or PROJECT_REPORT_PROFILES["郁南项目"]
        document = _generate_summary_document(project_name, cycle_name, snapshot)
        document.save(output_dir / f"{project_name}-{cycle_name}-{profile['shortName']}{profile['titleSuffix']}汇总报告.docx")
    return output_dir


def _append_database_summary(session: Session, paths: list[Path], records: list[AssessmentRecord]) -> None:
    """Keep a traceable database-derived summary inside each generated DOCX."""
    from docx import Document

    by_town: dict[str, list[AssessmentRecord]] = {}
    for record in records:
        by_town.setdefault(record.town.name, []).append(record)
    for path in paths:
        target_town = path.name.split("2023")[0]
        selected = records if target_town not in by_town else by_town[target_town]
        if not selected:
            continue
        document = Document(path)
        document.add_heading("考核数据复核摘要", level=2)
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ["镇街", "已复核记录", "状态", "评分条目", "问卷", "水质", "照片"]):
            _set_cell_text(cell, text, bold=True)
        _mark_table_header(table.rows[0])
        for town, items in by_town.items():
            record_ids = [item.id for item in items]
            row = table.add_row().cells
            row[0].text = town
            row[1].text = str(len(items))
            row[2].text = "、".join(sorted({item.status for item in items}))
            row[3].text = str(sum(len(item.scores) for item in items))
            row[4].text = str(session.scalar(select(func.count(SurveyRecord.id)).where(SurveyRecord.record_id.in_(record_ids))) or 0)
            row[5].text = str(session.scalar(select(func.count(WaterQualityRecord.id)).where(WaterQualityRecord.record_id.in_(record_ids))) or 0)
            row[6].text = str(session.scalar(select(func.count(Attachment.id)).where(Attachment.record_id.in_(record_ids))) or 0)
        _enforce_document_fonts(document)
        document.save(path)


def _storage_root() -> Path:
    if settings.storage_dir.is_absolute():
        return settings.storage_dir
    return settings.backend_dir / settings.storage_dir


def _next_report_version(session: Session, *, name: str, cycle_id: str | None, town_id: str | None) -> int:
    existing = session.scalars(
        select(Report).where(
            Report.name == name,
            Report.cycle_id == cycle_id,
            Report.town_id == town_id,
        )
    ).all()
    return max([report.version or 1 for report in existing], default=0) + 1


def _versioned_report_path(task_id: str, version: int, source: Path) -> Path:
    output_dir = _storage_root() / "generated_reports" / "tasks" / task_id / f"v{version:03d}"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / source.name


def run_report_task(task_id: str) -> None:
    with SessionLocal() as session:
        task = session.get(ReportTask, task_id)
        if task is None:
            return
        task.status, task.progress, task.started_at, task.error = "running", 10, utcnow(), None
        session.commit()
        try:
            town_names = set(task.payload.get("townNames", []))
            outputs = set(task.payload.get("outputs", []))
            include_separate = "separate" in outputs
            include_summary = "summary" in outputs
            project_id = task.payload.get("projectId")
            cycle = session.get(AssessmentCycle, task.cycle_id) if task.cycle_id else None
            record_query = select(AssessmentRecord).where(AssessmentRecord.status.in_(["submitted", "reviewed", "locked"]))
            if project_id:
                record_query = record_query.where(AssessmentRecord.city_id == project_id)
            if task.cycle_id:
                record_query = record_query.where(AssessmentRecord.cycle_id == task.cycle_id)
            records = list(session.scalars(record_query))
            if town_names and not include_summary:
                records = [record for record in records if record.town.name in town_names]
            dataset_town_names = None if include_summary else (town_names or None)
            snapshot = build_report_dataset(session, cycle=cycle, town_names=dataset_town_names, city_id=project_id)
            if task.payload.get("source") in {"dashboard", "mobile"}:
                validate_report_dataset(snapshot)
            task.data_snapshot = snapshot
            task.dataset_hash = snapshot.get("hash")
            session.commit()
            output_dir = _generate_project_reports(task, snapshot)
            task.progress = 80
            names = town_names
            output_paths = []
            for path in output_dir.glob("*.docx"):
                report_town = path.stem.split("-", 1)[0]
                is_summary = path.stem.endswith("汇总报告")
                if is_summary:
                    if not include_summary:
                        continue
                else:
                    if not include_separate:
                        continue
                    if names and report_town not in names:
                        continue
                if names and report_town not in names and not (include_summary and is_summary):
                    continue
                output_paths.append(path)
            if not output_paths:
                raise RuntimeError("正式报告生成器没有产出匹配的 DOCX 文件。")
            task.progress = 90
            for path in output_paths:
                report_town = path.stem.split("-", 1)[0]
                town_query = select(Town).where(Town.name == report_town)
                if project_id:
                    town_query = town_query.where(Town.city_id == project_id)
                town = session.scalar(town_query)
                town_id = town.id if town else None
                version = _next_report_version(session, name=path.name, cycle_id=task.cycle_id, town_id=town_id)
                final_path = _versioned_report_path(task.id, version, path)
                shutil.copy2(path, final_path)
                town_records = [item for item in snapshot.get("records", []) if item.get("town") == report_town]
                if path.stem.endswith("汇总报告"):
                    town_records = snapshot.get("records", [])
                report_snapshot = {
                    "hash": task.dataset_hash,
                    "cycleId": snapshot.get("cycleId"),
                    "cycleName": snapshot.get("cycleName"),
                    "projectId": snapshot.get("projectId"),
                    "projectName": snapshot.get("projectName"),
                    "town": report_town,
                    "recordIds": [item["id"] for item in town_records],
                    "indicatorVersionIds": sorted({item["indicatorVersionId"] for item in town_records if item.get("indicatorVersionId")}),
                    "towns": snapshot.get("towns", []),
                }
                session.add(
                    Report(
                        task_id=task.id,
                        town_id=town_id,
                        cycle_id=task.cycle_id,
                        name=path.name,
                        storage_key=str(final_path),
                        size=final_path.stat().st_size,
                        version=version,
                        format="docx",
                        dataset_hash=task.dataset_hash,
                        data_snapshot=report_snapshot,
                        task_parameters=task.payload,
                    )
                )
            try:
                from app.services.agent import create_report_task_agent_run

                create_report_task_agent_run(session, task)
            except Exception as agent_exc:
                task.payload = {**task.payload, "agentWarning": str(agent_exc)}
            task.status, task.progress, task.completed_at = "completed", 100, utcnow()
        except Exception as exc:
            task.status = "failed"
            task.error = f"{exc}\n{traceback.format_exc(limit=5)}"
        session.commit()


def _add_source_toc(document, rows: list[list[Any]]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document.add_page_break()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_paragraph_format(title, indent=False)
    title_run = title.add_run("目录")
    _apply_run_font(title_run, REPORT_HEADING_FONT, 16, bold=True)
    toc = document.add_paragraph()
    _apply_paragraph_format(toc, indent=False)
    run = toc.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    document.add_page_break()


def _is_maonan_project(project_name: str) -> bool:
    return "茂南" in project_name or "鑼傚崡" in project_name


def _add_example_toc(document, project_name: str, *, is_summary: bool) -> None:
    _add_source_toc(document, _source_toc_rows(project_name, is_summary=is_summary))


def _town_chapter_code(town_data: dict[str, Any] | None, index: int) -> str:
    code = (town_data or {}).get("chapterCode")
    return str(code) if code else f"2.{index + 1}"


def _chapter_sort_key(town_data: dict[str, Any] | None, fallback_index: int) -> tuple[int, ...]:
    code = (town_data or {}).get("chapterCode")
    if not code:
        return (9999, fallback_index)
    parts: list[int] = []
    for part in str(code).split("."):
        try:
            parts.append(int(part))
        except ValueError:
            parts.append(9999)
    parts.append(fallback_index)
    return tuple(parts)


def _ordered_towns(towns: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    return [
        town
        for index, town in sorted(
            enumerate(towns or []),
            key=lambda item: _chapter_sort_key(item[1], item[0]),
        )
    ]


def _records_for_town(records: list[dict[str, Any]], town_name: str) -> list[dict[str, Any]]:
    return [record for record in records if record.get("town") == town_name]


def _section_records_for_town(records: list[dict[str, Any]], town_name: str, *, is_summary: bool, has_catalog_towns: bool) -> list[dict[str, Any]]:
    town_records = _records_for_town(records, town_name)
    if town_records:
        return town_records
    return [] if is_summary or has_catalog_towns else records


def _record_type_rows(records: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for index, record in enumerate(records, 1):
        rows.append([
            index,
            _record_point_name(record),
            STATUS_LABELS.get(record.get("status"), record.get("status") or "-"),
            f"{float(record.get('totalScore') or 0):.2f}",
        ])
    return rows


def _add_facility_result_table(document, records: list[dict[str, Any]]) -> None:
    if not records:
        _add_paragraph(document, "本期考核范围未包含该类考核对象。")
        return
    _add_simple_table(document, ["序号", "项目点", "状态", "得分"], _record_type_rows(records))
    deductions = _deduction_rows(records)
    if deductions:
        _add_simple_table(document, ["序号", "设施点", "评分条目", "满分", "扣分", "依据或说明"], deductions)
    else:
        _add_paragraph(document, "本类考核对象未形成扣分记录。")


def _add_yunan_facility_sections(document, *, records: list[dict[str, Any]], towns: list[dict[str, Any]], is_summary: bool) -> None:
    document.add_heading("第二章 镇级设施运维考核情况", level=1)
    if is_summary:
        document.add_heading("2.1 项目考核结果汇总", level=2)
        _add_yunan_summary_score_overview(document, records, towns)

    town_items = _ordered_towns(towns) or [{"town": records[0].get("town") if records else "本镇", "assessmentTargets": _facility_types(records)}]
    for index, town_data in enumerate(town_items, 1):
        town_name = town_data.get("town") or "本镇"
        town_records = _section_records_for_town(records, town_name, is_summary=is_summary, has_catalog_towns=bool(towns))
        code = f"2.{index + 1}" if is_summary else f"2.{index}"
        document.add_heading(f"{code} {town_name}污水处理设施考核情况", level=2)
        document.add_heading(f"{code}.1 总体评价和设施情况概览", level=3)
        _add_town_basic_intro(document, town_name=town_name, town_data=town_data, records=town_records, add_heading=False)
        if not town_records:
            _add_paragraph(document, f"{town_name}本期尚无经复核或锁定的考核数据，报告仅列示项目基础信息，不纳入本期平均得分、扣分合计和绩效系数统计。")
            continue
        stats = _score_stats(town_records)
        _add_paragraph(document, f"{town_name}本期纳入考核记录{int(stats['count'])}项，平均得分{stats['average']:.2f}分，累计扣分{stats['deduction']:.2f}分。")
        type_map = {
            "town_plant": "污水处理厂运维考核情况",
            "town_network": "污水收集管网运维考核情况",
            "rural_treatment": "农村污水处理设施考核情况",
        }
        present_types = [item for item in ["town_plant", "town_network", "rural_treatment"] if any(record.get("facilityType") == item for record in town_records)]
        for offset, facility_type in enumerate(present_types, 2):
            document.add_heading(f"{code}.{offset} {type_map[facility_type]}", level=3)
            typed_records = [record for record in town_records if record.get("facilityType") == facility_type]
            _add_facility_result_table(document, typed_records)
            coefficient = town_average_coefficient(
                [float(record.get("totalScore") or 0) for record in typed_records],
                project="yunan",
            )
            if coefficient is not None:
                _add_paragraph(document, f"按郁南项目绩效付费口径，考核得分不低于90分时运维绩效考核系数取1，低于90分时按得分除以90计算。本项运维绩效考核系数为{coefficient:.3f}。")


def _add_maonan_result_sections(document, *, records: list[dict[str, Any]], towns: list[dict[str, Any]], is_summary: bool) -> None:
    document.add_heading("第二章 城镇水质净化设施考核结果", level=1)
    if is_summary:
        document.add_heading("2.1 项目考核结果汇总", level=2)
        _add_yunan_summary_score_overview(document, records, towns)

    town_items = _ordered_towns(towns) or [{"town": records[0].get("town") if records else "本镇", "assessmentTargets": _facility_types(records)}]
    for index, town_data in enumerate(town_items, 1):
        town_name = town_data.get("town") or "本镇"
        town_records = _section_records_for_town(records, town_name, is_summary=is_summary, has_catalog_towns=bool(towns))
        code = f"2.{index + 1}" if is_summary else f"2.{index}"
        document.add_heading(f"{code} {town_name}考核结果", level=2)
        document.add_heading(f"{code}.1 总体评价和设施情况概览", level=3)
        _add_town_basic_intro(document, town_name=town_name, town_data=town_data, records=town_records, add_heading=False)
        if not town_records:
            _add_paragraph(document, f"{town_name}本期尚无经复核或锁定的考核数据，报告仅列示项目基础信息，不纳入本期平均得分、扣分合计和绩效系数统计。")
            continue
        _add_paragraph(document, f"经汇总，{town_name}本期平均得分为{_score_stats(town_records)['average']:.2f}分。")
        section_titles = {
            "town_plant": "水质净化厂运维考核情况",
            "town_network": "污水收集管网运维考核情况",
        }
        present_types = [item for item in ["town_plant", "town_network"] if any(record.get("facilityType") == item for record in town_records)]
        for offset, facility_type in enumerate(present_types, 2):
            document.add_heading(f"{code}.{offset} {section_titles[facility_type]}", level=3)
            _add_facility_result_table(document, [record for record in town_records if record.get("facilityType") == facility_type])


def _add_source_chapter_one(document, *, project_name: str, cycle_name: str, scope_name: str, profile: dict[str, Any]) -> None:
    document.add_heading("第一章 考核工作概述", level=1)
    headings = [
        ("1.1 考核目的", f"本次绩效考核旨在依据项目合同、补充协议及绩效考核标准，评价{scope_name}{cycle_name}污水处理设施运行维护、资料管理、水质达标和整改落实情况。"),
        ("1.2 考核要求", "考核工作按照合同约定、评分标准和现场核查要求开展，评分结果以已提交、已复核或已锁定的数据为依据。"),
        ("1.3 考核依据", profile["basis"]),
        ("1.4 考核成员及分工责任", "考核成员及分工按委托单位、考核单位、检测单位和项目公司确认的现场安排执行。"),
        ("1.5 考核频次", f"本报告对应{cycle_name}绩效考核周期。"),
        ("1.7 考核时间", f"本次考核时间以{cycle_name}现场检查及资料复核工作安排为准。"),
    ]
    for heading, body in headings[:-1]:
        document.add_heading(heading, level=2)
        _add_paragraph(document, body)
    document.add_heading("1.6 考核方法", level=2)
    method_details = [
        ("1.6.1 现场检查", "考核组对设施运行状态、厂区及站点环境、设备维护、安全防护、管网及检查井状况进行现场核查，并对发现的问题形成记录。"),
        ("1.6.2 查阅资料", "对生产运行、巡查巡检、维护维修、污泥处置、人员培训、安全管理和问题整改等资料进行核对，重点检查记录的完整性、连续性和可追溯性。"),
    ]
    has_survey = any(str(method).startswith("问卷调查") for method in profile["methods"])
    if has_survey:
        method_details.append(("1.6.3 问卷调查", "按照项目考核安排收集实施机构、镇街代表及群众意见，调查结果用于评价污水处理服务成效和问题整改情况。"))
    water_index = 4 if has_survey else 3
    method_details.extend([
        (f"1.6.{water_index} 水质检测", "依据本项目适用的排放标准开展取样检测，将实测结果与规定限值逐项比较；出现人工修正结论时，同时核查修正说明和支撑资料。"),
        (f"1.6.{water_index + 1} 考核评分及系数计算", "考核组按照现行绩效考核标准逐项评分。扣分事项须对应具体评分条目、扣分依据和现场或资料证据，并据考核得分计算运维绩效考核系数。"),
    ])
    for heading, body in method_details:
        document.add_heading(heading, level=3)
        _add_paragraph(document, body)
    document.add_heading(headings[-1][0], level=2)
    _add_paragraph(document, headings[-1][1])


def _add_source_appendices(document, *, project_name: str, records: list[dict[str, Any]], profile: dict[str, Any]) -> None:
    maonan = _is_maonan_project(project_name)
    appendix_titles = [
        "附件1 考核标准",
        "附件2 周期评分表" if maonan else "附件2 考核评分表",
        "附件3 现场检查照片" if maonan else "附件3 现场照片",
        "附件4 水质净化厂处理水量" if maonan else "附件4 各镇污水处理厂处理水量",
        "附件5 水质抽检汇总" if maonan else "附件5 水质抽检情况汇总表",
        "附件6 检测报告" if maonan else "附件6 公众调查情况",
        "附件7 资料清单" if maonan else "附件7 考核资料清单目录",
        "附件8 月平均值统计" if maonan else "附件8 抽检水质检测报告",
    ]
    for title in appendix_titles:
        document.add_heading(title, level=1)
        if "考核标准" in title:
            overview = _standards_overview_rows(project_name, records)
            _add_simple_table(document, ["序号", "考核对象", "一级指标数", "评分项数", "满分"], overview)
            standards = _project_standards(project_name)
            for facility_type in standards:
                rows = _standard_rows_for_type(project_name, facility_type)
                if not rows:
                    continue
                document.add_heading(REPORT_TYPE_LABELS.get(facility_type, facility_type), level=2)
                _add_simple_table(
                    document,
                    ["序号", "一级指标", "二级指标", "评分条目", "满分", "扣分办法", "评价标准"],
                    rows,
                )
        elif "评分表" in title:
            score_rows = _all_score_rows(records)
            if score_rows:
                _add_simple_table(document, ["序号", "项目点", "评分条目", "满分", "实得分", "扣分", "核查情况"], score_rows)
            else:
                _add_paragraph(document, "本期暂无可列示的评分明细。")
        elif "照片" in title:
            rows = _attachment_rows(records, detection=False)
            if rows:
                inserted, skipped = _add_attachment_pictures(document, records, detection=False)
                if inserted == 0:
                    _add_paragraph(document, "本期附件中没有可装订的有效现场照片。")
                elif skipped:
                    _add_paragraph(document, f"已装订{inserted}张现场图片；另有{skipped}张图片文件无法读取，请核对原始附件。")
            else:
                _add_paragraph(document, "本期未上传现场照片。")
        elif "处理水量" in title:
            rows = _operation_volume_rows(records)
            if rows:
                _add_simple_table(document, ["序号", "项目点", "设计规模（m³/d）", "实际处理水量", "统计周期"], rows)
            else:
                _add_paragraph(document, "本期资料未包含可用于付费核定的实际处理水量记录，待补充经确认的计量资料后列示。")
        elif "水质抽检" in title:
            _add_paragraph(document, profile["waterStandard"])
            rows = _water_quality_rows(records)
            if rows:
                _add_simple_table(document, ["序号", "项目点", "取样时间", "检测指标", "实测值", "限值", "单位", "自动判定", "最终判定", "备注"], rows)
            else:
                _add_paragraph(document, "本期暂无已提交的水质抽检结果。")
            _add_simple_table(document, ["序号", "对象", "指标", "限值", "单位"], [[index, *row] for index, row in enumerate(profile["waterRows"], 1)])
        elif "公众调查" in title:
            rows = _survey_rows(records)
            if rows:
                _add_simple_table(document, ["序号", "项目点", "调查类型", "对象", "得分"], rows)
            else:
                _add_paragraph(document, "本期未开展公众调查或暂无相关资料。")
        elif "检测报告" in title:
            rows = _attachment_rows(records, detection=True)
            if rows:
                _add_simple_table(document, ["序号", "项目点", "文件名", "评分记录", "扣分项", "大小"], rows)
                inserted, skipped = _add_attachment_pictures(document, records, detection=True)
                if inserted:
                    _add_paragraph(document, f"本附件共装订{inserted}页检测资料。")
                if skipped:
                    _add_paragraph(document, f"另有{skipped}份检测资料无法转换或读取，请核对原始文件。")
            else:
                _add_paragraph(document, "本期未提交可装订的水质检测报告，水质抽检数据以附件5汇总表为准。")
        elif "月平均值" in title:
            rows = _monthly_water_average_rows(records)
            if rows:
                _add_simple_table(document, ["序号", "项目点", "统计月份", "平均进水CODCr", "平均出水CODCr", "核定说明"], rows)
            else:
                _add_paragraph(document, "本期资料未包含完整的月平均进、出水CODCr统计值，因此暂不计算月度水质浓度系数。")
        elif "资料清单" in title:
            _add_simple_table(document, ["序号", "资料类别", "资料内容", "备注"], [
                [1, "现场考核资料", "现场检查记录、照片、签字确认资料", "按实际提交资料归档"],
                [2, "运行维护资料", "运行台账、巡查记录、维修记录、污泥处置资料", "按评分标准复核"],
                [3, "水质检测资料", "水质抽检结果、检测报告、限值依据", "按检测报告归档"],
            ])
        else:
            _add_paragraph(document, "本期资料暂未包含该项内容，具体以本期正式提交并经业主确认的资料为准。")


def _generate_town_document(project_name: str, cycle_name: str, town_data: dict[str, Any], records: list[dict[str, Any]]):
    profile = PROJECT_REPORT_PROFILES.get(project_name) or PROJECT_REPORT_PROFILES.get("郁南项目") or next(iter(PROJECT_REPORT_PROFILES.values()))
    town_name = town_data["town"]
    title = f"{project_name}{town_name}{cycle_name}{profile['titleSuffix']}"
    document = _prepare_document(title)
    maonan = _is_maonan_project(project_name)
    report_type = "城镇设施绩效考核报告" if maonan else "镇级及农村设施考核报告"
    _add_source_cover(document, project_name=project_name, cycle_name=cycle_name, title=title, report_type=report_type)
    _add_assessment_summary_text(document, project_name=project_name, cycle_name=cycle_name, records=records, towns=[town_data], profile=profile)
    _add_example_toc(document, project_name, is_summary=False)
    _add_source_chapter_one(document, project_name=project_name, cycle_name=cycle_name, scope_name=town_name, profile=profile)
    if maonan:
        _add_maonan_result_sections(document, records=records, towns=[town_data], is_summary=False)
        _add_maonan_payment_chapter(document, records)
        _add_problem_and_suggestion_chapter(document, records, maonan=True, scope_name=town_name, is_summary=False)
    else:
        _add_yunan_facility_sections(document, records=records, towns=[town_data], is_summary=False)
        _add_yunan_coefficient_chapter(document, records)
        _add_problem_and_suggestion_chapter(document, records, maonan=False, scope_name=town_name, is_summary=False)
    _add_source_appendices(document, project_name=project_name, records=records, profile=profile)
    _enforce_document_fonts(document)
    return document


def _generate_summary_document(project_name: str, cycle_name: str, snapshot: dict[str, Any]):
    profile = PROJECT_REPORT_PROFILES.get(project_name) or PROJECT_REPORT_PROFILES.get("郁南项目") or next(iter(PROJECT_REPORT_PROFILES.values()))
    title = f"{project_name}{cycle_name}{profile['titleSuffix']}汇总报告"
    document = _prepare_document(title)
    records = snapshot.get("records") or []
    towns = snapshot.get("towns") or []
    maonan = _is_maonan_project(project_name)
    report_type = "城镇设施绩效考核报告" if maonan else "镇级及农村设施考核报告"
    _add_source_cover(document, project_name=project_name, cycle_name=cycle_name, title=title, report_type=report_type)
    _add_assessment_summary_text(document, project_name=project_name, cycle_name=cycle_name, records=records, towns=towns, profile=profile)
    _add_example_toc(document, project_name, is_summary=True)
    _add_source_chapter_one(document, project_name=project_name, cycle_name=cycle_name, scope_name=f"{project_name}全部项目", profile=profile)
    if maonan:
        _add_maonan_result_sections(document, records=records, towns=towns, is_summary=True)
        _add_maonan_payment_chapter(document, records)
        _add_problem_and_suggestion_chapter(document, records, maonan=True, scope_name=project_name, is_summary=True)
    else:
        _add_yunan_facility_sections(document, records=records, towns=towns, is_summary=True)
        _add_yunan_coefficient_chapter(document, records)
        _add_problem_and_suggestion_chapter(document, records, maonan=False, scope_name=project_name, is_summary=True)
    _add_source_appendices(document, project_name=project_name, records=records, profile=profile)
    _enforce_document_fonts(document)
    return document




def _source_toc_rows(project_name: str, *, is_summary: bool) -> list[list[Any]]:
    if _is_maonan_project(project_name):
        return [
            [1, "第一章", "考核工作概述"],
            [2, "第二章", "城镇水质净化设施考核结果"],
            [3, "第三章", "绩效付费计算"],
            [4, "第四章", "主要改进点、主要问题和整改工作建议"],
            [5, "附件1", "考核标准"],
            [6, "附件2", "周期评分表"],
            [7, "附件3", "现场检查照片"],
            [8, "附件4", "水质净化厂处理水量"],
            [9, "附件5", "水质抽检汇总"],
            [10, "附件6", "检测报告"],
            [11, "附件7", "资料清单"],
            [12, "附件8", "月平均值统计"],
        ]
    return [
        [1, "第一章", "考核工作概述"],
        [2, "第二章", "镇级设施运维考核情况"],
        [3, "第三章", "考核评价系数的确定"],
        [4, "第四章", "主要问题及整改建议"],
        [5, "附件1", "考核标准"],
        [6, "附件2", "考核评分表"],
        [7, "附件3", "现场照片"],
        [8, "附件4", "各镇污水处理厂处理水量"],
        [9, "附件5", "水质抽检情况汇总表"],
        [10, "附件6", "公众调查情况"],
        [11, "附件7", "考核资料清单目录"],
        [12, "附件8", "抽检水质检测报告"],
    ]


def _add_yunan_coefficient_chapter(document, records: list[dict[str, Any]]) -> None:
    document.add_heading("第三章 绩效付费计算", level=1)
    document.add_heading("3.1 运维绩效考核系数", level=2)
    _add_paragraph(document, "根据郁南项目绩效付费口径，考核得分W达到90分及以上时，运维绩效考核系数E2取1；考核得分低于90分时，按W/90计算。本期付费采用上一考核周期已提交并经复核的数据，本期形成的考核得分作为下一付费周期的系数来源。")
    rows = []
    for index, record in enumerate(records, 1):
        score = float(record.get("totalScore") or 0)
        context = _payment_context(record)
        applied, source = _payment_coefficient(record, project="yunan")
        calculated = yunan_operation_coefficient(score)
        calculation = f"W={score:.2f}≥90，E2=1" if score >= 90 else f"E2=W/90={score:.2f}/90={calculated:.4f}"
        rows.append([
            index,
            _record_point_name(record),
            REPORT_TYPE_LABELS.get(record.get("facilityType"), record.get("facilityType") or "-"),
            f"{score:.2f}",
            f"{calculated:.3f}",
            calculation,
            context.get("currentScoreAppliesTo") or "下一付费周期",
            f"{applied:.3f}" if applied is not None else "暂不核定",
            source,
        ])
    _add_simple_table(document, ["序号", "考核对象", "设施类型", "本期得分W", "本期得分折算E2", "折算过程", "本期得分适用周期", "本期付费采用E2", "采用依据"], rows)

    document.add_heading("3.2 水质浓度系数", level=2)
    _add_paragraph(document, "水质浓度系数Kq根据进水CODCr、出水CODCr及出水达标情况确定。进水CODCr低于140mg/L时，按进出水浓度差折算；进水CODCr达到140mg/L但出水不达标时，按项目既有付费口径取0.9；其余区间按郁南项目既有付费资料执行。")
    _add_paragraph(document, "镇区污水处理设施处理水量按附件18约定设置1.1倍设计规模上限；镇区管网负荷系数KQ按实际日均水量与设计规模折算，最高取1。县城区管网按附件18县城区管网基数列示，需补充售水量或经确认的负荷基数后核定KQ。")
    _add_paragraph(document, "镇区污水处理设施月服务费公式：PCi=P0×QB×（3/4×Kq+1/4×E2）。其中，P0为处理单价（元/立方米），QB为核定月处理水量（万立方米），计算结果单位为万元。")
    _add_paragraph(document, "镇区污水收集管网月服务费公式：Pwi=P0/12×（5/6×KQ×Kq+1/6×E1×E2）。其中，P0为管网年服务费基数（万元/年），KQ为负荷系数，Kq为水质浓度系数，E1为枯水期进水水质系数，E2为运维绩效考核系数。")
    _add_paragraph(document, payment_source_summary("郁南项目"))

    document.add_heading("3.3 金额基础表", level=2)
    selected_town_names = {_record_point_name(record) for record in records if record.get("facilityType") in {"town_plant", "town_network"}}
    selected_rural_towns = {str(record.get("town") or "") for record in records if record.get("facilityType") == "rural_treatment"}
    town_rows = []
    for row in yunan_town_payment_basis_rows():
        if row["pointName"] in selected_town_names:
            town_rows.append([
                len(town_rows) + 1,
                row["pointName"],
                "-" if row.get("treatmentOperationUnitPriceYuanPerCubicMeter") is None else f"{float(row['treatmentOperationUnitPriceYuanPerCubicMeter']):.2f}",
                "-" if row.get("designScaleCubicMetersPerDay") is None else f"{float(row['designScaleCubicMetersPerDay']):.0f}",
                f"{float(row['networkAvailabilityFeeTenThousandYuanPerYear']):.2f}",
                f"{float(row['networkOperationFeeTenThousandYuanPerYear']):.2f}",
            ])
    if any(str(record.get("town") or "") == "都城镇" and record.get("facilityType") == "town_network" for record in records):
        county = yunan_county_network_basis()
        town_rows.append([
            len(town_rows) + 1,
            "县城区",
            "-",
            "-",
            f"{float(county.get('networkAvailabilityFeeTenThousandYuanPerYear') or 0):.2f}",
            f"{float(county.get('networkOperationFeeTenThousandYuanPerYear') or 0):.2f}",
        ])
    if town_rows:
        _add_simple_table(document, ["序号", "镇区项目点", "处理单价（元/m³）", "设计规模（m³/d）", "管网可用性付费（万元/年）", "管网运维养护费（万元/年）"], town_rows)
    rural_rows = []
    for row in yunan_rural_payment_basis_rows():
        if row["pointName"] in selected_rural_towns:
            rural_rows.append([
                len(rural_rows) + 1,
                row["pointName"],
                f"{float(row['availabilityFeeTenThousandYuanPerYear']):.2f}",
                f"{float(row['operationFeeTenThousandYuanPerYear']):.2f}",
            ])
    if rural_rows:
        _add_simple_table(document, ["序号", "镇街", "农村设施可用性付费（万元/年）", "农村设施运维养护费（万元/年）"], rural_rows)

    document.add_heading("3.4 本期付费测算", level=2)
    payment_rows = []
    for record in records:
        point_name = _record_point_name(record)
        facility_type = record.get("facilityType") or ""
        context = _payment_context(record)
        coefficient, coefficient_source = _payment_coefficient(record, project="yunan")
        months = _payment_month_rows(record)
        basis = context.get("basis") if isinstance(context.get("basis"), dict) else None
        if not basis and facility_type in {"town_plant", "town_network"}:
            basis = yunan_town_payment_basis_for_point(point_name)
        if not basis and facility_type == "town_network" and str(record.get("town") or "") == "都城镇":
            basis = yunan_county_network_basis()

        if facility_type == "rural_treatment":
            payment_rows.append([
                len(payment_rows) + 1,
                point_name,
                REPORT_TYPE_LABELS.get(facility_type, facility_type),
                "-",
                "-",
                "农村设施按全镇设施汇总Kq、KQ和E2核定，单个村点记录不直接形成镇级服务费",
                "待全镇月度汇总数据齐全后核定",
            ])
            continue

        dry_days = sum(int(_optional_float(row.get("influentCodDaysOver160")) or 0) for row in months)
        dry_e1 = yunan_dry_season_quality_coefficient(dry_days)
        design_scale = _optional_float(context.get("designScaleCubicMetersPerDay"))
        rows_for_record: list[str] = []
        for month in months:
            month_label = str(month.get("month") or "-")
            missing: list[str] = []
            kq = _month_kq(month, project="yunan")
            if kq is None:
                missing.append("月均进出水COD")
            if coefficient is None:
                missing.append("上一期E2")
            amount_text = "暂不核定"
            if facility_type == "town_plant":
                unit_price = _optional_float(context.get("adjustedTreatmentUnitPriceYuanPerCubicMeter"))
                if unit_price is None and basis:
                    unit_price = _optional_float(basis.get("treatmentOperationUnitPriceYuanPerCubicMeter"))
                volume = _month_volume(month)
                if unit_price is None:
                    missing.append("处理单价")
                if volume is None:
                    missing.append("月处理水量")
                applied_volume = volume
                if volume is not None and design_scale is not None:
                    applied_volume = bounded_monthly_volume(
                        actual_volume_ten_thousand_cubic_meters=volume,
                        design_scale_cubic_meters_per_day=design_scale,
                        month=month_label,
                        maximum_factor=1.1,
                    )["applied"]
                if not missing and unit_price is not None and applied_volume is not None and kq is not None and coefficient is not None:
                    amount = yunan_town_treatment_monthly_fee(
                        operation_unit_price=unit_price,
                        monthly_volume_ten_thousand_cubic_meters=applied_volume,
                        water_quality_coefficient=kq,
                        operation_coefficient=coefficient,
                    )
                    amount_text = (
                        "代入："
                        f"PCi={unit_price:.4f}×{applied_volume:.4f}×"
                        f"（3/4×{kq:.4f}+1/4×{coefficient:.4f}）={amount:.4f}万元"
                    )
                input_text = f"{month_label}：Kq={kq:.4f}" if kq is not None else f"{month_label}：Kq缺失"
                if volume is not None:
                    input_text += f"，QB={volume:.4f}万m³"
                    if applied_volume is not None and applied_volume != volume:
                        input_text += f"，核定QB={applied_volume:.4f}万m³"
                rows_for_record.append(f"{input_text}，{_format_missing(missing)}，{amount_text}")
            elif facility_type == "town_network":
                annual_fee = None
                if basis:
                    annual_fee = _optional_float(basis.get("networkAvailabilityFeeTenThousandYuanPerYear"))
                    operation_fee = _optional_float(context.get("adjustedNetworkOperationFeeTenThousandYuanPerYear"))
                    if operation_fee is None:
                        operation_fee = _optional_float(basis.get("networkOperationFeeTenThousandYuanPerYear"))
                    if annual_fee is not None and operation_fee is not None:
                        annual_fee += operation_fee
                average_daily = _month_average_daily_volume(month)
                if annual_fee is None:
                    missing.append("管网年服务费基数")
                if average_daily is None:
                    missing.append("日均水量")
                if design_scale is None:
                    missing.append("设计规模或确认负荷基数")
                load = None
                if average_daily is not None and design_scale is not None:
                    load = yunan_town_network_load_coefficient(average_daily, design_scale)
                if not missing and annual_fee is not None and load is not None and kq is not None and coefficient is not None:
                    amount = yunan_network_monthly_fee(
                        annual_network_fee_ten_thousand_yuan=annual_fee,
                        load_coefficient=load,
                        water_quality_coefficient=kq,
                        dry_season_quality_coefficient=dry_e1,
                        operation_coefficient=coefficient,
                    )
                    amount_text = (
                        "代入："
                        f"Pwi={annual_fee:.4f}/12×"
                        f"（5/6×{load:.4f}×{kq:.4f}+1/6×{dry_e1:.4f}×{coefficient:.4f}）"
                        f"={amount:.4f}万元"
                    )
                input_text = f"{month_label}：Kq={kq:.4f}" if kq is not None else f"{month_label}：Kq缺失"
                if load is not None:
                    input_text += f"，KQ={load:.4f}"
                input_text += f"，E1={dry_e1:.4f}"
                rows_for_record.append(f"{input_text}，{_format_missing(missing)}，{amount_text}")
        payment_rows.append([
            len(payment_rows) + 1,
            point_name,
            REPORT_TYPE_LABELS.get(facility_type, facility_type or "-"),
            f"{coefficient:.4f}" if coefficient is not None else "暂不核定",
            coefficient_source,
            "；".join(rows_for_record) if rows_for_record else "未录入月度付费基础数据",
            "逐月金额齐全后汇总形成应付服务费结论",
        ])
    _add_simple_table(document, ["序号", "考核对象", "设施类型", "本期采用E2", "系数依据", "月度测算", "处理意见"], payment_rows)

    document.add_heading("3.5 金额核定条件", level=2)
    _add_simple_table(document, ["序号", "必需输入", "当前处理原则", "说明"], [
        [1, "附件18金额基数", "已按郁南附件18结构化引用", "不引用茂南或其他项目金额资料"],
        [2, "上一期考核得分", "用于本期付费E2", "首个付费周期可按1执行，但需在平台端明确勾选"],
        [3, "月度水量、月均COD和达标结论", "逐月计算Kq、KQ和服务费", "缺失时只列公式和缺失条件，不输出最终应付金额"],
        [4, "农村设施全镇汇总数据", "全镇汇总后核定", "单个村点记录不替代全镇农村设施付费测算"],
    ])


def _optional_float(value: Any) -> float | None:
    if value in (None, ""):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _payment_context(record: dict[str, Any]) -> dict[str, Any]:
    value = record.get("paymentContext")
    return value if isinstance(value, dict) else {}


def _payment_month_rows(record: dict[str, Any]) -> list[dict[str, Any]]:
    context = _payment_context(record)
    rows = context.get("months")
    if isinstance(rows, list):
        return [dict(item) for item in rows if isinstance(item, dict)]
    raw = record.get("rawPayload") or {}
    payment_data = raw.get("paymentData") if isinstance(raw.get("paymentData"), dict) else {}
    rows = payment_data.get("months")
    if isinstance(rows, list):
        return [dict(item) for item in rows if isinstance(item, dict)]
    return []


def _payment_coefficient(record: dict[str, Any], *, project: str) -> tuple[float | None, str]:
    context = _payment_context(record)
    coefficient = _optional_float(context.get("appliedOperationCoefficient"))
    if coefficient is not None:
        source = context.get("coefficientSourcePeriod") or context.get("coefficientStatus") or "已确认来源"
        return coefficient, str(source)
    if context.get("firstPaymentPeriod"):
        return 1.0, "首个付费周期按1执行"
    return None, "缺少上一期已提交考核结果，暂不核定本期付费金额"


def _month_kq(row: dict[str, Any], *, project: str) -> float | None:
    influent = _optional_float(row.get("influentCod"))
    effluent = _optional_float(row.get("effluentCod"))
    if influent is None or effluent is None:
        return None
    if project == "yunan":
        return yunan_water_quality_coefficient(
            influent,
            effluent,
            effluent_qualified=bool(row.get("effluentQualified", True)),
        )
    return maonan_water_quality_coefficient(influent, effluent)


def _month_volume(row: dict[str, Any]) -> float | None:
    return _optional_float(row.get("monthlyVolumeTenThousandCubicMeters"))


def _month_average_daily_volume(row: dict[str, Any]) -> float | None:
    return _optional_float(row.get("averageDailyVolumeCubicMeters"))


def _format_missing(items: list[str]) -> str:
    return "；".join(items) if items else "资料齐全"


def _maonan_monthly_payment_inputs(record: dict[str, Any]) -> tuple[float | None, float | None]:
    raw = record.get("rawPayload") or {}
    water = raw.get("waterQuality") or {}
    if not isinstance(water, dict):
        water = {}
    influent = _optional_float(raw.get("monthlyInfluentCod"))
    if influent is None:
        influent = _optional_float(water.get("monthlyInfluentCod"))
    effluent = _optional_float(raw.get("monthlyEffluentCod"))
    if effluent is None:
        effluent = _optional_float(water.get("monthlyEffluentCod"))
    kq = maonan_water_quality_coefficient(influent, effluent) if influent is not None and effluent is not None else None

    volume = _optional_float(raw.get("monthlyVolumeTenThousandTons"))
    if volume is None:
        volume = _optional_float(water.get("monthlyVolumeTenThousandTons"))
    if volume is None:
        unit = str(raw.get("monthlyVolumeUnit") or water.get("monthlyVolumeUnit") or "").strip().lower()
        if unit in {"万吨", "万吨/月", "10k_m3", "10k_tons"}:
            volume = _optional_float(raw.get("monthlyVolume") or water.get("monthlyVolume"))
    return kq, volume


def _add_maonan_payment_chapter(document, records: list[dict[str, Any]]) -> None:
    document.add_heading("第三章 绩效付费计算", level=1)
    document.add_heading("3.1 运维考核系数", level=2)
    _add_paragraph(document, "按茂南项目绩效付费口径，考核得分W不低于70分时，运维考核系数E1取1；低于70分时，E1=W/70。本期付费采用上一考核周期已提交并经复核的数据，本期考核得分作为下一付费周期服务费测算依据。")
    coefficient_rows = []
    for index, record in enumerate(records, 1):
        score = float(record.get("totalScore") or 0)
        context = _payment_context(record)
        applied, source = _payment_coefficient(record, project="maonan")
        calculated = maonan_operation_coefficient(score)
        calculation = f"W={score:.2f}≥70，E1=1" if score >= 70 else f"E1=W/70={score:.2f}/70={calculated:.4f}"
        coefficient_rows.append([
            index,
            _record_point_name(record),
            REPORT_TYPE_LABELS.get(record.get("facilityType"), record.get("facilityType") or "-"),
            f"{score:.2f}",
            f"{calculated:.3f}",
            calculation,
            context.get("currentScoreAppliesTo") or "下一付费周期",
            f"{applied:.3f}" if applied is not None else "暂不核定",
            source,
        ])
    _add_simple_table(document, ["序号", "考核对象", "设施类型", "本期得分W", "本期得分折算E1", "折算过程", "本期得分适用周期", "本期付费采用E1", "采用依据"], coefficient_rows)

    document.add_heading("3.2 付费公式", level=2)
    _add_paragraph(document, "城镇水质净化厂月服务费公式：Pz=Py×QB×（3/5+Kq/10+3E1/10）+Pk/12×（2/3+E1/3）。其中，Py为污水处理运营服务费单价（元/立方米），QB为当月核定处理水量（万立方米），Kq为水质浓度系数，Pk为水质净化设施可用性付费（万元/年），计算结果单位为万元。")
    _add_paragraph(document, "污水收集管网月服务费公式：Pw=Pk/12×（2/3+E1/3）+Py/12×（3/5+Kq/10+3E1/10）。其中，Pk为管网可用性付费（万元/年），Py为调整后管网年运营维护费（万元/年）。项目服务费由可用性付费和运营维护费组成；未提供新金额表时沿用茂南项目既有历史表，不引用其他项目或通用金额基础表。")
    _add_paragraph(document, payment_source_summary("茂南项目"))

    document.add_heading("3.3 金额基础表", level=2)
    basis_rows = maonan_payment_basis_rows()
    selected_types: dict[str, set[str]] = {}
    for record in records:
        basis = maonan_payment_basis_for_point(_record_point_name(record))
        if basis:
            selected_types.setdefault(basis["pointName"], set()).add(record.get("facilityType") or "")
    plant_rows = []
    network_rows = []
    for row in basis_rows:
        point_types = selected_types.get(row["pointName"], set())
        if "town_plant" in point_types and row.get("treatmentAvailabilityFeeTenThousandYuanPerYear") is not None:
            plant_rows.append([
                len(plant_rows) + 1,
                row["pointName"],
                f"{float(row['treatmentAvailabilityFeeTenThousandYuanPerYear']):.2f}",
                f"{float(row['treatmentOperationUnitPriceYuanPerCubicMeter']):.2f}",
                "茂南例文表4-1",
            ])
        if "town_network" in point_types and row.get("networkAvailabilityFeeTenThousandYuanPerYear") is not None:
            network_rows.append([
                len(network_rows) + 1,
                row["pointName"],
                f"{float(row['networkAvailabilityFeeTenThousandYuanPerYear']):.2f}",
                f"{float(row['originalNetworkOperationFeeTenThousandYuanPerYear']):.2f}",
                f"{float(row['adjustedNetworkOperationFeeYuanPerYear']):.2f}",
                "茂南例文表4-1、表4-2",
            ])
    if plant_rows:
        _add_simple_table(document, ["序号", "水质净化厂", "可用性付费（万元/年）", "运营服务费单价（元/m³）", "来源"], plant_rows)
    if network_rows:
        _add_simple_table(document, ["序号", "污水收集管网", "可用性付费（万元/年）", "原运营维护费（万元/年）", "调整后运营维护费（元/年）", "来源"], network_rows)
    _add_paragraph(document, "上述金额基数在未提供新合同金额表或正式调整文件时沿用；如后续提供新表，以新确认资料覆盖本表，不改变评分和系数计算规则。")

    document.add_heading("3.4 当期付费测算", level=2)
    payment_rows = []
    for record in records:
        point_name = _record_point_name(record)
        facility_type = record.get("facilityType") or ""
        score = float(record.get("totalScore") or 0)
        e1, coefficient_source = _payment_coefficient(record, project="maonan")
        basis = maonan_payment_basis_for_point(point_name)
        basis_text = "未匹配到本项目金额基数"
        result_text = "待补齐本期输入后计算"
        if basis and facility_type == "town_plant":
            pk = float(basis["treatmentAvailabilityFeeTenThousandYuanPerYear"])
            py = float(basis["treatmentOperationUnitPriceYuanPerCubicMeter"])
            max_year = maonan_annual_maximum_treatment_fee(
                annual_availability_fee_ten_thousand_yuan=pk,
                operation_unit_price_yuan_per_cubic_meter=py,
                design_scale_cubic_meters_per_day=float(basis.get("designScaleCubicMetersPerDay") or 0),
            )
            basis_text = f"Pk={pk:.2f}万元/年；Py={py:.2f}元/m³；年封顶={max_year:.2f}万元"
        elif basis and facility_type == "town_network":
            pk = float(basis["networkAvailabilityFeeTenThousandYuanPerYear"])
            adjusted_py = float(basis["adjustedNetworkOperationFeeYuanPerYear"]) / 10000
            basis_text = f"Pk={pk:.2f}万元/年；调整后Py={adjusted_py:.4f}万元/年"
        month_lines: list[str] = []
        months = _payment_month_rows(record)
        if not months:
            legacy_kq, legacy_qb = _maonan_monthly_payment_inputs(record)
            months = [{"month": "本期", "monthlyVolumeTenThousandCubicMeters": legacy_qb, "influentCod": None, "effluentCod": None}]
            if legacy_kq is not None:
                months[0]["legacyKq"] = legacy_kq
        for month in months:
            month_label = str(month.get("month") or "-")
            kq = _optional_float(month.get("legacyKq"))
            if kq is None:
                kq = _month_kq(month, project="maonan")
            qb = _month_volume(month)
            missing: list[str] = []
            if e1 is None:
                missing.append("上一期E1")
            if kq is None:
                missing.append("月均进出水COD")
            if facility_type == "town_plant" and qb is None:
                missing.append("QB")
            amount_text = "暂不核定"
            if basis and facility_type == "town_plant":
                pk = float(basis["treatmentAvailabilityFeeTenThousandYuanPerYear"])
                py = float(basis["treatmentOperationUnitPriceYuanPerCubicMeter"])
                applied_qb = qb
                design_scale = _optional_float(basis.get("designScaleCubicMetersPerDay"))
                if qb is not None and design_scale is not None and "-" in month_label:
                    bounded = bounded_monthly_volume(
                        actual_volume_ten_thousand_cubic_meters=qb,
                        design_scale_cubic_meters_per_day=design_scale,
                        month=month_label,
                        guaranteed_factor=_optional_float(basis.get("guaranteedVolumeFactor")),
                        maximum_factor=_optional_float(basis.get("maximumVolumeFactor")) or 1.2,
                    )
                    applied_qb = bounded["applied"]
                if not missing and kq is not None and applied_qb is not None and e1 is not None:
                    amount = maonan_treatment_monthly_fee(
                        operation_unit_price=py,
                        monthly_volume_ten_thousand_tons=applied_qb,
                        water_quality_coefficient=kq,
                        operation_coefficient=e1,
                        annual_availability_fee_ten_thousand_yuan=pk,
                    )
                    amount_text = (
                        "代入："
                        f"Pz={py:.4f}×{applied_qb:.4f}×（3/5+{kq:.4f}/10+3×{e1:.4f}/10）"
                        f"+{pk:.4f}/12×（2/3+{e1:.4f}/3）={amount:.4f}万元"
                    )
                input_text = f"{month_label}：Kq={kq:.4f}" if kq is not None else f"{month_label}：Kq缺失"
                if qb is not None:
                    input_text += f"，QB={qb:.4f}万吨"
                    if applied_qb is not None and applied_qb != qb:
                        input_text += f"，核定QB={applied_qb:.4f}万吨"
                month_lines.append(f"{input_text}，{_format_missing(missing)}，{amount_text}")
            elif basis and facility_type == "town_network":
                pk = float(basis["networkAvailabilityFeeTenThousandYuanPerYear"])
                adjusted_py = float(basis["adjustedNetworkOperationFeeYuanPerYear"]) / 10000
                if not missing and kq is not None and e1 is not None:
                    amount = maonan_network_monthly_fee(
                        annual_availability_fee_ten_thousand_yuan=pk,
                        annual_operation_fee_ten_thousand_yuan=adjusted_py,
                        water_quality_coefficient=kq,
                        operation_coefficient=e1,
                    )
                    amount_text = (
                        "代入："
                        f"Pw={pk:.4f}/12×（2/3+{e1:.4f}/3）"
                        f"+{adjusted_py:.4f}/12×（3/5+{kq:.4f}/10+3×{e1:.4f}/10）"
                        f"={amount:.4f}万元"
                    )
                input_text = f"{month_label}：Kq={kq:.4f}" if kq is not None else f"{month_label}：Kq缺失"
                month_lines.append(f"{input_text}，{_format_missing(missing)}，{amount_text}")
        if month_lines:
            result_text = "；".join(month_lines)
        payment_rows.append([
            len(payment_rows) + 1,
            point_name,
            REPORT_TYPE_LABELS.get(facility_type, facility_type or "-"),
            f"{score:.2f}",
            f"{e1:.4f}" if e1 is not None else "暂不核定",
            basis_text,
            coefficient_source,
            result_text,
        ])
    _add_simple_table(document, ["序号", "考核对象", "设施类型", "本期得分W", "本期采用E1", "沿用基数", "系数依据", "月度测算结果"], payment_rows)
    _add_paragraph(document, "当期月均进出水COD、处理水量或上一期运维系数未录入时，报告保留已确认金额基数和公式，不将缺失值按0处理，也不输出虚假的最终应付金额。")

    document.add_heading("3.5 金额核定条件", level=2)
    _add_simple_table(document, ["序号", "必需输入", "当前状态", "处理原则"], [
        [1, "可用性付费及运营维护费基数", "已结构化茂南例文表4-1、表4-2", "未提供新表时沿用；新表确认后覆盖。"],
        [2, "各月实际处理水量QB", "需由本期数据或本项目历史付费资料提供", "按出水计量数据并结合合同上下限核定。"],
        [3, "各月平均进、出水COD浓度", "需由本期数据或本项目历史付费资料提供", "用于逐月计算Kq。"],
    ])
    _add_paragraph(document, "金额基数已具备时，仍需本期QB和月均COD数据才能形成完整服务费结论；资料不完整时仅输出可复核的金额基础、公式和已具备系数。")


def _add_problem_and_suggestion_chapter(
    document,
    records: list[dict[str, Any]],
    *,
    maonan: bool,
    scope_name: str,
    is_summary: bool,
) -> None:
    if maonan and is_summary:
        document.add_heading("第四章 主要改进点、主要问题和整改工作建议", level=1)
        document.add_heading("4.1 主要改进点", level=2)
        _add_paragraph(document, "本期主要改进情况根据现场复核记录、历史问题整改情况和经确认的考核资料进行归纳。")
        document.add_heading("4.2 主要问题", level=2)
    elif maonan:
        document.add_heading(f"第四章 {scope_name}主要问题和整改工作建议", level=1)
        document.add_heading("4.1 主要问题", level=2)
    else:
        document.add_heading("第四章 主要问题及整改建议", level=1)
        document.add_heading("4.1 主要问题", level=2)

    if is_summary:
        _add_paragraph(document, f"本章汇总分析{scope_name}本期所有纳入报告镇街的扣分问题和共性整改要求。")
    else:
        _add_paragraph(document, f"本章仅分析{scope_name}本期已提交并纳入报告的考核记录，不引用其他镇街的问题或结论。")

    deductions = _deduction_rows(records)
    if deductions:
        _add_deduction_narrative(document, records)
        _add_simple_table(document, ["序号", "设施点", "评分条目", "满分", "扣分", "依据或说明"], deductions)
    else:
        _add_deduction_narrative(document, records)

    water_issues = _unqualified_water_rows(records)
    if water_issues:
        document.add_heading("水质抽检异常情况", level=3)
        _add_paragraph(document, "以下仅列示自动判定或最终判定为不达标的水质抽检项目；全部达标和不达标结果详见附件5水质抽检汇总。")
        _add_simple_table(
            document,
            ["序号", "项目点", "取样时间", "检测指标", "实测值", "限值", "单位", "自动判定", "最终判定", "备注"],
            [[index, *row[1:]] for index, row in enumerate(water_issues, 1)],
        )

    document.add_heading("4.3 整改工作建议" if maonan and is_summary else "4.2 整改工作建议" if maonan else "4.2 工作建议", level=2)
    _add_yunan_suggestion_section(document, records)
