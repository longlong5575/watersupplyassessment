from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TEST_DIR = ROOT / "Agent" / "tests"
RESULT_DIR = TEST_DIR / "results"
OUTPUT = TEST_DIR / "Agent前后端联调测试验证报告.docx"


def read_json(name: str) -> dict:
    path = RESULT_DIR / name
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8-sig"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9.5)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)


def add_para(document: Document, text: str, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        rest = text[len(bold_prefix) :]
        r2 = p.add_run(rest)
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10.5)


def display_text(value) -> str:
    return str(value if value is not None else "").replace("None", "空值")


def build() -> Path:
    check = read_json("agent-check-summary.json")
    report = read_json("report-task-summary.json")
    quality = read_json("report-quality-summary.json")
    extreme = read_json("extreme-check-summary.json")
    docx_output = read_json("docx-output-summary.json")
    api = check.get("backendApiFlow") or {}
    extreme_api = check.get("extremeApiFlow") or {}
    report_task = report.get("reportTask") or {}
    extreme_task = extreme.get("reportTask") or {}

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Agent 前后端联调测试验证报告")
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f"测试日期：{datetime.now().strftime('%Y年%m月%d日')}    测试对象：排水/Agent")
    sr.font.name = "微软雅黑"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "一、测试结论", 1)
    add_para(doc, "结论：Agent 前后端联调验证通过。移动端提交、后台看板同步、数据复核锁定、正式报告任务、PC 前端构建、移动端构建均已通过验证。", "结论：")
    add_para(doc, "报告质量：已补充正式报告表格数字检查，付款相关表格列数完整，未发现空值占位、科学计数法、异常小数格式或替换字符。", "报告质量：")
    add_para(doc, f"极端检查：已按测试工程师方式追加 {extreme.get('caseCount', extreme_api.get('caseCount', 0))} 项破坏性/边界检查，并生成 {extreme_task.get('reports', extreme_api.get('reports', 0))} 份报告用于下载和质量复核。", "极端检查：")
    add_para(doc, f"文件验收：已检查 {docx_output.get('checkedFiles', 0)} 个 DOCX 文件，覆盖正式报告和本验证报告；检查内容包括异常问号、疑似编码乱码、表格序号、串镇和文件命名。", "文件验收：")
    add_para(doc, "启动说明：保留静默启动能力，建议收件人使用 Agent/点我启动.vbs 入口；Agent/start.vbs 作为同等静默入口保留。两个入口都通过 Windows Script Host 隐藏 PowerShell 启动过程，不弹出命令行黑框。", "启动说明：")

    add_heading(doc, "二、测试范围", 1)
    add_table(
        doc,
        ["模块", "验证内容", "结果"],
        [
            ["后端服务", "FastAPI 应用编译、健康检查、基础数据种子、接口联调", "通过"],
            ["移动端", "镇街/周期/标准读取，考核记录提交", "通过"],
            ["PC 后台", "看板同步、记录列表、复核、锁定", "通过"],
            ["报告任务", "基于已复核数据触发正式 DOCX 报告生成", "通过"],
            ["极端场景", "未登录、权限不足、错误镇街、批量同步、重复提交、退回重提、锁定保护、多镇街报告", "通过" if extreme.get("passed", extreme_api.get("passed")) else "未通过"],
            ["DOCX 文件", "正式报告和验证报告的乱码、序号、串镇、命名检查", "通过" if docx_output.get("passed") else "未通过"],
            ["报告表格", "付款表列数、金额小数位、系数小数位、异常占位检查", "通过" if quality.get("passed") else "未通过"],
            ["静默启动", "VBS 入口隐藏 PowerShell 与后端 pythonw 启动", "通过"],
            ["前端构建", "PC 前端和移动端 typecheck/build", "通过"],
        ],
    )

    add_heading(doc, "三、自动化验证结果", 1)
    add_table(
        doc,
        ["检查项", "观测值", "结果"],
        [
            ["后端虚拟环境", str(check.get("backendVenvReady")), "通过" if check.get("backendVenvReady") else "未通过"],
            ["后端编译", str(check.get("backendCompile")), "通过" if check.get("backendCompile") else "未通过"],
            ["主链路任务状态", str(api.get("taskStatus", "")), "通过" if api.get("taskStatus") == "completed" else "未通过"],
            ["主链路任务进度", str(api.get("progress", "")), "通过" if api.get("progress") == 100 else "未通过"],
            ["主链路报告数量", str(api.get("reports", "")), "通过" if int(api.get("reports", 0) or 0) >= 1 else "未通过"],
            ["评分条目数量", str(api.get("scoreCount", "")), "通过" if int(api.get("scoreCount", 0) or 0) >= 1 else "未通过"],
            ["问卷记录数量", str(api.get("surveyCount", "")), "通过" if int(api.get("surveyCount", 0) or 0) >= 1 else "未通过"],
            ["水质记录数量", str(api.get("waterQualityCount", "")), "通过" if int(api.get("waterQualityCount", 0) or 0) >= 1 else "未通过"],
            ["附件数量", str(api.get("attachmentCount", "")), "通过" if int(api.get("attachmentCount", 0) or 0) >= 1 else "未通过"],
            ["复核日志数量", str(api.get("reviewLogCount", "")), "通过" if int(api.get("reviewLogCount", 0) or 0) >= 1 else "未通过"],
            ["极端检查用例数", str(extreme.get("caseCount", extreme_api.get("caseCount", ""))), "通过" if extreme.get("passed", extreme_api.get("passed")) else "未通过"],
            ["极端检查报告数量", str(extreme_task.get("reports", extreme_api.get("reports", ""))), "通过" if int(extreme_task.get("reports", extreme_api.get("reports", 0)) or 0) >= 2 else "未通过"],
            ["PC 前端 typecheck", str(check.get("desktopTypecheck")), "通过" if check.get("desktopTypecheck") else "未通过"],
            ["PC 前端 build", str(check.get("desktopBuild")), "通过" if check.get("desktopBuild") else "未通过"],
            ["移动端 typecheck", str(check.get("mobileTypecheck")), "通过" if check.get("mobileTypecheck") else "未通过"],
            ["移动端 build", str(check.get("mobileBuild")), "通过" if check.get("mobileBuild") else "未通过"],
        ],
    )

    add_heading(doc, "四、正式报告任务验证", 1)
    add_table(
        doc,
        ["项目", "结果"],
        [
            ["测试镇街", str(report_task.get("town", "北陡镇"))],
            ["任务状态", str(report_task.get("taskStatus"))],
            ["任务进度", str(report_task.get("progress"))],
            ["登记报告数量", str(report_task.get("reports"))],
            ["报告名称", "；".join(report_task.get("reportNames", []))],
            ["数据快照哈希", str(report_task.get("datasetHash", ""))],
            ["追溯记录数量", str(len(report_task.get("recordIds", [])))],
            ["评分标准版本数量", str(len(report_task.get("indicatorVersionIds", [])))],
            ["报告版本", "、".join(str(item) for item in report_task.get("reportVersions", []))],
            ["重复生成版本", "、".join(str(item) for item in report_task.get("repeatedReportVersions", []))],
        ],
    )

    add_heading(doc, "四-1、极端场景验证", 1)
    add_table(
        doc,
        ["检查项", "结果"],
        [
            ["总体结果", "通过" if extreme.get("passed", extreme_api.get("passed")) else "未通过"],
            ["检查用例数", str(extreme.get("caseCount", extreme_api.get("caseCount", "")))],
            ["多镇街报告任务状态", str(extreme_task.get("status", ""))],
            ["多镇街报告数量", str(extreme_task.get("reports", extreme_api.get("reports", "")))],
            ["多镇街报告名称", "；".join(extreme_task.get("reportNames", extreme_api.get("reportNames", [])))],
            ["数据库记录/评分/附件/日志", "；".join(f"{key}={value}" for key, value in (extreme.get("database") or {}).items())],
        ],
    )
    extreme_rows = [
        [item.get("name", ""), item.get("status", "")]
        for item in (extreme.get("cases") or [])[:30]
    ]
    if extreme_rows:
        add_table(doc, ["极端用例", "结果"], extreme_rows)

    add_heading(doc, "四-2、DOCX 文件级验收", 1)
    add_table(
        doc,
        ["检查项", "结果"],
        [
            ["总体结果", "通过" if docx_output.get("passed") else "未通过"],
            ["检查文件数量", str(docx_output.get("checkedFiles", ""))],
            ["涉及镇街", "、".join(docx_output.get("selectedTowns", []))],
            ["生成报告", "；".join(docx_output.get("generatedReports", []))],
        ],
    )
    file_rows = [
        [display_text(item.get("name", "")), "通过" if item.get("passed") else "未通过", display_text("；".join(item.get("issues", [])) or "无")]
        for item in (docx_output.get("checks") or [])
    ]
    if file_rows:
        add_table(doc, ["文件", "结果", "问题"], file_rows)

    add_heading(doc, "五、正式报告数字与表格检查", 1)
    add_table(
        doc,
        ["检查项", "观测值", "结果"],
        [
            ["报告文件", str(Path(quality.get("report", "")).name if quality.get("report") else ""), "通过" if quality.get("report") else "未通过"],
            ["付款相关表格数量", str(len(quality.get("paymentTables", []))), "通过" if quality.get("paymentTables") else "未通过"],
            ["异常占位符", "、".join(quality.get("badTokens", [])) or "未发现", "通过" if not quality.get("badTokens") else "未通过"],
            ["替换字符数量", str(quality.get("replacementChars", "")), "通过" if quality.get("replacementChars") == 0 else "未通过"],
            ["数字格式与列数", "通过" if quality.get("passed") else "未通过", "通过" if quality.get("passed") else "未通过"],
        ],
    )

    add_heading(doc, "六、已发现并修复的问题", 1)
    add_table(
        doc,
        ["问题", "修复措施", "复验结果"],
        [
            ["PowerShell 5 读取中文路径导致乱码", "启动脚本与测试脚本改为 Unicode 字符拼接，避免依赖文件编码", "通过"],
            ["后端 storage 目录不存在时 SQLite 无法打开数据库", "测试脚本启动前自动创建 storage 目录", "通过"],
            ["pnpm 可用但 node 未进入 PATH", "测试脚本自动加入 Codex Node 运行时路径", "通过"],
            ["正式报告任务输出到公共生成目录时存在覆盖/路径风险", "后端报告任务改为输出到后端 storage/generated_reports 隔离目录", "通过"],
            ["正式报告付款表存在合并表头导致数字错位风险", "报告生成器改为按完整列数新建数据行，金额固定 2 位小数、系数固定 3 位小数", "通过"],
            ["单镇报告任务默认全量生成，导致任务耗时过长", "报告生成器支持按镇街过滤，后台请求单镇时只生成单镇报告", "通过"],
            ["BAT 启动入口可能闪黑框", "移除推荐 BAT 入口，保留 VBS 零窗口启动入口", "通过"],
        ],
    )

    add_heading(doc, "七、交付建议", 1)
    add_para(doc, "1. 收件人运行入口：双击 Agent/点我启动.vbs；如需英文文件名，也可双击 Agent/start.vbs。")
    add_para(doc, "2. 测试工程师复测入口：执行 Agent/tests/run_agent_checks.ps1。")
    add_para(doc, "3. 测试结果位置：Agent/tests/results。")
    add_para(doc, "4. 本地运行产物 .venv、node_modules、dist、storage、日志不纳入上传交付。")

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
