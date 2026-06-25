from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


PAYMENT_TABLE_MARKERS = [
    ["农村污水处理设施点数", "设计处理规模"],
    ["可用性付费基数Pk3", "第九批"],
    ["建设期考核系数E1", "运维服务绩效考核系数Ec1"],
    ["可用性付费基数（元/月）Pk3", "可用性付费(元/月)"],
    ["第一批运维服务费(元/月)", "运维服务费合计"],
    ["第八批运维服务费(元/月)", "第九批运维服务费(元/月)"],
    ["每月合计扣减费用", "第一~七批运维服务费"],
]

BAD_TOKENS = ["None", "nan", "NaN", "Decimal(", "E+", "e+"]
MONEY_RE = re.compile(r"^-?\d+(?:\.\d{2})?$")
COEFFICIENT_RE = re.compile(r"^-?\d+\.\d{3,4}$")


def table_text(table) -> str:
    return "|".join(cell.text.strip().replace("\n", "") for row in table.rows for cell in row.cells)


def row_text(row) -> list[str]:
    return [cell.text.strip() for cell in row.cells]


def is_numeric_like(text: str) -> bool:
    return bool(re.fullmatch(r"-?\d+(?:\.\d+)?", text))


def main() -> None:
    backend = Path(__file__).resolve().parents[1] / "backend"
    report_dir = backend / "storage" / "generated_reports"
    candidates = sorted(report_dir.glob("北陡镇2023年下半年度村级设施考核报告（正文）.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        candidates = sorted(report_dir.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        raise FileNotFoundError(f"未找到生成报告：{report_dir}")

    path = candidates[0]
    doc = Document(path)
    full_text = "\n".join([p.text for p in doc.paragraphs] + [table_text(table) for table in doc.tables])
    bad_tokens = [token for token in BAD_TOKENS if token in full_text]
    replacement_chars = full_text.count("\ufffd")

    payment_checks = []
    for markers in PAYMENT_TABLE_MARKERS:
        matched = None
        for table in doc.tables:
            text = table_text(table)
            if all(marker in text for marker in markers):
                matched = table
                break
        if matched is None:
            payment_checks.append({"markers": markers, "found": False})
            continue
        rows = [row_text(row) for row in matched.rows]
        data_rows = rows[1:]
        expected_columns = len(matched.columns)
        data_row_lengths_ok = all(len(row) == expected_columns for row in data_rows)
        numeric_cells = [cell for row in data_rows for cell in row if is_numeric_like(cell)]
        inconsistent_numeric = [
            cell
            for cell in numeric_cells
            if "." in cell and not (MONEY_RE.match(cell) or COEFFICIENT_RE.match(cell))
        ]
        payment_checks.append(
            {
                "markers": markers,
                "found": True,
                "columns": expected_columns,
                "dataRows": len(data_rows),
                "dataRowLengthsOk": data_row_lengths_ok,
                "numericCells": len(numeric_cells),
                "inconsistentNumeric": inconsistent_numeric[:20],
            }
        )

    passed = not bad_tokens and replacement_chars == 0 and all(
        item.get("found") and item.get("dataRows", 0) > 0 and item.get("dataRowLengthsOk") and not item.get("inconsistentNumeric")
        for item in payment_checks
    )
    result = {
        "report": str(path),
        "passed": passed,
        "badTokens": bad_tokens,
        "replacementChars": replacement_chars,
        "paymentTables": payment_checks,
    }
    output = Path(__file__).resolve().parent / "results" / "report-quality-summary.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
