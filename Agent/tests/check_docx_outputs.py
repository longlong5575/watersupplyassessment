from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document


FORBIDDEN_LABELS = ["套话", "框架", "合并版", "技能", "高仿", "专项", "保留版", "生成版"]
BAD_TOKENS = ["????", "???", "\ufffd", "None", "nan", "NaN", "Decimal(", "E+", "e+"]
MOJIBAKE_RE = re.compile(r"[ÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖ×ØÙÚÛÜÝÞßàáâãäåæçèéêëìíîïðñòóôõö÷øùúûüýþÿ]{2,}")
REPORT_NAME_RE = re.compile(r"^(.+?)2023年下半年度村级设施考核报告（正文）\.docx$")
SECTION_RE = re.compile(r"([^，。、；\s]+(?:镇|街道|乡|村|区|县|市))农村污水处理设施考核情况")


def read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8-sig"))


def cell_text(cell) -> str:
    return "".join(paragraph.text for paragraph in cell.paragraphs).strip()


def table_text(table) -> str:
    return "\n".join("|".join(cell_text(cell) for cell in row.cells) for row in table.rows)


def doc_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    parts.extend(table_text(table) for table in doc.tables)
    return "\n".join(parts)


def report_target(path: Path) -> str:
    match = REPORT_NAME_RE.match(path.name)
    return match.group(1) if match else path.stem.split("2023年")[0]


def unique(values: list[str]) -> list[str]:
    return sorted(dict.fromkeys(value for value in values if value))


def bad_text_issues(text: str) -> list[str]:
    issues: list[str] = []
    found_tokens = [token for token in BAD_TOKENS if token in text]
    if found_tokens:
        issues.append("发现异常占位/乱码：" + "、".join(found_tokens))
    mojibake = unique(MOJIBAKE_RE.findall(text))
    if mojibake:
        issues.append("发现疑似编码乱码：" + "、".join(mojibake[:10]))
    return issues


def sequence_issues(doc: Document) -> list[str]:
    issues: list[str] = []
    for table_index, table in enumerate(doc.tables, start=1):
        header_row_index = None
        sequence_col = None
        for row_index, row in enumerate(table.rows[:4]):
            values = [cell_text(cell).replace(" ", "") for cell in row.cells]
            for col_index, value in enumerate(values):
                if value == "序号":
                    header_row_index = row_index
                    sequence_col = col_index
                    break
            if header_row_index is not None:
                break
        if header_row_index is None or sequence_col is None:
            continue
        expected = 1
        seen_numbers = 0
        for row in table.rows[header_row_index + 1 :]:
            value = cell_text(row.cells[sequence_col]).strip()
            if value in {"", "合计", "共计", "总计"}:
                continue
            if not re.fullmatch(r"\d+", value):
                continue
            seen_numbers += 1
            if int(value) != expected:
                issues.append(f"表{table_index}序号为{value}，应为{expected}")
                break
            expected += 1
        if seen_numbers == 1 and expected != 2:
            issues.append(f"表{table_index}单行序号异常")
    return issues


def section_towns(text: str) -> list[str]:
    return unique([match.group(1) for match in SECTION_RE.finditer(text)])


def validate_docx(path: Path, selected_towns: list[str]) -> dict[str, Any]:
    doc = Document(path)
    text = doc_text(doc)
    issues = bad_text_issues(text)
    target = report_target(path)
    if any(label in path.name for label in FORBIDDEN_LABELS):
        issues.append("文件名含制作痕迹")
    if path.name != "Agent前后端联调测试验证报告.docx" and not REPORT_NAME_RE.match(path.name):
        issues.append("文件名不符合正式报告命名")
    if target in selected_towns:
        crossed = [town for town in selected_towns if town != target and f"{town}农村污水处理设施考核情况" in text]
        if crossed:
            issues.append("单镇报告串镇：" + "、".join(crossed))
        if f"{target}农村污水处理设施考核情况" not in text:
            issues.append("单镇报告缺少本镇考核章节")
    elif target == "项目":
        missing = [town for town in selected_towns if f"{town}农村污水处理设施考核情况" not in text]
        if missing:
            issues.append("项目汇总报告缺少镇街章节：" + "、".join(missing))
        extra = [town for town in section_towns(text) if town not in selected_towns]
        if extra:
            issues.append("项目汇总报告含非本次镇街章节：" + "、".join(extra))
    issues.extend(sequence_issues(doc))
    return {
        "file": str(path),
        "name": path.name,
        "passed": not issues,
        "issues": issues,
        "tables": len(doc.tables),
        "paragraphs": len(doc.paragraphs),
    }


def main() -> None:
    agent = Path(__file__).resolve().parents[1]
    result_dir = Path(__file__).resolve().parent / "results"
    generated_dir = agent / "backend" / "storage" / "generated_reports"
    extreme = read_json(result_dir / "extreme-check-summary.json")
    report_task = extreme.get("reportTask") or {}
    report_names = report_task.get("reportNames") or []
    selected_towns = unique([name.split("2023年")[0] for name in report_names if name.endswith(".docx") and not name.startswith("项目")])

    generated_files = sorted(generated_dir.glob("*.docx"))
    expected_missing = [name for name in report_names if not (generated_dir / name).exists()]
    validation_report = agent / "tests" / "Agent前后端联调测试验证报告.docx"
    files = generated_files + ([validation_report] if validation_report.exists() else [])
    checks = [validate_docx(path, selected_towns) for path in files]
    if expected_missing:
        checks.append({"file": str(generated_dir), "name": "输出目录", "passed": False, "issues": ["缺少报告：" + "、".join(expected_missing)]})

    result = {
        "passed": all(item["passed"] for item in checks),
        "checkedFiles": len(checks),
        "selectedTowns": selected_towns,
        "generatedReports": [path.name for path in generated_files],
        "checks": checks,
    }
    output = result_dir / "docx-output-summary.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False))
    if not result["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
