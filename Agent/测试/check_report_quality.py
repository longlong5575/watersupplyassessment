from __future__ import annotations

import argparse
import json
import os
import re
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


BAD_TOKENS = ["None", "nan", "NaN", "Decimal(", "reviewed", "submitted", "locked", "returned", "E+", "e+", "原文例文", "例文口径", "原文镇名", "原文报告", "模板版", "系统生成版", "极端重复扣分", "第二原因", "极端问卷", "后台极端", "封顶验证", "修改后再次同步", "按项目目录列示全部镇街"]
BAD_AMOUNT_TEXT = ["使用通用金额基础表", "按通用金额基础表", "合同单价 × 核定处理水量", "奖励金额", "季度奖励金"]
PROJECT_EXPECTATIONS = {
    "郁南": ["镇级及农村设施考核报告", "项目人员组成", "公众调查", "农村污水处理设施", "DB44/2208-2019", "TP", "附件1 考核标准", "附件2 考核评分表"],
    "茂南": ["城镇设施绩效考核报告", "项目人员组成", "水质净化厂", "城镇污水处理设施", "TP", "附件1 考核标准", "附件8 月平均值统计"],
}


def table_text(table) -> str:
    return "\n".join(cell.text.strip() for row in table.rows for cell in row.cells)


def all_text(document: Document) -> str:
    return "\n".join([p.text for p in document.paragraphs] + [table_text(table) for table in document.tables])


def serial_errors(document: Document) -> list[dict[str, object]]:
    errors: list[dict[str, object]] = []
    for index, table in enumerate(document.tables, 1):
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if "序号" not in headers:
            continue
        serial_index = headers.index("序号")
        serials = [row.cells[serial_index].text.strip() for row in table.rows[1:]]
        start = int(serials[0]) if serials and all(value.isdigit() for value in serials) else 1
        expected = [str(number) for number in range(start, start + len(serials))]
        if serials != expected:
            errors.append({"table": index, "serials": serials, "expected": expected})
    return errors


def chapter_order_errors(document: Document) -> list[dict[str, object]]:
    chapters: list[tuple[int, str]] = []
    subsections: dict[int, list[tuple[int, str]]] = {}
    for paragraph in document.paragraphs:
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if not re.fullmatch(r"Heading [1-3]|标题 [1-3]", style_name):
            continue
        text = paragraph.text.strip()
        match = re.match(r"^2\.(\d+)\s+", text)
        if match:
            chapters.append((int(match.group(1)), text))
        subsection = re.match(r"^2\.(\d+)\.(\d+)\s+", text)
        if subsection:
            subsections.setdefault(int(subsection.group(1)), []).append((int(subsection.group(2)), text))
    numbers = [number for number, _ in chapters]
    errors: list[dict[str, object]] = []
    expected = list(range(1, len(numbers) + 1))
    if numbers != expected:
        errors.append({"chapters": [text for _, text in chapters], "numbers": numbers, "expectedOrder": expected})
    for parent, items in subsections.items():
        child_numbers = [number for number, _ in items]
        child_expected = list(range(1, len(child_numbers) + 1))
        if child_numbers != child_expected:
            errors.append({"parent": f"2.{parent}", "chapters": [text for _, text in items], "numbers": child_numbers, "expectedOrder": child_expected})
    return errors


def font_errors(document: Document, path: Path) -> list[dict[str, str]]:
    errors: list[dict[str, str]] = []
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    attributes = [f"{{{namespace['w']}}}{name}" for name in ("ascii", "hAnsi", "eastAsia", "cs")]
    with ZipFile(path) as package:
        styles = ET.fromstring(package.read("word/styles.xml"))
        default_fonts = styles.find(".//w:docDefaults/w:rPrDefault/w:rPr/w:rFonts", namespace)
        if default_fonts is None or any(default_fonts.get(attribute) != "宋体" for attribute in attributes):
            errors.append({"location": "styles:docDefaults", "font": "缺少完整宋体设置"})
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml") or name == "word/styles.xml":
                continue
            try:
                root = ET.fromstring(package.read(name))
            except ET.ParseError:
                continue
            for index, run in enumerate(root.findall(".//w:r", namespace), 1):
                if run.find("w:t", namespace) is None:
                    continue
                fonts = run.find("w:rPr/w:rFonts", namespace)
                if fonts is None or any(fonts.get(attribute) != "宋体" for attribute in attributes):
                    errors.append({"location": f"{name}:run:{index}", "font": "未完整设置宋体"})
                    if len(errors) >= 20:
                        return errors
    return errors


def table_header_errors(document: Document) -> list[dict[str, object]]:
    errors: list[dict[str, object]] = []
    for index, table in enumerate(document.tables, 1):
        if not table.rows:
            continue
        row = table.rows[0]
        if row._tr.get_or_add_trPr().find(qn("w:tblHeader")) is None:
            errors.append({"table": index, "reason": "表头未设置跨页重复"})
        if any(paragraph.paragraph_format.keep_with_next is not True for cell in row.cells for paragraph in cell.paragraphs):
            errors.append({"table": index, "reason": "表头未与下一行同页"})
    return errors


def adjacent_table_errors(document: Document) -> list[dict[str, object]]:
    children = list(document._element.body.iterchildren())
    return [
        {"position": index, "reason": "不同表头的表格之间缺少分隔"}
        for index in range(1, len(children))
        if children[index - 1].tag == qn("w:tbl") and children[index].tag == qn("w:tbl")
    ]


def score_table_indicator_errors(document: Document) -> list[dict[str, object]]:
    errors = []
    for index, table in enumerate(document.tables, 1):
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if "评分条目" in headers and "实得分" in headers and "指标编号" in headers:
            errors.append({"table": index, "headers": headers})
    return errors


def water_summary_errors(document: Document) -> list[dict[str, object]]:
    for index, table in enumerate(document.tables, 1):
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers[:4] == ["序号", "项目点", "取样时间", "检测指标"]:
            required = {"自动判定", "最终判定", "备注"}
            if not required.issubset(headers):
                return [{"table": index, "headers": headers, "reason": "水质汇总缺少判定字段"}]
            results = {row.cells[headers.index("最终判定")].text.strip() for row in table.rows[1:]}
            if not results.issubset({"达标", "不达标", "待判定"}):
                return [{"table": index, "results": sorted(results), "reason": "水质结论不规范"}]
            return []
    return []


def latest_reports(report_root: Path) -> list[Path]:
    candidates = sorted(report_root.rglob("*.docx"), key=lambda path: path.stat().st_mtime, reverse=True)
    selected: list[Path] = []
    for keyword in ("郁南", "茂南"):
        for is_summary in (False, True):
            match = next(
                (
                    path
                    for path in candidates
                    if keyword in path.name and (("汇总" in path.name) == is_summary)
                ),
                None,
            )
            if match:
                selected.append(match)
    return selected


def inspect_report(path: Path) -> dict[str, object]:
    document = Document(path)
    text = all_text(document)
    project_key = "郁南" if "郁南" in path.name or "郁南" in text else "茂南" if "茂南" in path.name or "茂南" in text else "未知"
    expected = PROJECT_EXPECTATIONS.get(project_key, [])
    missing = [item for item in expected if item not in text]
    bad_tokens = [token for token in BAD_TOKENS if token in text]
    bad_amount_text = [token for token in BAD_AMOUNT_TEXT if token in text]
    replacement_chars = text.count("\ufffd")
    sequence_errors = serial_errors(document)
    chapter_errors = chapter_order_errors(document)
    report_font_errors = font_errors(document, path)
    header_errors = table_header_errors(document)
    table_separation_errors = adjacent_table_errors(document)
    indicator_errors = score_table_indicator_errors(document)
    water_errors = water_summary_errors(document)
    unit_errors = re.findall(r"(?i)m\s*(?:\^\s*)?3", text)
    is_summary = "汇总" in path.name
    if project_key == "茂南":
        required_sections = ["摘  要", "目录", "第一章 考核工作概述", "1.6.1 现场检查", "1.6.2 查阅资料", "1.6.3 水质检测", "第二章 城镇水质净化设施考核结果", "第三章 绩效付费计算", "3.3 金额基础表", "主要问题和整改工作建议", "附件1 考核标准", "附件2 周期评分表", "附件3 现场检查照片", "附件5 水质抽检汇总", "附件8 月平均值统计"]
        amount_boundary_missing = "不引用其他项目" not in text and "本项目既有例文和历史付费表" not in text
    else:
        required_sections = ["摘要", "目录", "第一章 考核工作概述", "1.6.1 现场检查", "1.6.2 查阅资料", "1.6.3 问卷调查", "1.6.4 水质检测", "第二章 镇级设施运维考核情况", "第三章 绩效付费计算", "3.3 金额基础表", "第四章 主要问题及整改建议", "附件1 考核标准", "附件2 考核评分表", "附件3 现场照片", "附件5 水质抽检情况汇总表"]
        amount_boundary_missing = "不引用茂南或其他项目金额资料" not in text
    missing_sections = [item for item in required_sections if item not in text]
    with ZipFile(path) as package:
        document_xml = package.read("word/document.xml")
        settings_xml = package.read("word/settings.xml")
        footer_xml = b"".join(
            package.read(name) for name in package.namelist() if name.startswith("word/footer") and name.endswith(".xml")
        )
    toc_field_missing = b"PAGEREF " not in document_xml or b"bookmarkStart" not in document_xml
    toc_text_missing = text.count("第一章 考核工作概述") < 2
    update_fields_missing = b"updateFields" not in settings_xml
    page_field_missing = b" PAGE " not in footer_xml
    weird_numbers = re.findall(r"\d+\.\d{5,}", text)
    summary_scope_missing = is_summary and ("已提交、已复核或已锁定" not in text or "草稿、退回和未提交" not in text)
    min_tables = (18 if project_key == "茂南" else 20) if is_summary else 12
    min_paragraphs = (95 if project_key == "茂南" else 105) if is_summary else 70
    too_short = len(document.paragraphs) < min_paragraphs or len(document.tables) < min_tables
    passed = not missing and not bad_tokens and not bad_amount_text and not amount_boundary_missing and not toc_field_missing and not toc_text_missing and not update_fields_missing and not page_field_missing and replacement_chars == 0 and not sequence_errors and not chapter_errors and not report_font_errors and not header_errors and not table_separation_errors and not indicator_errors and not water_errors and not unit_errors and not summary_scope_missing and not missing_sections and not weird_numbers and not too_short
    return {
        "report": str(path),
        "project": project_key,
        "passed": passed,
        "missingExpectedText": missing,
        "missingSections": missing_sections,
        "badTokens": bad_tokens,
        "badAmountText": bad_amount_text,
        "amountBoundaryMissing": amount_boundary_missing,
        "tocFieldMissing": toc_field_missing,
        "tocTextMissing": toc_text_missing,
        "updateFieldsMissing": update_fields_missing,
        "pageFieldMissing": page_field_missing,
        "replacementChars": replacement_chars,
        "serialErrors": sequence_errors,
        "chapterOrderErrors": chapter_errors,
        "fontErrors": report_font_errors,
        "tableHeaderErrors": header_errors,
        "tableSeparationErrors": table_separation_errors,
        "indicatorCodeErrors": indicator_errors,
        "waterSummaryErrors": water_errors,
        "unitErrors": unit_errors,
        "summaryScopeMissing": summary_scope_missing,
        "overlongDecimals": weird_numbers[:20],
        "tooShort": too_short,
        "tableCount": len(document.tables),
        "paragraphCount": len(document.paragraphs),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="检查郁南和茂南正式报告内容、结构与格式。")
    parser.add_argument("--report-root", type=Path, help="指定需要检查的报告目录。")
    parser.add_argument("--all-reports", action="store_true", help="检查目录内全部 DOCX，而非每项目最新正文和汇总。")
    args = parser.parse_args()
    root = Path(__file__).resolve().parent
    agent_root = root.parent
    base = agent_root.parent.parent if agent_root.parent.name.lower() == "watersupplyassessment" else agent_root.parent
    runtime_root = Path(os.environ.get("WATERSUPPLY_RUNTIME_DIR") or base / "运行脚本" / "watersupply-agent-runtime")
    result_root = runtime_root / "test-results"
    report_root = args.report_root or (result_root / "project-pipeline" / "storage" / "generated_reports")
    reports = sorted(report_root.rglob("*.docx")) if args.all_reports else latest_reports(report_root)
    minimum = 1 if args.report_root else 4
    if len(reports) < minimum:
        raise FileNotFoundError(f"未找到足够的郁南和茂南报告：{report_root}")
    items = [inspect_report(path) for path in reports]
    result = {"passed": all(item["passed"] for item in items), "items": items}
    output = result_root / ("report-quality-summary-all.json" if args.all_reports else "report-quality-summary.json")
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False))
    if not result["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
