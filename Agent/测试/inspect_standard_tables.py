from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent / "标准提取"
SCORE_RE = re.compile(r"^\s*(\d+(?:\.\d+)?)\s*(?:分)?\s*$")


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\u3000", " ")).strip()


def score_value(text: str) -> float:
    match = SCORE_RE.match(clean(text))
    return float(match.group(1)) if match else 0.0


def table_text(table) -> str:
    return "\n".join(clean(cell.text) for row in table.rows for cell in row.cells)


def table_rows(table) -> list[list[str]]:
    return [[clean(cell.text) for cell in row.cells] for row in table.rows]


def row_signature(row: list[str]) -> tuple[str, str, float]:
    nonempty = [cell for cell in row if cell]
    first = nonempty[0] if nonempty else ""
    second = nonempty[1] if len(nonempty) > 1 else ""
    return first, second, score_value(row[-1] if row else "")


def leaf_total(rows: list[list[str]]) -> tuple[float, int]:
    seen: set[tuple[str, str, float]] = set()
    total = 0.0
    count = 0
    for row in rows[1:]:
        if not row:
            continue
        score = score_value(row[-1])
        if score <= 0:
            continue
        sig = row_signature(row)
        if sig in seen:
            continue
        seen.add(sig)
        total += score
        count += 1
    return total, count


def looks_like_standard(text: str) -> bool:
    keywords = ["评价标准", "评分标准", "考核标准", "评分方法", "数据来源", "分值", "扣分"]
    return sum(1 for word in keywords if word in text) >= 3


def inspect(docx_name: str) -> None:
    doc = Document(ROOT / docx_name)
    print(f"\n## {docx_name} tables={len(doc.tables)}")
    for index, table in enumerate(doc.tables):
        rows = table_rows(table)
        text = table_text(table)
        total, count = leaf_total(rows)
        if not looks_like_standard(text) and total not in {100.0, 98.0, 95.0, 92.0, 83.0}:
            continue
        headers = rows[0] if rows else []
        preview = " | ".join(headers[:8])
        title = ""
        for paragraph in doc.paragraphs:
            # python-docx does not expose a direct table title link here; keep table content preview instead.
            pass
        first_rows = []
        for row in rows[1:4]:
            first_rows.append(" | ".join(row[: min(6, len(row))]))
        print(f"table={index} rows={len(rows)} cols={len(rows[0]) if rows else 0} total={total:g} count={count}")
        print(f"  header: {preview}")
        print(f"  sample: {' || '.join(first_rows)}")


def main() -> None:
    inspect("yunan_source.docx")
    inspect("maonan_source.docx")


if __name__ == "__main__":
    main()
