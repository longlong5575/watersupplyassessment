from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


BACKEND = Path(__file__).resolve().parents[1] / "backend"
sys.path.insert(0, str(BACKEND))

from app.services.payment import (  # noqa: E402
    maonan_network_monthly_fee,
    maonan_network_monthly_operation_fee,
    maonan_operation_coefficient,
    maonan_treatment_monthly_fee,
    maonan_water_quality_coefficient,
    town_average_coefficient,
    yunan_dry_season_quality_coefficient,
    yunan_network_monthly_fee,
    yunan_operation_coefficient,
    yunan_town_treatment_monthly_fee,
    yunan_water_quality_coefficient,
)
from app.services.payment_basis import (  # noqa: E402
    load_payment_basis,
    maonan_payment_basis_for_point,
    maonan_payment_basis_rows,
    payment_source_summary,
    yunan_county_network_basis,
    yunan_rural_payment_basis_rows,
    yunan_town_payment_basis_for_point,
    yunan_town_payment_basis_rows,
)
from app.services.payment_context import adjacent_period_name, months_for_period  # noqa: E402
from app.services.project_catalog import MAONAN_TOWNS, town_scoring_policy  # noqa: E402
from app.services.report_tasks import _add_maonan_payment_chapter, _format_report_time  # noqa: E402
from app.services.scoring_policy import calculate_policy_score, scoring_policy  # noqa: E402


def close(actual: float, expected: float, tolerance: float = 1e-8) -> None:
    assert abs(actual - expected) <= tolerance, (actual, expected)


def main() -> None:
    assert _format_report_time("2030-12-15T09:30:00") == "2030-12-15 09:30"
    assert _format_report_time("2026-07-14T07:20:00", assume_utc=True) == "2026-07-14 15:20"
    assert _format_report_time("2026-07-14T07:20:00+00:00", assume_utc=True) == "2026-07-14 15:20"
    close(maonan_operation_coefficient(70), 1)
    close(maonan_operation_coefficient(69), 69 / 70)
    close(yunan_operation_coefficient(90), 1)
    close(yunan_operation_coefficient(81), 0.9)
    close(town_average_coefficient([90, 81], project="yunan") or 0, 0.95)
    yunan_network = scoring_policy("郁南项目", "桂圩镇", "town_network")
    assert yunan_network["applicableMaxScore"] == 88
    assert yunan_network["excludedScore"] == 12
    assert calculate_policy_score(yunan_network, [{"score": 68.5, "deduction": 19.5}])["percentScore"] == 77.8
    maonan_network = scoring_policy("茂南项目", "鳌头镇", "town_network")
    assert maonan_network["applicableMaxScore"] == 82
    assert maonan_network["excludedScore"] == 18
    assert calculate_policy_score(maonan_network, [{"score": 77.4, "deduction": 4.6}])["percentScore"] == 94.4
    assert scoring_policy("郁南项目", "建城镇", "town_network")["mode"] == "direct_100"
    for town_name in ["桂圩镇", "罗顺片区", "通门镇", "大方镇", "河口镇", "东坝镇"]:
        configured = town_scoring_policy("郁南项目", town_name, "town_network")
        assert configured and configured["excludedGroupNames"] == ["泵站运行维护质量"]
    for town_name in ["金塘镇", "鳌头镇"]:
        configured = town_scoring_policy("茂南项目", town_name, "town_network")
        assert configured and configured["mode"] == "scaled_applicable"
    assert town_scoring_policy("郁南项目", "建城镇", "town_network") is None

    # 茂南例文表4-3：108.29/12.21 按出水最低30计算，Kq=0.87（显示值）。
    kq = maonan_water_quality_coefficient(108.29, 12.21)
    close(round(kq, 2), 0.87)
    # 表4-7：Py=1.69、Kq=0.87、E1=1，运营单价折算显示为1.67。
    operation_part = 1.69 * (3 / 5 + kq / 10 + 3 / 10)
    close(round(operation_part, 2), 1.67)
    # 表4-7山阁镇2025年6月：Kq显示0.92，但金额使用(113.12-30)/90的未舍入值。
    shange_kq = maonan_water_quality_coefficient(113.12, 10.91)
    close(maonan_network_monthly_operation_fee(
        annual_operation_fee_yuan=302277.33,
        water_quality_coefficient=shange_kq,
        operation_coefficient=1,
    ), 24997.33, tolerance=0.2)
    close(maonan_network_monthly_operation_fee(
        annual_operation_fee_yuan=120000,
        water_quality_coefficient=0.8,
        operation_coefficient=0.9,
    ), 120000 * (3 / 5 + 0.8 / 10 + 3 * 0.9 / 10) / 12)
    close(maonan_network_monthly_fee(
        annual_availability_fee_ten_thousand_yuan=168.48,
        annual_operation_fee_ten_thousand_yuan=30.227733,
        water_quality_coefficient=shange_kq,
        operation_coefficient=1,
    ), 168.48 / 12 + 30.227733 * (3 / 5 + shange_kq / 10 + 3 / 10) / 12)

    # 污水处理费同时包含运营费和可用性付费，单位为万元/月。
    close(maonan_treatment_monthly_fee(
        operation_unit_price=1.69,
        monthly_volume_ten_thousand_tons=3,
        water_quality_coefficient=0.87,
        operation_coefficient=1,
        annual_availability_fee_ten_thousand_yuan=33.59,
    ), 1.69 * 3 * (3 / 5 + 0.87 / 10 + 3 / 10) + 33.59 / 12)

    close(yunan_water_quality_coefficient(210, 35), 1.05)
    close(yunan_water_quality_coefficient(230, 35), 1.1)
    close(yunan_water_quality_coefficient(160, 35, effluent_qualified=False), 0.9)
    basis = load_payment_basis()
    assert set(basis) == {"yunan", "maonan"}
    assert len(maonan_payment_basis_rows()) == 7
    assert maonan_payment_basis_for_point("茂南区水质净化厂")["treatmentOperationUnitPriceYuanPerCubicMeter"] == 1.22
    assert maonan_payment_basis_for_point("中科云粤西产业园污水收集管网")["adjustedNetworkOperationFeeYuanPerYear"] == 861768.32
    assert basis["yunan"]["basisStatus"] == "合同附件18金额基数已结构化"
    assert basis["maonan"]["basisStatus"] == "历史金额基数已结构化"
    assert len(yunan_town_payment_basis_rows()) == 15
    assert len(yunan_rural_payment_basis_rows()) == 15
    assert yunan_county_network_basis()["networkAvailabilityFeeTenThousandYuanPerYear"] == 649.84
    assert yunan_town_payment_basis_for_point("建城镇污水处理厂")["treatmentOperationUnitPriceYuanPerCubicMeter"] == 1.46
    close(sum(row["networkAvailabilityFeeTenThousandYuanPerYear"] for row in yunan_town_payment_basis_rows()), 938.64, tolerance=0.01)
    close(sum(row["operationFeeTenThousandYuanPerYear"] for row in yunan_rural_payment_basis_rows()), 372.21, tolerance=0.01)
    close(yunan_town_treatment_monthly_fee(
        operation_unit_price=1.46,
        monthly_volume_ten_thousand_cubic_meters=3,
        water_quality_coefficient=1.0,
        operation_coefficient=0.9,
    ), 1.46 * 3 * (3 / 4 + 1 / 4 * 0.9))
    close(yunan_dry_season_quality_coefficient(20), 89 / 90)
    close(yunan_dry_season_quality_coefficient(21), 1)
    close(yunan_network_monthly_fee(
        annual_network_fee_ten_thousand_yuan=54.51,
        load_coefficient=0.8,
        water_quality_coefficient=1.0,
        dry_season_quality_coefficient=1.0,
        operation_coefficient=0.9,
    ), 54.51 / 12 * (5 / 6 * 0.8 + 1 / 6 * 0.9))
    assert adjacent_period_name("2026年第1季度", -1) == "2025年第4季度"
    assert adjacent_period_name("2026年第4季度", 1) == "2027年第1季度"
    assert months_for_period("2026年下半年度") == ["2026-07", "2026-08", "2026-09", "2026-10", "2026-11", "2026-12"]
    mapped_targets = 0
    for town in MAONAN_TOWNS:
        row = maonan_payment_basis_for_point(town["name"])
        assert row is not None, town["name"]
        for target in town["assessmentTargets"]:
            mapped_targets += 1
            if target == "town_plant":
                assert row["treatmentAvailabilityFeeTenThousandYuanPerYear"] is not None
                assert row["treatmentOperationUnitPriceYuanPerCubicMeter"] is not None
            if target == "town_network":
                assert row["networkAvailabilityFeeTenThousandYuanPerYear"] is not None
                assert row["adjustedNetworkOperationFeeYuanPerYear"] is not None
    assert mapped_targets == 12
    assert "茂南项目既有绩效考核报告" in payment_source_summary("茂南项目")
    assert "表4-1 水质净化设施项目服务费" in payment_source_summary("茂南项目")

    report = Document()
    report_records = [
        {
            "town": "金塘镇",
            "facilityType": "town_plant",
            "totalScore": 60,
            "rawPayload": {
                "paymentData": {
                    "months": [
                        {
                            "month": "2026-04",
                            "influentCod": 108.29,
                            "effluentCod": 12.21,
                            "monthlyVolumeTenThousandCubicMeters": 3,
                        }
                    ]
                },
            },
            "paymentContext": {
                "appliedOperationCoefficient": 69 / 70,
                "coefficientSourcePeriod": "2026年第1季度",
                "currentScoreAppliesTo": "2026年第3季度",
                "months": [
                    {
                        "month": "2026-04",
                        "influentCod": 108.29,
                        "effluentCod": 12.21,
                        "monthlyVolumeTenThousandCubicMeters": 3,
                    }
                ],
            },
        },
        {
            "town": "山阁镇",
            "facilityType": "town_network",
            "totalScore": 60,
            "rawPayload": {
                "paymentData": {
                    "months": [
                        {"month": "2026-04", "influentCod": 113.12, "effluentCod": 10.91}
                    ]
                }
            },
            "paymentContext": {
                "appliedOperationCoefficient": 69 / 70,
                "coefficientSourcePeriod": "2026年第1季度",
                "currentScoreAppliesTo": "2026年第3季度",
                "months": [
                    {"month": "2026-04", "influentCod": 113.12, "effluentCod": 10.91}
                ],
            },
        },
    ]
    _add_maonan_payment_chapter(report, report_records)
    report_text = "\n".join(paragraph.text for paragraph in report.paragraphs)
    report_text += "\n" + "\n".join(cell.text for table in report.tables for row in table.rows for cell in row.cells)
    plant_e1 = 69 / 70
    plant_amount = maonan_treatment_monthly_fee(
        operation_unit_price=1.69,
        monthly_volume_ten_thousand_tons=3,
        water_quality_coefficient=kq,
        operation_coefficient=plant_e1,
        annual_availability_fee_ten_thousand_yuan=33.59,
    )
    network_amount = maonan_network_monthly_fee(
        annual_availability_fee_ten_thousand_yuan=168.48,
        annual_operation_fee_ten_thousand_yuan=30.227733,
        water_quality_coefficient=shange_kq,
        operation_coefficient=plant_e1,
    )
    assert "代入：Pz=1.6900×3.0000×" in report_text
    assert f"={plant_amount:.4f}万元" in report_text
    assert "代入：Pw=168.4800/12×" in report_text
    assert f"={network_amount:.4f}万元" in report_text
    assert "2026年第1季度" in report_text
    assert "缺少月均进出水COD" not in report_text
    print("PASS: 郁南/茂南扣分系数与金额公式反算")


if __name__ == "__main__":
    main()
