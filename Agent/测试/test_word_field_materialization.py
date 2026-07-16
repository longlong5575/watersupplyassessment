from __future__ import annotations

import os
import re
from pathlib import Path
from zipfile import ZipFile


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
import sys
sys.path.insert(0, str(BACKEND))

from docx import Document  # noqa: E402

from app.services.report_tasks import (  # noqa: E402
    _add_source_toc,
    _finalize_static_toc,
)


def default_runtime_root() -> Path:
    configured = os.environ.get("WATERSUPPLY_RUNTIME_DIR")
    if configured:
        return Path(configured)
    base = ROOT.parent.parent if ROOT.parent.name.lower() == "watersupplyassessment" else ROOT.parent
    return base / "运行脚本" / "watersupply-agent-runtime"


def main() -> None:
    output_dir = default_runtime_root() / "test-results" / "word-field-materialization"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "目录页码静默定稿验证.docx"

    document = Document()
    document.add_paragraph("封面页")
    _add_source_toc(document, [])
    document.add_heading("第一章 考核工作概述", level=1)
    document.add_paragraph("目录页码定稿验证正文。")
    _finalize_static_toc(document)
    document.save(output_path)

    with ZipFile(output_path) as archive:
        before_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
    assert "PAGEREF" not in before_xml
    assert "updateFields" not in settings_xml

    toc_lines = [paragraph.text.strip() for paragraph in Document(output_path).paragraphs]
    toc_line = next(text for text in toc_lines if "第一章 考核工作概述" in text)
    page_match = re.search(r"(\d+)$", toc_line)
    assert page_match, f"目录行末缺少数字页码：{toc_line!r}"
    assert int(page_match.group(1)) > 1, f"目录页码没有按实际分页写入：{toc_line!r}"

    print(f"PASS: Word目录页码静默定稿（{output_path}）")


if __name__ == "__main__":
    main()
