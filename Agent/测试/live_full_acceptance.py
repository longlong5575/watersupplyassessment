from __future__ import annotations

import os
import shutil
import sys
import json
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent.parent if ROOT.parent.name.lower() == "watersupplyassessment" else ROOT.parent
RUN_SCRIPTS_NAME = "".join(chr(c) for c in [0x8fd0, 0x884c, 0x811a, 0x672c])
RUNTIME = WORKSPACE / RUN_SCRIPTS_NAME / "watersupply-agent-runtime"
RESULTS = RUNTIME / "test-results" / "live-full-acceptance"
DB_PATH = RUNTIME / "storage" / "assessment.db"

os.environ["DATABASE_URL"] = f"sqlite:///{DB_PATH.as_posix()}"
os.environ["STORAGE_DIR"] = str((RUNTIME / "storage").resolve())
sys.path.insert(0, str(ROOT / "backend"))

import httpx
from sqlalchemy import delete, func, select

from app.core.database import SessionLocal
from app.models import AssessmentCycle, AssessmentRecord, City, Report, ReportTask


STATUS_PATH = WORKSPACE / RUN_SCRIPTS_NAME / "watersupply-agent-runtime" / "logs" / "startup-status.json"
BASE_URL = os.environ.get("WATERSUPPLY_BACKEND_URL") or "http://127.0.0.1:8000"
PERIOD = "".join(chr(c) for c in [0x32, 0x30, 0x33, 0x30, 0x5e74, 0x7b2c, 0x34, 0x5b63, 0x5ea6])
if not os.environ.get("WATERSUPPLY_BACKEND_URL") and STATUS_PATH.exists():
    try:
        BASE_URL = json.loads(STATUS_PATH.read_text(encoding="utf-8")).get("backendUrl") or BASE_URL
    except Exception:
        pass


def require(response: httpx.Response, label: str):
    if response.status_code >= 300:
        raise RuntimeError(f"{label}失败：{response.status_code} {response.text}")
    return response.json()


def prepare_cycles() -> dict[str, str]:
    cycle_ids: dict[str, str] = {}
    with SessionLocal() as session:
        project_names = ["".join(chr(c) for c in codes) for codes in ([0x90c1, 0x5357, 0x9879, 0x76ee], [0x8302, 0x5357, 0x9879, 0x76ee])]
        projects = session.scalars(select(City).where(City.name.in_(project_names))).all()
        for city in projects:
            existing = session.scalar(select(AssessmentCycle).where(AssessmentCycle.city_id == city.id, AssessmentCycle.name == PERIOD))
            if existing is not None:
                raise RuntimeError(f"{city.name}已存在{PERIOD}测试周期，请先清理旧测试数据。")
            cycle = AssessmentCycle(city_id=city.id, name=PERIOD, status="active")
            session.add(cycle)
            session.flush()
            cycle_ids[city.name] = cycle.id
        session.commit()
    return cycle_ids


def cleanup_live_data(cycle_ids: dict[str, str]) -> None:
    with SessionLocal() as session:
        project_ids = {
            city.name: city.id
            for city in session.scalars(select(City).where(City.name.in_(cycle_ids))).all()
        }
    try:
        with httpx.Client(base_url=BASE_URL, timeout=120) as client:
            login = client.post("/api/auth/login", json={"username": "admin", "password": "Admin@123456"})
            if login.status_code >= 300:
                return
            headers = {"Authorization": f"Bearer {login.json()['token']}"}
            for project_name, cycle_id in cycle_ids.items():
                project_id = project_ids.get(project_name)
                if project_id:
                    client.delete(
                        "/api/mobile/assessment-records",
                        headers=headers,
                        params={"city_id": project_id, "cycle_id": cycle_id, "period": PERIOD},
                    )
    except httpx.HTTPError:
        pass


def cleanup_schema(cycle_ids: dict[str, str]) -> None:
    with SessionLocal() as session:
        session.execute(delete(AssessmentCycle).where(AssessmentCycle.id.in_(cycle_ids.values())))
        session.commit()


def verify_cleanup(cycle_ids: dict[str, str]) -> None:
    ids = list(cycle_ids.values())
    if not ids:
        return
    with SessionLocal() as session:
        remaining = {
            "测试周期": session.scalar(select(func.count()).select_from(AssessmentCycle).where(AssessmentCycle.id.in_(ids))) or 0,
            "考核记录": session.scalar(select(func.count()).select_from(AssessmentRecord).where(AssessmentRecord.cycle_id.in_(ids))) or 0,
            "报告任务": session.scalar(select(func.count()).select_from(ReportTask).where(ReportTask.cycle_id.in_(ids))) or 0,
            "生成报告": session.scalar(select(func.count()).select_from(Report).where(Report.cycle_id.in_(ids))) or 0,
        }
    leftovers = {name: count for name, count in remaining.items() if count}
    if leftovers:
        raise RuntimeError(f"全真验收清理不完整：{leftovers}")

def choose_option(item: dict, index: int) -> list[dict]:
    options = item.get("deductionOptions") or []
    if not options or index % 5:
        return []
    option = options[0]
    entry = {
        "optionId": option["id"],
        "selection": "standard",
        "instances": min(2, int(option.get("maxInstances") or 2)),
        "note": f"现场核查发现：{option['name']}",
    }
    if option.get("type") == "range":
        low = float(option.get("min") or 0)
        high = float(option.get("max") or option.get("deduction") or 0)
        entry["rangeValue"] = round((low + high) / 2, 2)
    return [entry]


def make_payload(project: dict, cycle_id: str, town: dict, village: dict | None, facility_type: str, standards: dict) -> dict:
    leaves = [item for item in standards["items"] if item["level"] == 3]
    entries = {
        item["id"]: {
            "itemId": item["id"],
            "done": True,
            "options": choose_option(item, index),
        }
        for index, item in enumerate(leaves)
    }
    record = {
        "village": village["name"] if village else "",
        "primaryFacilityType": facility_type,
        "currentScore": 100,
        "entries": entries,
    }
    if facility_type == "rural_treatment":
        record["surveyEntries"] = {
            "satisfaction_villager1": {"score": 4, "comment": "整体满意，建议加强巡检", "completed": True},
            "sewage_collection_villager1": {"score": 4, "comment": "收集效果较好", "completed": True},
        }
        record["waterQuality"] = {
            "sampleTime": "2030-12-15T09:30:00",
            "dischargeStandard": "DB44/2208-2019",
            "processType": "一体化生化处理设施",
            "designScale": "50",
            "codValue": "45", "codLimit": "40",
            "nh3nValue": "4.2", "nh3nLimit": "5",
            "tpValue": "0.4", "tpLimit": "0.5",
            "conclusion": "unqualified",
            "completed": True,
            "manualOverride": False,
            "remark": "CODCr实测值超过限值",
        }
    return {
        "schemaVersion": "1.0",
        "cityId": project["id"],
        "cycleId": cycle_id,
        "indicatorVersionId": standards["version"]["id"],
        "city": project["name"],
        "period": PERIOD,
        "town": town["name"],
        "villages": [record],
    }


def main() -> None:
    if RESULTS.exists():
        shutil.rmtree(RESULTS)
    RESULTS.mkdir(parents=True)
    cycle_ids = prepare_cycles()
    reports: list[dict] = []
    try:
        with httpx.Client(base_url=BASE_URL, timeout=120) as client:
            inspector_token = require(client.post("/api/auth/login", json={"username": "inspector", "password": "Inspector@123456"}), "检查员登录")["token"]
            admin_token = require(client.post("/api/auth/login", json={"username": "admin", "password": "Admin@123456"}), "管理员登录")["token"]
            inspector = {"Authorization": f"Bearer {inspector_token}"}
            admin = {"Authorization": f"Bearer {admin_token}"}
            projects = require(client.get("/api/mobile/projects", headers=admin), "读取项目")["items"]
            by_name = {item["name"]: item for item in projects}

            scenarios = [
                ("郁南项目", "rural_treatment"),
                ("郁南项目", "town_network"),
                ("郁南项目", "town_plant"),
                ("茂南项目", "town_plant"),
                ("茂南项目", "town_network"),
                ("茂南项目", "town_plant"),
            ]
            used_towns: set[str] = set()
            towns_by_project: dict[str, list[str]] = {"郁南项目": [], "茂南项目": []}
            for scenario_index, (project_name, facility_type) in enumerate(scenarios, 1):
                project = by_name[project_name]
                towns = require(client.get("/api/mobile/towns", headers=admin, params={"city_id": project["id"]}), "读取镇街")["items"]
                town = next(item for item in towns if facility_type in item["assessmentTargets"] and item["name"] not in used_towns)
                used_towns.add(town["name"])
                towns_by_project[project_name].append(town["name"])
                village = None
                if facility_type == "rural_treatment":
                    village = require(client.get(f"/api/mobile/towns/{town['id']}/villages", headers=admin), "读取村点")["items"][0]
                standards = require(client.get("/api/mobile/indicator-standards", headers=admin, params={
                    "city_id": project["id"], "cycle_id": cycle_ids[project_name], "facility_type": facility_type,
                }), "读取评分标准")
                leaves = [item for item in standards["items"] if item["level"] == 3]
                total_full_score = round(sum(float(item["fullScore"]) for item in leaves), 2)
                assert leaves and total_full_score == 100, f"{project_name} {facility_type} 评分项合计为 {total_full_score}"
                assert all(item.get("deductionOptions") for item in leaves)
                assert all(item.get("evaluationStandard") and item.get("dataSource") for item in leaves)

                payload = make_payload(project, cycle_ids[project_name], town, village, facility_type, standards)
                created = require(client.post("/api/mobile/assessment-records", headers=inspector, json=payload), "保存考核")
                record_id = created["recordIds"][0]
                require(client.post(f"/api/mobile/assessment-records/{record_id}/submit", headers=inspector), "提交考核")
                detail = require(client.get(f"/api/records/{record_id}", headers=admin), "读取评分结果")
                deduction = round(sum(float(item["deduction"]) for item in detail["scores"]), 2)
                assert round(float(detail["totalScore"]), 2) == round(100 - deduction, 2)
                score = next(item for item in detail["scores"] if float(item["deduction"]) > 0)
                import fitz

                photo_document = fitz.open()
                photo_page = photo_document.new_page(width=1200, height=800)
                photo_page.draw_rect((40, 40, 1160, 760), color=(0.1, 0.35, 0.55), width=10)
                photo_pixmap = photo_page.get_pixmap(alpha=False)
                photo_bytes = photo_pixmap.tobytes("png")
                photo_document.close()
                require(client.post(
                    f"/api/mobile/assessment-records/{record_id}/attachments",
                    headers=inspector,
                    data={"score_id": score["id"], "deduction_option_id": score.get("deductionOptionId") or ""},
                    files={"file": (f"第{scenario_index}组现场问题照片.png", photo_bytes, "image/png")},
                ), "上传证据")

                pdf = fitz.open()
                page = pdf.new_page()
                page.draw_rect((72, 90, 520, 680), color=(0.1, 0.35, 0.55), width=4)
                pdf_bytes = pdf.tobytes()
                pdf.close()
                require(client.post(
                    f"/api/mobile/assessment-records/{record_id}/attachments",
                    headers=inspector,
                    files={"file": (f"{town['name']}水质检测报告.pdf", pdf_bytes, "application/pdf")},
                ), "上传检测报告")
                require(client.post(f"/api/records/{record_id}/review", headers=admin), "平台复核")
                require(client.post(f"/api/agent/records/{record_id}/analysis", headers=admin), "分析校验")

                task = require(client.post("/api/report-tasks", headers=admin, json={
                    "source": "dashboard", "projectId": project["id"], "period": PERIOD,
                    "townNames": [town["name"]], "outputs": ["separate"],
                }), "生成报告")
                result = require(client.get(f"/api/report-tasks/{task['id']}", headers=admin), "读取报告任务")
                assert result["status"] == "completed", result.get("error")
                report = result["reports"][0]
                preview = require(client.get(f"/api/reports/{report['id']}/preview", headers=admin), "预览报告")
                assert preview["content"]["paragraphCount"] >= 20 and preview["content"]["tableCount"] >= 6
                download = client.get(f"/api/reports/{report['id']}/download", headers=admin)
                if download.status_code != 200 or len(download.content) < 10000:
                    raise RuntimeError("报告下载失败或文件不完整")
                output = RESULTS / report["name"]
                output.write_bytes(download.content)
                with ZipFile(BytesIO(download.content)) as package:
                    document_xml = package.read("word/document.xml")
                    assert document_xml.count(b"<w:drawing") >= 2
                report_document = Document(BytesIO(download.content))
                report_text = "\n".join(
                    [paragraph.text for paragraph in report_document.paragraphs]
                    + [cell.text for table in report_document.tables for row in table.rows for cell in row.cells]
                )
                assert "指标编号" not in report_text
                assert "m3" not in report_text.lower()
                assert "总体评价和设施情况概览" in report_text
                photo_captions = [paragraph.text for paragraph in report_document.paragraphs if paragraph.text.startswith("序号：")]
                assert photo_captions and all("项目点：" in caption and ".png" not in caption.lower() for caption in photo_captions)
                if facility_type == "rural_treatment":
                    water_tables = [
                        table for table in report_document.tables
                        if table.rows and [cell.text.strip() for cell in table.rows[0].cells][:4] == ["序号", "项目点", "取样时间", "检测指标"]
                    ]
                    assert water_tables
                    water_table = water_tables[-1]
                    headers = [cell.text.strip() for cell in water_table.rows[0].cells]
                    assert {"自动判定", "最终判定", "备注"}.issubset(headers)
                    automatic_results = {row.cells[headers.index("自动判定")].text.strip() for row in water_table.rows[1:]}
                    assert {"达标", "不达标"}.issubset(automatic_results)
                reports.append({
                    "project": project_name, "town": town["name"], "facilityType": facility_type,
                    "score": detail["totalScore"], "deduction": deduction, "path": str(output),
                    "paragraphCount": preview["content"]["paragraphCount"],
                    "tableCount": preview["content"]["tableCount"],
                })

            for project_name, town_names in towns_by_project.items():
                project = by_name[project_name]
                task = require(client.post("/api/report-tasks", headers=admin, json={
                    "source": "dashboard", "projectId": project["id"], "period": PERIOD,
                    "townNames": town_names, "outputs": ["summary"],
                }), "生成汇总报告")
                result = require(client.get(f"/api/report-tasks/{task['id']}", headers=admin), "读取汇总报告任务")
                assert result["status"] == "completed", result.get("error")
                report = result["reports"][0]
                preview = require(client.get(f"/api/reports/{report['id']}/preview", headers=admin), "预览汇总报告")
                download = client.get(f"/api/reports/{report['id']}/download", headers=admin)
                if download.status_code != 200 or len(download.content) < 10000:
                    raise RuntimeError("汇总报告下载失败或文件不完整")
                output = RESULTS / report["name"]
                output.write_bytes(download.content)
                summary_document = Document(BytesIO(download.content))
                summary_text = "\n".join(
                    [paragraph.text for paragraph in summary_document.paragraphs]
                    + [cell.text for table in summary_document.tables for row in table.rows for cell in row.cells]
                )
                assert "指标编号" not in summary_text
                assert "m3" not in summary_text.lower()
                assert "2.1 项目考核结果汇总" in summary_text
                reports.append({
                    "project": project_name, "town": "汇总", "facilityType": "summary",
                    "score": None, "deduction": None, "path": str(output),
                    "paragraphCount": preview["content"]["paragraphCount"],
                    "tableCount": preview["content"]["tableCount"],
                })

            for project_name, cycle_id in cycle_ids.items():
                project = by_name[project_name]
                require(client.delete("/api/mobile/assessment-records", headers=admin, params={
                    "city_id": project["id"], "cycle_id": cycle_id, "period": PERIOD,
                }), "清理测试数据")
    finally:
        cleanup_live_data(cycle_ids)
        cleanup_schema(cycle_ids)
        verify_cleanup(cycle_ids)

    import json
    (RESULTS / "acceptance_result.json").write_text(json.dumps({"passed": True, "reports": reports}, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"passed": True, "reports": reports}, ensure_ascii=False))


if __name__ == "__main__":
    main()
