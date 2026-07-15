from __future__ import annotations

import calendar
import json
import mimetypes
import os
import re
import shutil
import sys
from collections import defaultdict
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from sqlalchemy import delete, func, select


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent.parent
RUNTIME = WORKSPACE / "运行脚本" / "watersupply-agent-runtime"
SOURCE_DIR = RUNTIME / "test-results" / "exemplar-replay" / "source-extract"
FIXTURE_PATH = RUNTIME / "test-results" / "exemplar-replay" / "fixture.json"
RESULTS = RUNTIME / "test-results" / "exemplar-replay" / "full-flow"
DELIVERY = WORKSPACE / "生成" / "例文数据全流程回归"
DB_PATH = RUNTIME / "storage" / "assessment.db"
STATUS_PATH = RUNTIME / "logs" / "startup-status.json"

os.environ["DATABASE_URL"] = f"sqlite:///{DB_PATH.as_posix()}"
os.environ["STORAGE_DIR"] = str((RUNTIME / "storage").resolve())
sys.path.insert(0, str(ROOT / "backend"))

from app.core.database import SessionLocal
from app.models import AssessmentCycle, AssessmentRecord, City, Report, ReportTask


PERIOD_SPECS = {
    ("郁南项目", "2025年第2季度"): "第2季度",
    ("茂南项目", "2025年上半年度"): "上半年度",
    ("茂南项目", "2025年下半年度"): "下半年度",
}


def require(response: httpx.Response, label: str) -> dict[str, Any]:
    if response.status_code >= 300:
        raise RuntimeError(f"{label}失败：{response.status_code} {response.text}")
    return response.json()


def base_url() -> str:
    if STATUS_PATH.exists():
        try:
            return json.loads(STATUS_PATH.read_text(encoding="utf-8")).get("backendUrl") or "http://127.0.0.1:8000"
        except (OSError, ValueError):
            pass
    return "http://127.0.0.1:8000"


def prepare_cycles() -> tuple[dict[tuple[str, str], str], dict[tuple[str, str], str]]:
    with SessionLocal() as session:
        projects = {item.name: item for item in session.scalars(select(City).where(City.name.in_(["郁南项目", "茂南项目"]))).all()}
        selected_year = None
        for year in range(2030, 2023, -1):
            requested = {
                (project, source_period): f"{year}年{suffix}"
                for (project, source_period), suffix in PERIOD_SPECS.items()
            }
            collision = False
            for (project_name, _), period in requested.items():
                project = projects[project_name]
                if session.scalar(select(AssessmentCycle.id).where(AssessmentCycle.city_id == project.id, AssessmentCycle.name == period)):
                    collision = True
                    break
            if not collision:
                selected_year = year
                break
        if selected_year is None:
            raise RuntimeError("2024年至2030年的测试周期均已存在，无法安全创建例文回放周期。")

        cycle_ids: dict[tuple[str, str], str] = {}
        periods: dict[tuple[str, str], str] = {}
        for key, suffix in PERIOD_SPECS.items():
            project_name, _ = key
            period = f"{selected_year}年{suffix}"
            cycle = AssessmentCycle(city_id=projects[project_name].id, name=period, status="active")
            session.add(cycle)
            session.flush()
            cycle_ids[key] = cycle.id
            periods[key] = period
        session.commit()
        return cycle_ids, periods


def cleanup(client: httpx.Client, admin: dict[str, str], cycle_ids: dict[tuple[str, str], str], projects: dict[str, dict[str, Any]]) -> None:
    for (project_name, _), cycle_id in cycle_ids.items():
        try:
            client.delete(
                "/api/mobile/assessment-records",
                headers=admin,
                params={"city_id": projects[project_name]["id"], "cycle_id": cycle_id},
                timeout=180,
            )
        except httpx.HTTPError:
            pass
    with SessionLocal() as session:
        session.execute(delete(AssessmentCycle).where(AssessmentCycle.id.in_(cycle_ids.values())))
        session.commit()


def verify_cleanup(cycle_ids: dict[tuple[str, str], str]) -> None:
    ids = list(cycle_ids.values())
    with SessionLocal() as session:
        remaining = {
            "测试周期": session.scalar(select(func.count()).select_from(AssessmentCycle).where(AssessmentCycle.id.in_(ids))) or 0,
            "考核记录": session.scalar(select(func.count()).select_from(AssessmentRecord).where(AssessmentRecord.cycle_id.in_(ids))) or 0,
            "报告任务": session.scalar(select(func.count()).select_from(ReportTask).where(ReportTask.cycle_id.in_(ids))) or 0,
            "生成报告": session.scalar(select(func.count()).select_from(Report).where(Report.cycle_id.in_(ids))) or 0,
        }
    leftovers = {key: value for key, value in remaining.items() if value}
    if leftovers:
        raise RuntimeError(f"例文回放数据清理不完整：{leftovers}")


def _number(value: Any) -> float | None:
    try:
        if value in (None, "", "/"):
            return None
        return float(str(value).replace("%", "").replace("*", "").strip())
    except ValueError:
        return None


def source_metadata() -> dict[str, Any]:
    yunan = json.loads((SOURCE_DIR / "yunan-full.json").read_text(encoding="utf-8"))
    maonan = json.loads((SOURCE_DIR / "maonan-full.json").read_text(encoding="utf-8"))
    result: dict[str, Any] = {
        "water": defaultdict(list),
        "survey": {},
        "payment": defaultdict(list),
        "facility": {},
    }

    yunan_summary = yunan["tables"][7]["rows"]
    for row in yunan_summary[1:]:
        if len(row) < 5 or not row[1]:
            continue
        result["facility"][("郁南项目", "2025年第2季度", row[1], "town_plant")] = {
            "designScale": _number(row[3]), "averageDailyVolume": _number(row[4]), "processType": row[2]
        }
    yunan_water = yunan["tables"][236]["rows"]
    current_town = ""
    for row in yunan_water[2:]:
        if len(row) < 13:
            continue
        current_town = row[1] or current_town
        if not current_town or not row[2]:
            continue
        sample = {
            "sampleTime": row[2], "influentCod": _number(row[3]), "influentNh3n": _number(row[4]),
            "influentTp": _number(row[5]), "codValue": _number(row[6]), "nh3nValue": _number(row[7]),
            "tpValue": _number(row[8]), "kq": _number(row[10]), "qualified": row[12] == "是",
        }
        result["water"][("郁南项目", "2025年第2季度", current_town, None)].append(sample)
    for key, samples in list(result["water"].items()):
        project, source_period, town, point = key
        if point is not None:
            continue
        if project != "郁南项目":
            continue
        facility = result["facility"].get((project, source_period, town, "town_plant"), {})
        daily = facility.get("averageDailyVolume")
        for sample in samples:
            month_match = re.search(r"(\d{4})年(\d{1,2})月", sample["sampleTime"])
            if not month_match:
                continue
            year, month = map(int, month_match.groups())
            result["payment"][(project, source_period, town)].append({
                "month": f"{year}-{month:02d}",
                "influentCod": sample["influentCod"], "effluentCod": sample["codValue"],
                "effluentQualified": sample["qualified"], "legacyKq": sample["kq"],
                "averageDailyVolumeCubicMeters": daily,
                "monthlyVolumeTenThousandCubicMeters": round(daily * calendar.monthrange(year, month)[1] / 10000, 4) if daily else None,
            })

    current_town = ""
    current_village = ""
    for row in yunan["tables"][237]["rows"]:
        if len(row) < 12 or not str(row[0] or "").strip().isdigit():
            continue
        current_town = str(row[1] or current_town).strip()
        current_village = str(row[2] or current_village).strip()
        point = str(row[3] or "").strip()
        if not current_town or not point:
            continue
        result["water"][("郁南项目", "2025年第2季度", current_town, point)].append({
            "sampleTime": row[4],
            "administrativeVillage": current_village,
            "influentCod": _number(row[5]),
            "influentNh3n": _number(row[6]),
            "codValue": _number(row[7]),
            "nh3nValue": _number(row[8]),
            "kq": _number(row[9]),
            "qualified": str(row[10]).strip() == "是",
            "influentQualified": str(row[11]).strip() == "是",
        })

    for table in yunan["tables"]:
        context = " > ".join(table.get("contextBefore") or [])
        matches = re.findall(r"([^ >]+镇)农村污水处理设施公众调查情况汇总表", context)
        if not matches:
            continue
        town = matches[-1]
        current_village = ""
        for row in table.get("rows") or []:
            if len(row) < 11:
                continue
            current_village = str(row[0] or current_village).strip()
            point = str(row[1] or "").strip()
            count = _number(row[2])
            satisfaction_score = _number(row[6])
            collection_score = _number(row[10])
            if not point or point == "自然村" or count is None or count <= 0:
                continue
            result["survey"][("郁南项目", "2025年第2季度", town, point)] = {
                "administrativeVillage": current_village,
                "count": int(count),
                "satisfaction": {
                    "a": int(_number(row[3]) or 0), "b": int(_number(row[4]) or 0),
                    "c": int(_number(row[5]) or 0), "score": satisfaction_score,
                },
                "sewageCollection": {
                    "a": int(_number(row[7]) or 0), "b": int(_number(row[8]) or 0),
                    "c": int(_number(row[9]) or 0), "score": collection_score,
                },
            }

    for table_index, source_period in ((7, "2025年上半年度"), (21, "2025年下半年度")):
        for row in maonan["tables"][table_index - 1]["rows"][1:]:
            if len(row) < 7 or not row[1]:
                continue
            town = row[1].replace("水质净化厂", "").strip()
            result["facility"][("茂南项目", source_period, town, "town_plant")] = {
                "designScale": _number(row[3]), "averageDailyVolume": _number(row[5]), "processType": row[2]
            }
    for table_index, source_period in ((38, "2025年上半年度"), (39, "2025年下半年度")):
        current_town = ""
        for row in maonan["tables"][table_index - 1]["rows"][1:]:
            if len(row) < 7:
                continue
            if row[1]:
                current_town = row[1].replace("水质净化厂及配套管网", "").strip()
            if not current_town or not row[2]:
                continue
            month_match = re.search(r"(\d{4})年(\d{1,2})月", row[2])
            if not month_match:
                continue
            year, month = map(int, month_match.groups())
            facility = result["facility"].get(("茂南项目", source_period, current_town, "town_plant"), {})
            daily = facility.get("averageDailyVolume")
            result["payment"][("茂南项目", source_period, current_town)].append({
                "month": f"{year}-{month:02d}", "influentCod": _number(row[3]), "effluentCod": _number(row[4]),
                "legacyKq": _number(row[6]), "effluentQualified": True,
                "averageDailyVolumeCubicMeters": daily,
                "monthlyVolumeTenThousandCubicMeters": round(daily * calendar.monthrange(year, month)[1] / 10000, 4) if daily else None,
            })
    for row in maonan["tables"][104]["rows"][2:]:
        if len(row) < 12 or not row[1]:
            continue
        result["water"][("茂南项目", "2025年上半年度", row[1], None)].append({
            "sampleTime": row[2], "influentCod": _number(row[3]), "influentNh3n": _number(row[4]),
            "influentTp": _number(row[5]), "influentSs": _number(row[6]), "codValue": _number(row[7]),
            "nh3nValue": _number(row[8]), "tpValue": _number(row[9]), "ssValue": _number(row[10]),
            "qualified": row[11] == "是",
        })
        result["water"][("茂南项目", "2025年下半年度", row[1], None)] = list(
            result["water"][("茂南项目", "2025年上半年度", row[1], None)]
        )
    return result


def source_images() -> list[Path]:
    images: list[Path] = []
    for folder in (SOURCE_DIR / "yunan-media", SOURCE_DIR / "maonan-media"):
        if not folder.exists():
            continue
        for path in folder.iterdir():
            if path.suffix.lower() in {".jpg", ".jpeg", ".png"} and 5_000 < path.stat().st_size <= 500_000:
                images.append(path)
    if not images:
        raise RuntimeError("例文未提取到可用于全流程验证的图片。")
    return images


def _fresh_reason(item_name: str, deduction: float, point: str) -> str:
    return f"本期对{point}进行现场与资料复核时，确认“{item_name}”项存在扣分事实，按现行标准核减{deduction:g}分。"


def _option_entries(item: dict[str, Any], deduction: float, note: str) -> tuple[list[dict[str, Any]], str]:
    options = item.get("deductionOptions") or []
    target = round(float(deduction), 4)
    for option in options:
        if option.get("type") == "range":
            low = float(option.get("min") or 0)
            high = float(option.get("max") or item.get("fullScore") or target)
            if low - 0.0001 <= target <= high + 0.0001:
                return [{"optionId": option["id"], "selection": "standard", "rangeValue": target, "note": note}], "range"
    for option in options:
        value = float(option.get("deduction") or 0)
        if value <= 0:
            continue
        count = round(target / value)
        maximum = int(option.get("maxInstances") or max(count, 1))
        if 1 <= count <= maximum and abs(value * count - target) < 0.0001:
            return [{"optionId": option["id"], "selection": "standard", "instances": count, "note": note}], "fixed"
    if options:
        return [{"optionId": options[0]["id"], "selection": "standard", "adjustedScore": target, "adjustNote": note}], "adjusted"
    return [], "legacy"


def make_entries(record: dict[str, Any], standards: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    leaves = [item for item in standards["items"] if item["level"] == 3]
    by_code = {item["code"]: item for item in leaves}
    deductions: dict[str, float] = defaultdict(float)
    source_reasons: dict[str, list[str]] = defaultdict(list)
    source_items: dict[str, list[str]] = defaultdict(list)
    for item in record["deductions"]:
        deductions[item["itemCode"]] += float(item["deduction"])
        reason = str(item.get("reason") or "").strip()
        if reason and reason not in source_reasons[item["itemCode"]]:
            source_reasons[item["itemCode"]].append(reason)
        source_item = str(item.get("sourceItem") or "").strip()
        if source_item and source_item not in source_items[item["itemCode"]]:
            source_items[item["itemCode"]].append(source_item)
    diagnostics: list[dict[str, Any]] = []
    entries: dict[str, Any] = {}
    for item in leaves:
        target = min(round(deductions.get(item["code"], 0), 4), float(item["fullScore"]))
        entry = {"itemId": item["id"], "done": True, "options": []}
        if target > 0:
            reasons = source_reasons.get(item["code"]) or []
            source_item_names = source_items.get(item["code"]) or []
            if reasons:
                note = "；".join(reasons)
            elif source_item_names:
                note = f"例文评分表“{'、'.join(source_item_names)}”记录扣分{target:g}分。"
            else:
                note = _fresh_reason(item["name"], target, record["point"])
            options, mode = _option_entries(item, target, note)
            if options:
                entry["options"] = options
            else:
                entry["deduction"] = target
                entry["reason"] = note
            diagnostics.append({
                "item": item["name"], "deduction": target, "mode": mode,
                "usedSourceReason": bool(source_reasons.get(item["code"])),
                "usedSourceEvidence": bool(source_reasons.get(item["code"]) or source_items.get(item["code"])),
            })
        entries[item["id"]] = entry
    missing = sorted(set(deductions) - set(by_code))
    if missing:
        raise RuntimeError(f"{record['project']} {record['point']}存在未匹配评分项：{missing}")
    return entries, diagnostics


def _iso_sample_time(value: str, fallback_period: str) -> str:
    match = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", str(value or ""))
    if match:
        year, month, day = map(int, match.groups())
        return f"{year:04d}-{month:02d}-{day:02d}T09:30:00"
    year_match = re.search(r"(\d{4})年", fallback_period)
    year = year_match.group(1) if year_match else "2030"
    month_day_match = re.search(r"(\d{1,2})月(\d{1,2})日", str(value or ""))
    if month_day_match:
        month, day = map(int, month_day_match.groups())
        return f"{year}-{month:02d}-{day:02d}T09:30:00"
    return f"{year}-06-30T09:30:00"


def water_payload(record: dict[str, Any], metadata: dict[str, Any], period: str) -> dict[str, Any] | None:
    point = record["point"] if record["facilityType"] == "rural_treatment" else None
    samples = metadata["water"].get((record["project"], record["period"], record["town"], point), [])
    if record["facilityType"] not in {"town_plant", "rural_treatment"} or not samples:
        return None
    maonan = record["project"] == "茂南项目"
    rural = record["facilityType"] == "rural_treatment"
    facility = metadata["facility"].get((record["project"], record["period"], record["town"], "town_plant"), {})
    normalized: list[dict[str, Any]] = []
    for sample in samples:
        item = {
            "sampleTime": _iso_sample_time(sample["sampleTime"], period),
            "dischargeStandard": "DB44/2208-2019" if rural else "城镇污水处理厂污染物排放标准一级A及项目适用标准",
            "processType": "农村生活污水处理设施" if rural else facility.get("processType") or "以项目台账为准",
            "designScale": None if rural else facility.get("designScale"),
            "influentCod": sample.get("influentCod"),
            "influentNh3n": sample.get("influentNh3n"),
            "influentTp": sample.get("influentTp"),
            "codValue": sample.get("codValue"), "codLimit": 60 if rural else 40,
            "nh3nValue": sample.get("nh3nValue"), "nh3nLimit": 8 if rural else 5,
            "tpValue": sample.get("tpValue"), "tpLimit": None if rural else 0.5,
            "ssValue": sample.get("ssValue") if maonan else None,
            "ssLimit": 10 if maonan else None,
            "conclusion": "qualified" if sample.get("qualified") else "unqualified",
            "completed": True,
            "note": "实测值、取样时间和达标结论取自对应项目例文。",
        }
        normalized.append({key: value for key, value in item.items() if value is not None})
    result = dict(normalized[-1])
    result["samples"] = normalized
    result["note"] = f"共回填例文水质记录{len(normalized)}条；各条记录按实测值与适用限值逐项核对。"
    return result


def survey_payload(record: dict[str, Any], metadata: dict[str, Any]) -> dict[str, Any] | None:
    source = metadata["survey"].get((record["project"], record["period"], record["town"], record["point"]))
    if not source:
        return None
    count = source["count"]
    satisfaction = source["satisfaction"]
    collection = source["sewageCollection"]
    satisfaction_comment = f"例文调查{count}份：满意{satisfaction['a']}份，基本满意{satisfaction['b']}份，不满意{satisfaction['c']}份。"
    collection_comment = f"例文调查{count}份：有改善{collection['a']}份，一般{collection['b']}份，没有改善{collection['c']}份。"
    result: dict[str, Any] = {}
    if satisfaction.get("score") is not None:
        for respondent in ("villager1", "villager2"):
            result[f"satisfaction_{respondent}"] = {
                "score": satisfaction["score"], "comment": satisfaction_comment, "completed": True,
            }
    if collection.get("score") is not None:
        score = round(float(collection["score"]) / 2, 2)
        for respondent in ("villager1", "villager2", "gov_rep", "assessment_team"):
            result[f"sewage_collection_{respondent}"] = {
                "score": score, "comment": collection_comment, "completed": True,
            }
    return result or None


def make_payload(record: dict[str, Any], project: dict[str, Any], cycle_id: str, period: str, town: dict[str, Any], village: dict[str, Any] | None, standards: dict[str, Any], metadata: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    entries, diagnostics = make_entries(record, standards)
    facility = metadata["facility"].get((record["project"], record["period"], record["town"], record["facilityType"]), {})
    months = metadata["payment"].get((record["project"], record["period"], record["town"]), [])
    item: dict[str, Any] = {
        "village": village["name"] if village else "", "primaryFacilityType": record["facilityType"],
        "currentScore": record["currentRuleScore"], "entries": entries,
        "actualTreatmentVolume": facility.get("averageDailyVolume"), "volumePeriod": period,
        "designScale": facility.get("designScale"), "sourceScore": record.get("sourceScore"),
        "sourceTableIndex": record["sourceTableIndex"],
        "paymentData": {
            "firstPaymentPeriod": record["period"] in {"2025年第2季度", "2025年上半年度"},
            "designScaleCubicMetersPerDay": facility.get("designScale"), "months": months,
            "note": "金额输入取自对应项目例文中的结构化水量及水质数据。",
        },
    }
    water = water_payload(record, metadata, period)
    if water:
        item["waterQuality"] = water
    surveys = survey_payload(record, metadata)
    if surveys:
        item["surveyEntries"] = surveys
    return ({
        "schemaVersion": "1.0", "cityId": project["id"], "cycleId": cycle_id,
        "indicatorVersionId": standards["version"]["id"], "city": project["name"], "period": period,
        "town": town["name"], "villages": [item],
    }, diagnostics)


def _normalize_paragraph(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).strip()


def source_paragraphs(project: str) -> set[str]:
    filename = "yunan-full.json" if project == "郁南项目" else "maonan-full.json"
    data = json.loads((SOURCE_DIR / filename).read_text(encoding="utf-8"))
    return {
        text
        for item in data["paragraphs"]
        if len(text := _normalize_paragraph(item["text"])) >= 60
    }


def inspect_report(content: bytes, project: str, expected_points: list[str]) -> dict[str, Any]:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = [cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells]
    text = "\n".join([*paragraphs, *table_text])
    required = ["目录", "第一章 考核工作概述", "第二章", "第三章", "第四章", "附件1", "附件2", "附件3", "附件5"]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"报告缺少必要模块：{missing}")
    if "指标编号" in text:
        raise RuntimeError("报告中仍有禁用字段。")
    absent_points = [point for point in expected_points if point not in text]
    if absent_points:
        raise RuntimeError(f"报告遗漏项目点：{absent_points[:10]}")
    copied = sorted({
        paragraph for paragraph in paragraphs
        if len(_normalize_paragraph(paragraph)) >= 60 and _normalize_paragraph(paragraph) in source_paragraphs(project)
    })
    if copied:
        raise RuntimeError(f"生成报告仍直接复用了例文段落：{copied[:3]}")
    return {
        "paragraphCount": len(paragraphs), "tableCount": len(document.tables),
        "embeddedImageCount": sum(1 for rel in document.part.rels.values() if "image" in rel.reltype),
        "copiedSourceParagraphCount": len(copied), "fileSize": len(content),
    }


def generate_reports(client: httpx.Client, admin: dict[str, str], project: dict[str, Any], period: str, town_names: list[str], expected_points: list[str]) -> list[dict[str, Any]]:
    reports: list[dict[str, Any]] = []
    for outputs, selected_towns in ((["summary"], town_names), (["separate"], [town_names[0]])):
        payload = {"source": "dashboard", "projectId": project["id"], "period": period, "townNames": selected_towns, "outputs": outputs}
        precheck = require(client.post("/api/report-tasks/precheck", headers=admin, json=payload, timeout=180), "报告预检")
        if not precheck.get("ok"):
            raise RuntimeError(f"报告预检未通过：{precheck}")
        task = require(client.post("/api/report-tasks", headers=admin, json=payload, timeout=600), "生成报告")
        result = require(client.get(f"/api/report-tasks/{task['id']}", headers=admin, timeout=180), "读取报告任务")
        if result.get("status") != "completed":
            raise RuntimeError(f"报告生成失败：{result.get('error')}")
        for report in result.get("reports") or []:
            download = client.get(f"/api/reports/{report['id']}/download", headers=admin, timeout=300)
            if download.status_code != 200 or len(download.content) < 10_000:
                raise RuntimeError(f"报告下载失败：{report.get('name')}")
            output = RESULTS / report["name"]
            output.write_bytes(download.content)
            points = expected_points if outputs == ["summary"] else [point for point in expected_points if point == selected_towns[0]] or [selected_towns[0]]
            quality = inspect_report(download.content, project["name"], points)
            reports.append({"name": report["name"], "type": outputs[0], "path": str(output), **quality})
    return reports


def score_variances(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    variances: list[dict[str, Any]] = []
    for record in records:
        source = record.get("sourceScore")
        current = record.get("currentRuleScore")
        if source is None or current is None or abs(float(source) - float(current)) <= 0.001:
            continue
        expected = (
            record["project"] == "郁南项目"
            and record["town"] == "宋桂镇"
            and record["point"] == "井上村"
            and int(record.get("sourceTableIndex") or 0) == 145
            and abs(float(source) - 64.0) <= 0.01
            and abs(float(current) - 94.0) <= 0.01
        )
        reason = (
            "例文井上村评分表仅列明资料扣1分、公众调查扣5分，按100分制应得94分；"
            "同结构的两头村评分表亦按相同算法得94分，因此例文末行64分按源表异常记录，不作为系统算法。"
            if expected
            else "例文分数与当前评分标准计算结果不一致，尚未找到可验证的换算依据。"
        )
        variances.append({
            "project": record["project"], "sourcePeriod": record["sourcePeriod"],
            "town": record["town"], "point": record["point"], "facilityType": record["facilityType"],
            "sourceScore": source, "currentRuleScore": current,
            "sourceDeductionTotal": record.get("sourceDeductionTotal"),
            "expectedSourceAnomaly": expected, "reason": reason,
        })
    return variances


def main() -> None:
    fixture = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))
    metadata = source_metadata()
    images = source_images()
    if RESULTS.exists():
        shutil.rmtree(RESULTS)
    RESULTS.mkdir(parents=True)
    cycle_ids, periods = prepare_cycles()
    url = base_url()
    created_records: list[dict[str, Any]] = []
    diagnostics: list[dict[str, Any]] = []
    reports: list[dict[str, Any]] = []
    projects: dict[str, dict[str, Any]] = {}
    cleanup_error: Exception | None = None
    try:
        with httpx.Client(base_url=url, timeout=180) as client:
            inspector_token = require(client.post("/api/auth/login", json={"username": "inspector", "password": "Inspector@123456"}), "检查员登录")["token"]
            admin_token = require(client.post("/api/auth/login", json={"username": "admin", "password": "Admin@123456"}), "管理员登录")["token"]
            inspector = {"Authorization": f"Bearer {inspector_token}"}
            admin = {"Authorization": f"Bearer {admin_token}"}
            projects = {item["name"]: item for item in require(client.get("/api/mobile/projects", headers=admin), "读取项目")["items"]}
            towns_cache: dict[str, dict[str, Any]] = {}
            villages_cache: dict[str, dict[str, Any]] = {}
            standards_cache: dict[tuple[str, str, str], dict[str, Any]] = {}
            for index, record in enumerate(fixture["records"]):
                key = (record["project"], record["period"])
                project = projects[record["project"]]
                if record["project"] not in towns_cache:
                    towns = require(client.get("/api/mobile/towns", headers=admin, params={"city_id": project["id"]}), "读取镇街")["items"]
                    towns_cache[record["project"]] = {item["name"]: item for item in towns}
                town = towns_cache[record["project"]][record["town"]]
                village = None
                if record["facilityType"] == "rural_treatment":
                    if town["id"] not in villages_cache:
                        villages = require(client.get(f"/api/mobile/towns/{town['id']}/villages", headers=admin), "读取项目点")["items"]
                        villages_cache[town["id"]] = {item["name"]: item for item in villages}
                    village = villages_cache[town["id"]][record["point"]]
                standard_key = (record["project"], record["period"], record["facilityType"])
                if standard_key not in standards_cache:
                    standards_cache[standard_key] = require(client.get("/api/mobile/indicator-standards", headers=admin, params={
                        "city_id": project["id"], "cycle_id": cycle_ids[key], "facility_type": record["facilityType"],
                    }), "读取评分标准")
                payload, item_diagnostics = make_payload(
                    record, project, cycle_ids[key], periods[key], town, village, standards_cache[standard_key], metadata
                )
                created = require(client.post("/api/mobile/assessment-records", headers=inspector, json=payload, timeout=180), "保存移动端考核")
                record_id = created["recordIds"][0]
                image = images[index % len(images)]
                content_type = mimetypes.guess_type(image.name)[0] or "image/jpeg"
                require(client.post(
                    f"/api/mobile/assessment-records/{record_id}/attachments", headers=inspector,
                    files={"file": (f"{record['point']}现场照片{image.suffix.lower()}", image.read_bytes(), content_type)}, timeout=180,
                ), "上传现场照片")
                if record["facilityType"] == "town_plant":
                    detection = images[(index + 1) % len(images)]
                    detection_type = mimetypes.guess_type(detection.name)[0] or "image/jpeg"
                    require(client.post(
                        f"/api/mobile/assessment-records/{record_id}/attachments", headers=inspector,
                        files={"file": (f"{record['point']}水质检测页{detection.suffix.lower()}", detection.read_bytes(), detection_type)}, timeout=180,
                    ), "上传检测资料")
                require(client.post(f"/api/mobile/assessment-records/{record_id}/submit", headers=inspector), "提交考核")
                detail = require(client.get(f"/api/records/{record_id}", headers=admin), "读取考核结果")
                actual = round(float(detail["totalScore"]), 2)
                expected = round(float(record["currentRuleScore"]), 2)
                if actual != expected:
                    raise RuntimeError(f"{record['project']} {record['point']}分数不一致：期望{expected}，实际{actual}")
                require(client.post(f"/api/records/{record_id}/review", headers=admin), "平台复核")
                created_records.append({
                    "project": record["project"], "sourcePeriod": record["period"], "testPeriod": periods[key],
                    "town": record["town"], "point": record["point"], "facilityType": record["facilityType"],
                    "sourceScore": record.get("sourceScore"), "currentRuleScore": expected,
                    "sourceDeductionTotal": record["sourceDeductionTotal"],
                    "sourceTableIndex": record["sourceTableIndex"], "recordId": record_id,
                    "waterQualityCount": len(detail.get("waterQuality") or []),
                    "surveyRecordCount": len(detail.get("surveys") or []),
                })
                diagnostics.extend({"project": record["project"], "point": record["point"], **item} for item in item_diagnostics)

            for key in cycle_ids:
                project_name, source_period = key
                period_records = [item for item in fixture["records"] if item["project"] == project_name and item["period"] == source_period]
                town_names = list(dict.fromkeys(item["town"] for item in period_records))
                point_names = list(dict.fromkeys(item["point"] for item in period_records))
                reports.extend(generate_reports(client, admin, projects[project_name], periods[key], town_names, point_names))
            cleanup(client, admin, cycle_ids, projects)
    except Exception:
        if projects:
            try:
                with httpx.Client(base_url=url, timeout=180) as cleanup_client:
                    token = require(cleanup_client.post("/api/auth/login", json={"username": "admin", "password": "Admin@123456"}), "清理登录")["token"]
                    cleanup(cleanup_client, {"Authorization": f"Bearer {token}"}, cycle_ids, projects)
            except Exception as exc:
                cleanup_error = exc
        raise
    finally:
        try:
            verify_cleanup(cycle_ids)
        except Exception as exc:
            cleanup_error = cleanup_error or exc
        if cleanup_error:
            raise cleanup_error

    adjusted_count = sum(1 for item in diagnostics if item["mode"] == "adjusted")
    legacy_count = sum(1 for item in diagnostics if item["mode"] == "legacy")
    source_reason_count = sum(1 for item in diagnostics if item.get("usedSourceReason"))
    source_evidence_count = sum(1 for item in diagnostics if item.get("usedSourceEvidence"))
    water_quality_count = sum(item["waterQualityCount"] for item in created_records)
    survey_record_count = sum(item["surveyRecordCount"] for item in created_records)
    if legacy_count:
        raise RuntimeError(f"有{legacy_count}条扣分无法通过移动端扣分选项表达。")
    variances = score_variances(created_records)
    unexpected_variances = [item for item in variances if not item["expectedSourceAnomaly"]]
    if unexpected_variances:
        raise RuntimeError(f"仍有{len(unexpected_variances)}处分数差异无法解释：{unexpected_variances[:3]}")
    comparison = {
        "passed": True, "recordCount": len(created_records), "deductionCount": len(diagnostics),
        "adjustedOptionCount": adjusted_count, "legacyOptionCount": legacy_count,
        "sourceReasonCount": source_reason_count,
        "sourceEvidenceCount": source_evidence_count,
        "waterQualityRecordCount": water_quality_count,
        "surveyRecordCount": survey_record_count,
        "sourceScoreVarianceCount": len(variances),
        "expectedSourceScoreVarianceCount": len(variances) - len(unexpected_variances),
        "sourceDataAnomalyCount": len(variances) - len(unexpected_variances),
        "unexpectedSourceScoreVarianceCount": len(unexpected_variances),
        "scoreVariances": variances,
        "periods": {f"{project}|{source_period}": period for (project, source_period), period in periods.items()},
        "records": created_records, "reports": reports,
        "generatedAt": datetime.now().isoformat(timespec="seconds"),
    }
    (RESULTS / "全流程对比结果.json").write_text(json.dumps(comparison, ensure_ascii=False, indent=2), encoding="utf-8")
    if DELIVERY.exists():
        shutil.rmtree(DELIVERY)
    DELIVERY.mkdir(parents=True)
    for report in reports:
        if report["type"] == "summary":
            shutil.copy2(report["path"], DELIVERY / Path(report["path"]).name)
    shutil.copy2(RESULTS / "全流程对比结果.json", DELIVERY / "全流程对比结果.json")
    print(json.dumps({
        "passed": True, "recordCount": len(created_records), "deductionCount": len(diagnostics),
        "sourceReasonCount": source_reason_count, "sourceEvidenceCount": source_evidence_count,
        "waterQualityRecordCount": water_quality_count,
        "surveyRecordCount": survey_record_count, "adjustedOptionCount": adjusted_count,
        "sourceScoreVarianceCount": len(variances),
        "unexpectedSourceScoreVarianceCount": len(unexpected_variances),
        "reportCount": len(reports), "delivery": str(DELIVERY),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
