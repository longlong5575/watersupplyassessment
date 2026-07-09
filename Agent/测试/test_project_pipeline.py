from __future__ import annotations

import os
import re
import sys
import base64
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"


def default_runtime_root() -> Path:
    if os.environ.get("WATERSUPPLY_RUNTIME_DIR"):
        return Path(os.environ["WATERSUPPLY_RUNTIME_DIR"])
    base = ROOT.parent.parent if ROOT.parent.name.lower() == "watersupplyassessment" else ROOT.parent
    return base / "运行脚本" / "watersupply-agent-runtime"


RESULTS = default_runtime_root() / "test-results" / "project-pipeline"
RESULTS.mkdir(parents=True, exist_ok=True)
DB_PATH = RESULTS / "project-pipeline-test.db"
if DB_PATH.exists():
    DB_PATH.unlink()
os.environ["DATABASE_URL"] = f"sqlite:///{DB_PATH.as_posix()}"
os.environ["STORAGE_DIR"] = str((RESULTS / "storage").resolve())
os.environ["CELERY_TASK_ALWAYS_EAGER"] = "true"
sys.path.insert(0, str(BACKEND))

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


INDICATOR_GROUPS: dict[str, list[dict]] = {}


def assert_ok(response, label: str):
    assert response.status_code < 300, f"{label}: {response.status_code} {response.text}"
    return response.json()


def login(client: TestClient, username: str) -> dict[str, str]:
    data = assert_ok(client.post("/api/auth/login", json={"username": username}), f"login {username}")
    return {"Authorization": f"Bearer {data['token']}"}


def leaf_standards(client: TestClient, project_id: str, cycle_id: str, facility_type: str):
    data = assert_ok(
        client.get("/api/mobile/indicator-standards", params={"city_id": project_id, "cycle_id": cycle_id, "facility_type": facility_type}),
        f"standards {facility_type}",
    )
    leaves = [item for item in data["items"] if item["level"] == 3]
    assert leaves and round(sum(float(item["fullScore"]) for item in leaves), 2) == 100
    return data, leaves


def record_payload(*, project, cycle, town, village, facility_type, indicator, note="现场复核发现该项存在不符合评分标准要求的情况"):
    indicators = INDICATOR_GROUPS.get(indicator["id"], [indicator])
    payload = {
        "schemaVersion": "1.0",
        "cityId": project["id"],
        "cycleId": cycle["id"],
        "city": project["name"],
        "period": cycle["name"],
        "town": town,
        "villages": [{
            "village": village,
            "primaryFacilityType": facility_type,
            "currentScore": 0,
            "waterQuality": {
                "sampleTime": "2026-06-30T09:00:00",
                "dischargeStandard": "自动带出标准",
                "codValue": "39",
                "codLimit": "40",
                "nh3nValue": "4.8",
                "nh3nLimit": "5（8）",
                "tpValue": "0.4",
                "tpLimit": "0.5",
                "conclusion": "qualified",
                "completed": True,
            },
            "entries": {
                item["id"]: {
                    "itemId": item["id"],
                    "done": True,
                    "options": ([
                        {"optionId": item["deductionOptions"][0]["id"], "selection": "standard", "instances": 99, "note": note},
                        {"optionId": item["deductionOptions"][0]["id"], "selection": "standard", "instances": 99, "note": "资料记录与现场情况不一致，需整改闭环"},
                    ] if item["id"] == indicator["id"] and item.get("deductionOptions") else []),
                }
                for item in indicators
            },
        }],
    }
    if facility_type == "rural_treatment":
        payload["villages"][0]["surveyEntries"] = {
            "satisfaction_villager1": {"score": 5, "comment": "满意", "completed": True},
            "sewage_collection_villager1": {"score": 4, "comment": "有改善", "completed": True},
        }
    return payload


def complete_payload(*, project, cycle, town, village, facility_type, indicators):
    return {
        "schemaVersion": "1.0",
        "cityId": project["id"],
        "cycleId": cycle["id"],
        "city": project["name"],
        "period": cycle["name"],
        "town": town,
        "villages": [{
            "village": village,
            "primaryFacilityType": facility_type,
            "currentScore": 100,
            "entries": {
                item["id"]: {"itemId": item["id"], "done": True, "options": []}
                for item in indicators
            },
        }],
    }


def create_record(client: TestClient, headers: dict[str, str], *, project, cycle, town, village, facility_type, indicator):
    payload = record_payload(project=project, cycle=cycle, town=town, village=village, facility_type=facility_type, indicator=indicator)
    created = assert_ok(client.post("/api/mobile/assessment-records", json=payload, headers=headers), "create assessment")
    record_id = created["recordIds"][0]
    assert_ok(client.post(f"/api/mobile/assessment-records/{record_id}/submit", headers=headers), "submit assessment")
    return record_id


def check_docx(path: Path, town: str, project_name: str):
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    all_text = f"{text}\n{table_text}"
    assert town in all_text
    assert ("摘要" in all_text) or ("摘  要" in all_text)
    assert "目录" in all_text
    assert "附件1 考核标准" in all_text
    assert "项目人员组成" in all_text
    assert "1.6.1 现场检查" in all_text
    assert "1.6.2 查阅资料" in all_text
    if project_name == "郁南项目":
        assert "镇级及农村设施考核报告" in all_text
        assert "公众调查" in all_text
        assert "农村污水处理设施" in all_text
        assert "DB44/2208-2019" in all_text
        assert "第一章 考核工作概述" in all_text
        assert "第二章 镇级设施运维考核情况" in all_text
        assert "第三章 考核评价系数的确定" in all_text
        assert "第四章 主要问题及整改建议" in all_text
        assert "运维绩效考核系数" in all_text
        assert "不引用其他项目金额代算" in all_text
        assert "附件2 考核评分表" in all_text
        assert "附件3 现场照片" in all_text
        assert "附件5 水质抽检情况汇总表" in all_text
    if project_name == "茂南项目":
        assert "城镇设施绩效考核报告" in all_text
        assert "水质净化厂" in all_text
        assert "问卷调查（村级考核有）" not in all_text
        assert "第一章 考核工作概述" in all_text
        assert "第二章 城镇水质净化设施考核结果" in all_text
        assert "第三章 绩效付费计算" in all_text
        assert "第四章 主要改进点、主要问题和整改工作建议" in all_text
        assert "附件2 周期评分表" in all_text
        assert "附件3 现场检查照片" in all_text
        assert "附件5 水质抽检汇总" in all_text
        assert "附件8 月平均值统计" in all_text
    assert "fixture-photo.png" in all_text
    assert "Agent辅助校验" not in all_text
    assert "系统采集" not in all_text
    for table in document.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if "序号" in headers:
            index = headers.index("序号")
            serials = [row.cells[index].text.strip() for row in table.rows[1:]]
            assert serials == [str(number) for number in range(1, len(serials) + 1)]


def main():
    with TestClient(app) as client:
        inspector = login(client, "inspector")
        admin = login(client, "admin")
        projects = assert_ok(client.get("/api/mobile/projects"), "projects")["items"]
        assert {item["name"] for item in projects} == {"郁南项目", "茂南项目"}
        by_name = {item["name"]: item for item in projects}

        yunan = by_name["郁南项目"]
        maonan = by_name["茂南项目"]
        yunan_cycle = assert_ok(client.get("/api/mobile/assessment-cycles", params={"city_id": yunan["id"]}), "yunan cycles")["items"][0]
        maonan_cycle = assert_ok(client.get("/api/mobile/assessment-cycles", params={"city_id": maonan["id"]}), "maonan cycles")["items"][0]

        yunan_towns = assert_ok(client.get("/api/mobile/towns", params={"city_id": yunan["id"]}), "yunan towns")["items"]
        maonan_towns = assert_ok(client.get("/api/mobile/towns", params={"city_id": maonan["id"]}), "maonan towns")["items"]
        assert len(yunan_towns) == 16 and len(maonan_towns) == 7
        yunan_by_name = {item["name"]: item for item in yunan_towns}
        maonan_by_name = {item["name"]: item for item in maonan_towns}
        assert yunan_by_name["桂圩镇"]["chapterCode"] == "2.3"
        assert set(yunan_by_name["桂圩镇"]["assessmentTargets"]) == {"town_plant", "town_network", "rural_treatment"}
        assert maonan_by_name["金塘镇"]["assessmentTargets"] == ["town_plant"]
        assert maonan_by_name["中科云粤西产业园"]["assessmentTargets"] == ["town_network"]
        assert sum(len(item["assessmentTargets"]) for item in maonan_towns) == 12

        village_total = 0
        for town in yunan_towns:
            villages = assert_ok(client.get(f"/api/mobile/towns/{town['id']}/villages"), f"villages {town['name']}")["items"]
            village_total += len(villages)
        assert village_total == 59
        guiwei_villages = assert_ok(client.get(f"/api/mobile/towns/{yunan_by_name['桂圩镇']['id']}/villages"), "guiwei villages")["items"]
        assert [(item["administrativeVillage"], item["name"]) for item in guiwei_villages] == [
            ("新塘村", "山禾地村"), ("䓣口村", "赤坭村"), ("䓣口村", "高寨村"), ("䓣口村", "平山村"), ("䓣口村", "道枝村")
        ]
        template = assert_ok(client.get(f"/api/mobile/projects/{yunan['id']}/report-template"), "project template")
        guiwei_template = next(item for item in template["towns"] if item["name"] == "桂圩镇")
        assert guiwei_template["reportTemplate"]["assessmentObjectSection"] == "2.3.1"

        standards = {}
        for project, cycle, types in [
            (yunan, yunan_cycle, ["town_plant", "town_network", "rural_treatment"]),
            (maonan, maonan_cycle, ["town_plant", "town_network"]),
        ]:
            for facility_type in types:
                standards[(project["name"], facility_type)] = leaf_standards(client, project["id"], cycle["id"], facility_type)

        for _, leaves in standards.values():
            option_names = [option["name"] for item in leaves for option in item.get("deductionOptions", [])]
            assert not any("检查单元" in name or "抽查5个井段" in name for name in option_names), "抽检规则不应成为扣分选项"
            assert all(item.get("evaluationStandard") or item.get("description") for item in leaves), "每个评分点必须有知识库说明"
            missing_options = [item["name"] for item in leaves if not item.get("deductionOptions")]
            assert not missing_options, f"每个评分点必须有可选择的扣分原因: {missing_options}"
        maonan_network = standards[("茂南项目", "town_network")][1]
        overflow_item = next(item for item in maonan_network if "无污水冒出" in item["name"])
        overflow_options = [option["name"] for option in overflow_item["deductionOptions"]]
        assert overflow_options == ["发现一处扣0.2分"], overflow_options
        unit_options = [
            option
            for _, leaves in standards.values()
            for item in leaves
            for option in item.get("deductionOptions") or []
            if option.get("unit")
        ]
        assert unit_options and all(option.get("maxInstances") for option in unit_options)
        all_options = [
            option
            for _, leaves in standards.values()
            for item in leaves
            for option in item.get("deductionOptions") or []
        ]
        assert all(option.get("name", "").strip() and float(option.get("deduction") or 0) > 0 for option in all_options)
        count_markers = re.compile(r"每(?:缺少|发现|出现|增加|有|个|一|处|项|次|座|人|岗位)")
        count_options = [option for option in all_options if count_markers.search(option["name"])]
        assert count_options and all(option.get("unit") and option.get("maxInstances") for option in count_options)
        for project_name in ("郁南项目", "茂南项目"):
            plant_items = standards[(project_name, "town_plant")][1]
            water_quality = next(item for item in plant_items if item["name"] == "污水处理质量")
            full_score = float(water_quality["fullScore"])
            assert any(
                "判定为不合格扣" in option["name"] and float(option["deduction"]) == full_score
                for option in water_quality["deductionOptions"]
            )
        split_option_items = [
            item
            for _, leaves in standards.values()
            for item in leaves
            if len(item.get("deductionOptions") or []) >= 2
        ]
        assert split_option_items
        split_options = split_option_items[0]["deductionOptions"]
        assert all(not re.match(r"^\d+\.\s", option["name"]) for option in split_options)
        assert all(float(option["deduction"]) > 0 for option in split_options)
        for _, leaves in standards.values():
            for leaf in leaves:
                INDICATOR_GROUPS[leaf["id"]] = leaves

        # Scoring stress cases: full score, capped legacy input, capped counts,
        # water-quality full deduction, survey replacement, and review overrides.
        stress_towns = [town for town in yunan_towns if "town_plant" in town["assessmentTargets"]]
        assert len(stress_towns) >= 4
        plant_leaves = standards[("郁南项目", "town_plant")][1]

        full_payload = complete_payload(
            project=yunan, cycle=yunan_cycle, town=stress_towns[0]["name"], village="",
            facility_type="town_plant", indicators=plant_leaves,
        )
        full_created = assert_ok(client.post("/api/mobile/assessment-records", json=full_payload, headers=inspector), "full-score assessment")
        full_detail = assert_ok(client.get(f"/api/records/{full_created['recordIds'][0]}"), "full-score detail")
        assert float(full_detail["totalScore"]) == 100
        assert all(float(item["deduction"]) == 0 for item in full_detail["scores"])

        legacy_payload = complete_payload(
            project=yunan, cycle=yunan_cycle, town=stress_towns[1]["name"], village="",
            facility_type="town_plant", indicators=plant_leaves,
        )
        legacy_indicator = plant_leaves[0]
        legacy_payload["villages"][0]["entries"][legacy_indicator["id"]]["deduction"] = 999999
        legacy_created = assert_ok(client.post("/api/mobile/assessment-records", json=legacy_payload, headers=inspector), "legacy overflow assessment")
        legacy_detail = assert_ok(client.get(f"/api/records/{legacy_created['recordIds'][0]}"), "legacy overflow detail")
        legacy_score = next(item for item in legacy_detail["scores"] if item["indicatorId"] == legacy_indicator["id"])
        assert float(legacy_score["deduction"]) == float(legacy_indicator["fullScore"])
        assert float(legacy_score["score"]) == 0
        assert float(legacy_detail["totalScore"]) == 100 - float(legacy_indicator["fullScore"])

        count_indicator = next(
            item for item in plant_leaves
            if any(option.get("unit") and option.get("maxInstances") for option in item.get("deductionOptions") or [])
        )
        count_option = next(option for option in count_indicator["deductionOptions"] if option.get("unit") and option.get("maxInstances"))
        count_payload = complete_payload(
            project=yunan, cycle=yunan_cycle, town=stress_towns[2]["name"], village="",
            facility_type="town_plant", indicators=plant_leaves,
        )
        count_payload["villages"][0]["entries"][count_indicator["id"]]["options"] = [{
            "optionId": count_option["id"], "selection": "standard",
            "instances": int(count_option["maxInstances"]) + 1000, "rangeValue": 0,
            "note": "现场发现多处同类问题，按评分标准上限扣分",
        }]
        count_created = assert_ok(client.post("/api/mobile/assessment-records", json=count_payload, headers=inspector), "count overflow assessment")
        count_detail = assert_ok(client.get(f"/api/records/{count_created['recordIds'][0]}"), "count overflow detail")
        count_score = next(item for item in count_detail["scores"] if item["indicatorId"] == count_indicator["id"])
        expected_count_deduction = min(
            float(count_indicator["fullScore"]),
            float(count_option["deduction"]) * int(count_option["maxInstances"]),
        )
        assert float(count_score["deduction"]) == expected_count_deduction
        adjusted_payload = complete_payload(
            project=yunan, cycle=yunan_cycle, town=stress_towns[2]["name"], village="",
            facility_type="town_plant", indicators=plant_leaves,
        )
        adjusted_payload["villages"][0]["entries"][count_indicator["id"]]["options"] = [{
            "optionId": count_option["id"], "selection": "standard", "instances": 999,
            "adjustedScore": 0.7, "note": "经复核按实际影响程度调整扣分",
        }]
        adjusted_created = assert_ok(client.post("/api/mobile/assessment-records", json=adjusted_payload, headers=inspector), "adjusted score assessment")
        assert adjusted_created["recordIds"] == count_created["recordIds"]
        adjusted_detail = assert_ok(client.get(f"/api/records/{adjusted_created['recordIds'][0]}"), "adjusted score detail")
        adjusted_score = next(item for item in adjusted_detail["scores"] if item["indicatorId"] == count_indicator["id"])
        assert float(adjusted_score["deduction"]) == 0.7

        water_indicator = next(item for item in plant_leaves if item["name"] == "污水处理质量")
        water_option = next(option for option in water_indicator["deductionOptions"] if float(option["deduction"]) == float(water_indicator["fullScore"]))
        water_payload = complete_payload(
            project=yunan, cycle=yunan_cycle, town=stress_towns[3]["name"], village="",
            facility_type="town_plant", indicators=plant_leaves,
        )
        water_payload["villages"][0]["waterQuality"] = {
            "sampleTime": "2026-06-30T23:59:59", "codValue": "9999", "codLimit": "40",
            "nh3nValue": "9999", "nh3nLimit": "5", "tpValue": "9999", "tpLimit": "0.5",
            "conclusion": "unqualified", "completed": True,
        }
        water_payload["villages"][0]["entries"][water_indicator["id"]]["options"] = [{
            "optionId": water_option["id"], "selection": "standard", "instances": 1,
            "note": "水质抽检结果不符合本项考核要求",
        }]
        water_created = assert_ok(client.post("/api/mobile/assessment-records", json=water_payload, headers=inspector), "water-quality assessment")
        water_detail = assert_ok(client.get(f"/api/records/{water_created['recordIds'][0]}"), "water-quality detail")
        water_score = next(item for item in water_detail["scores"] if item["indicatorId"] == water_indicator["id"])
        assert float(water_score["deduction"]) == float(water_indicator["fullScore"])
        assert float(water_detail["totalScore"]) == 100 - float(water_indicator["fullScore"])

        review_score_id = legacy_score["id"]
        reviewed_scores = assert_ok(client.put(
            f"/api/records/{legacy_created['recordIds'][0]}/scores", headers=admin,
            json={"scores": [{"id": review_score_id, "deduction": 999999, "reason": "后台复核按评分标准上限扣分"}], "reason": "复核扣分上限"},
        ), "review score cap")
        reviewed_item = next(item for item in reviewed_scores["scores"] if item["id"] == review_score_id)
        assert float(reviewed_item["deduction"]) == float(reviewed_item["indicatorFullScore"])
        assert float(reviewed_item["score"]) == 0
        assert 0 <= float(reviewed_scores["totalScore"]) <= 100

        rural_town = next(town for town in yunan_towns if town["name"] != "桂圩镇" and "rural_treatment" in town["assessmentTargets"])
        rural_villages = assert_ok(client.get(f"/api/mobile/towns/{rural_town['id']}/villages"), "stress rural villages")["items"]
        assert rural_villages
        rural_leaves = standards[("郁南项目", "rural_treatment")][1]
        survey_payload = complete_payload(
            project=yunan, cycle=yunan_cycle, town=rural_town["name"], village=rural_villages[0]["name"],
            facility_type="rural_treatment", indicators=rural_leaves,
        )
        survey_payload["villages"][0]["surveyEntries"] = {
            f"{category}_{respondent}": {"score": score, "comment": "问卷结果已按评分规则回填", "completed": True}
            for category, respondents, score in [
                ("sewage_collection", ["villager1", "villager2", "gov_rep", "assessment_team"], 1),
                ("overall_effect", ["villager1", "villager2", "gov_rep", "assessment_team"], 3),
                ("satisfaction", ["villager1", "villager2", "gov_rep", "implementation_org"], 5),
            ]
            for respondent in respondents
        }
        survey_created = assert_ok(client.post("/api/mobile/assessment-records", json=survey_payload, headers=inspector), "survey backfill assessment")
        survey_detail = assert_ok(client.get(f"/api/records/{survey_created['recordIds'][0]}"), "survey backfill detail")
        indicator_ids = [item["indicatorId"] for item in survey_detail["scores"]]
        assert len(indicator_ids) == len(set(indicator_ids)), "问卷回填后不应重复生成同一评分点"
        assert set(indicator_ids) == {item["id"] for item in rural_leaves}, "问卷评分不得串用其他项目或考核对象的指标"
        assert round(float(survey_detail["totalScore"]), 2) == round(sum(float(item["score"]) for item in survey_detail["scores"]), 2)
        assert 0 < float(survey_detail["totalScore"]) < 100
        survey_total = float(survey_detail["totalScore"])
        entries_only_update = assert_ok(client.put(
            f"/api/records/{survey_created['recordIds'][0]}", headers=admin,
            json={"data": {"entries": survey_payload["villages"][0]["entries"]}, "reason": "仅更新评分明细"},
        ), "preserve survey scores after entries update")
        assert any(item["source"] == "survey" for item in entries_only_update["scores"])
        assert round(float(entries_only_update["totalScore"]), 2) == round(survey_total, 2)
        cleared_surveys = assert_ok(client.put(
            f"/api/mobile/assessment-records/{survey_created['recordIds'][0]}/surveys", headers=inspector, json={},
        ), "clear survey scores")
        cleared_detail = assert_ok(client.get(f"/api/records/{survey_created['recordIds'][0]}"), "cleared survey detail")
        assert not any(item["source"] == "survey" for item in cleared_detail["scores"])
        assert {item["indicatorId"] for item in cleared_detail["scores"]} == {item["id"] for item in rural_leaves}
        assert float(cleared_surveys["totalScore"]) == 100

        cleared_scores = assert_ok(client.put(
            f"/api/mobile/assessment-records/{full_created['recordIds'][0]}/scores", headers=inspector, json={"entries": {}},
        ), "clear score entries")
        assert float(cleared_scores["totalScore"]) == 0

        # Clearing submitted data is scoped to one project and cycle and removes
        # dependent review/agent/report rows plus managed files.
        adjusted_record_id = adjusted_created["recordIds"][0]
        assert_ok(client.post(f"/api/mobile/assessment-records/{adjusted_record_id}/submit", headers=inspector), "submit clear fixture")
        attachment = client.post(
            f"/api/mobile/assessment-records/{adjusted_record_id}/attachments",
            headers=inspector,
            data={"score_id": adjusted_score["id"], "deduction_option_id": adjusted_score["deductionOptionId"] or ""},
            files={"file": ("clear-fixture.jpg", b"clear fixture", "image/jpeg")},
        )
        assert_ok(attachment, "upload clear fixture")
        assert_ok(client.post(f"/api/records/{adjusted_record_id}/review", headers=admin), "review clear fixture")
        clear_agent = assert_ok(client.post(f"/api/agent/records/{adjusted_record_id}/analysis", headers=admin), "clear fixture agent run")
        assert clear_agent["recordId"] == adjusted_record_id
        clear_task = assert_ok(client.post("/api/report-tasks", headers=admin, json={
            "source": "dashboard", "projectId": yunan["id"], "period": yunan_cycle["name"],
            "townNames": [stress_towns[2]["name"]], "outputs": ["separate", "summary"],
        }), "create clear fixture reports")
        clear_result = assert_ok(client.get(f"/api/report-tasks/{clear_task['id']}"), "clear fixture report result")
        assert clear_result["status"] == "completed"
        from app.core.database import SessionLocal
        from app.models import Attachment, Report
        with SessionLocal() as session:
            clear_report_paths = [Path(session.get(Report, item["id"]).storage_key) for item in clear_result["reports"]]
            clear_attachment_path = Path(session.get(Attachment, attachment.json()["id"]).storage_key)
        assert all(path.is_file() for path in [*clear_report_paths, clear_attachment_path])

        wrong_scope = client.delete("/api/mobile/assessment-records", headers=inspector, params={
            "city_id": yunan["id"], "cycle_id": maonan_cycle["id"], "period": yunan_cycle["name"],
        })
        assert wrong_scope.status_code == 422
        cleared = assert_ok(client.delete("/api/mobile/assessment-records", headers=inspector, params={
            "city_id": yunan["id"], "cycle_id": yunan_cycle["id"], "period": yunan_cycle["name"],
        }), "clear project cycle data")
        assert cleared["recordCount"] >= 5
        assert cleared["reportCount"] == len(clear_result["reports"])
        assert all(not path.exists() for path in [*clear_report_paths, clear_attachment_path])
        after_clear = assert_ok(client.get("/api/mobile/assessment-records", headers=inspector, params={
            "city_id": yunan["id"], "cycle_id": yunan_cycle["id"],
        }), "records after project clear")
        assert after_clear["items"] == []

        invalid_payload = {"cityId": yunan["id"], "cycleId": yunan_cycle["id"], "town": "金塘镇", "primaryFacilityType": "town_plant"}
        invalid = client.post("/api/mobile/assessment-records", json=invalid_payload, headers=inspector)
        assert invalid.status_code == 422

        yunan_indicator = standards[("郁南项目", "rural_treatment")][1][0]
        maonan_indicator = standards[("茂南项目", "town_plant")][1][0]
        maonan_network_indicator = standards[("茂南项目", "town_network")][1][0]
        yunan_record = create_record(client, inspector, project=yunan, cycle=yunan_cycle, town="桂圩镇", village="山禾地村", facility_type="rural_treatment", indicator=yunan_indicator)
        maonan_record = create_record(client, inspector, project=maonan, cycle=maonan_cycle, town="金塘镇", village="", facility_type="town_plant", indicator=maonan_indicator)

        incomplete_payload = record_payload(project=maonan, cycle=maonan_cycle, town="山阁镇", village="", facility_type="town_plant", indicator=maonan_indicator)
        incomplete_payload["villages"][0]["entries"][maonan_indicator["id"]]["done"] = False
        incomplete_created = assert_ok(client.post("/api/mobile/assessment-records", json=incomplete_payload, headers=inspector), "create incomplete assessment")
        incomplete_submit = client.post(f"/api/mobile/assessment-records/{incomplete_created['recordIds'][0]}/submit", headers=inspector)
        assert incomplete_submit.status_code == 409

        # A fresh browser/device must recover all submitted targets from the backend.
        maonan_plant = create_record(client, inspector, project=maonan, cycle=maonan_cycle, town="茂南区", village="", facility_type="town_plant", indicator=maonan_indicator)
        maonan_network = create_record(client, inspector, project=maonan, cycle=maonan_cycle, town="茂南区", village="", facility_type="town_network", indicator=maonan_network_indicator)
        maonan_submitted_only = create_record(client, inspector, project=maonan, cycle=maonan_cycle, town="镇盛镇", village="", facility_type="town_plant", indicator=maonan_indicator)
        submitted_only_detail = assert_ok(client.get(f"/api/records/{maonan_submitted_only}"), "submitted-only detail")
        assert submitted_only_detail["status"] == "submitted"
        restored = assert_ok(client.get("/api/mobile/assessment-records", headers=inspector, params={"city_id": maonan["id"], "cycle_id": maonan_cycle["id"]}), "restore mobile records")["items"]
        restored_maonan = [item for item in restored if item["town"] == "茂南区"]
        assert {item["raw"]["primaryFacilityType"] for item in restored_maonan} == {"town_plant", "town_network"}
        assert all(item["editable"] is True for item in restored_maonan)

        # Re-submitting an editable target updates the same record instead of creating a duplicate.
        updated_payload = record_payload(project=maonan, cycle=maonan_cycle, town="茂南区", village="", facility_type="town_plant", indicator=maonan_indicator, note="复核后已更新现场检查记录")
        updated = assert_ok(client.post("/api/mobile/assessment-records", json=updated_payload, headers=inspector), "update existing assessment")
        assert updated["recordIds"] == [maonan_plant]
        restored_after_update = assert_ok(client.get("/api/mobile/assessment-records", headers=inspector, params={"city_id": maonan["id"], "cycle_id": maonan_cycle["id"]}), "restore updated records")["items"]
        assert len([item for item in restored_after_update if item["town"] == "茂南区"]) == 2

        # Locked records remain visible on mobile but cannot be edited or submitted again.
        assert_ok(client.post(f"/api/records/{maonan_network}/review", headers=admin), "review network")
        reviewed_items = assert_ok(client.get("/api/mobile/assessment-records", headers=inspector, params={"city_id": maonan["id"], "cycle_id": maonan_cycle["id"]}), "restore reviewed records")["items"]
        reviewed_network = next(item for item in reviewed_items if item["id"] == maonan_network)
        assert reviewed_network["status"] == "reviewed" and reviewed_network["editable"] is False
        assert_ok(client.post(f"/api/records/{maonan_network}/lock", headers=admin), "lock network")
        locked_items = assert_ok(client.get("/api/mobile/assessment-records", headers=inspector, params={"city_id": maonan["id"], "cycle_id": maonan_cycle["id"]}), "restore locked records")["items"]
        locked_network = next(item for item in locked_items if item["id"] == maonan_network)
        assert locked_network["status"] == "locked" and locked_network["editable"] is False
        locked_update = client.post("/api/mobile/assessment-records", json=record_payload(project=maonan, cycle=maonan_cycle, town="茂南区", village="", facility_type="town_network", indicator=maonan_network_indicator), headers=inspector)
        assert locked_update.status_code == 422

        for record_id in [yunan_record, maonan_record]:
            detail = assert_ok(client.get(f"/api/records/{record_id}"), "record detail")
            assert len(detail["scores"]) >= 1
            manual_score = next(item for item in detail["scores"] if item["source"] == "manual")
            assert float(manual_score["deduction"]) <= float(manual_score["indicatorFullScore"])
            total_deduction = sum(float(item["deduction"]) for item in detail["scores"])
            assert round(float(detail["totalScore"]), 2) == round(max(100 - total_deduction, 0), 2)
            upload = client.post(
                f"/api/mobile/assessment-records/{record_id}/attachments",
                headers=inspector,
                data={"score_id": manual_score["id"], "deduction_option_id": manual_score["deductionOptionId"] or ""},
                files={"file": (
                    "fixture-photo.png",
                    base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="),
                    "image/png",
                )},
            )
            assert_ok(upload, "upload attachment")
            assert_ok(client.post(f"/api/records/{record_id}/review", headers=admin), "review")
            blocked = client.post(f"/api/mobile/assessment-records/{record_id}/submit", headers=inspector)
            assert blocked.status_code == 409
            assert_ok(client.post(f"/api/records/{record_id}/return", headers=admin, json={"reason": "资料需补充后重新提交", "data": {}}), "return")
            returned_items = assert_ok(client.get("/api/mobile/assessment-records", headers=inspector), "restore returned record")["items"]
            returned_record = next(item for item in returned_items if item["id"] == record_id)
            assert returned_record["status"] == "returned" and returned_record["editable"] is True
            assert_ok(client.post(f"/api/mobile/assessment-records/{record_id}/submit", headers=inspector), "resubmit returned")
            assert_ok(client.post(f"/api/records/{record_id}/review", headers=admin), "review after return")
            run = assert_ok(client.post(f"/api/agent/records/{record_id}/analysis", headers=admin), "agent record")
            assert run["output"]["evidenceRefs"]
            confirmed = assert_ok(client.post(f"/api/agent/runs/{run['id']}/confirm", headers=admin, json={"accepted": True}), "agent confirm")
            assert confirmed["accepted"] is True

        for project, cycle, town in [(yunan, yunan_cycle, "桂圩镇"), (maonan, maonan_cycle, "金塘镇")]:
            precheck = assert_ok(client.post("/api/report-tasks/precheck", headers=admin, json={
                "source": "dashboard", "projectId": project["id"], "period": cycle["name"], "townNames": [town], "outputs": ["separate", "summary"]
            }), f"report precheck {town}")
            assert precheck["ok"] is True
            assert precheck["summary"]["recordCount"] >= 1
            task = assert_ok(client.post("/api/report-tasks", headers=admin, json={
                "source": "dashboard", "projectId": project["id"], "period": cycle["name"], "townNames": [town], "outputs": ["separate", "summary"]
            }), f"report {town}")
            result = assert_ok(client.get(f"/api/report-tasks/{task['id']}"), f"report result {town}")
            assert result["status"] == "completed", result.get("error")
            assert result["dataSnapshot"]["towns"][0]["assessmentObject"]
            summary_report = next(item for item in result["reports"] if "汇总报告" in item["name"])
            snapshot_towns = {item["town"] for item in result["dataSnapshot"]["towns"]}
            if project["name"] == "郁南项目":
                expected_summary_towns = {"桂圩镇"}
                excluded_summary_towns = {item["name"] for item in yunan_towns} - expected_summary_towns
            if project["name"] == "茂南项目":
                expected_summary_towns = {"金塘镇", "镇盛镇", "茂南区"}
                excluded_summary_towns = {item["name"] for item in maonan_towns} - expected_summary_towns
            assert snapshot_towns == expected_summary_towns, "汇总报告只能包含本期已提交、已复核或已锁定的镇街"
            town_report = next(item for item in result["reports"] if item.get("town") == town)
            from app.core.database import SessionLocal
            from app.models import Report
            with SessionLocal() as session:
                path = Path(session.get(Report, town_report["id"]).storage_key)
                summary_path = Path(session.get(Report, summary_report["id"]).storage_key)
            assert path.is_file() and path.stat().st_size > 10000
            assert summary_path.is_file() and summary_path.stat().st_size > 10000
            summary_doc = Document(str(summary_path))
            summary_text = "\n".join(
                [paragraph.text for paragraph in summary_doc.paragraphs]
                + ["\t".join(cell.text for cell in row.cells) for table in summary_doc.tables for row in table.rows]
            )
            assert all(name in summary_text for name in expected_summary_towns)
            assert all(name not in summary_text for name in excluded_summary_towns)
            assert "未提交/未复核" not in summary_text
            check_docx(path, town, project["name"])
            preview = assert_ok(client.get(f"/api/reports/{town_report['id']}/preview"), f"report preview {town}")
            assert preview["content"]["paragraphCount"] > 0
            assert preview["content"]["tableCount"] > 0
            download = client.get(f"/api/reports/{town_report['id']}/download")
            assert download.status_code == 200 and len(download.content) > 10000

        summary_only = assert_ok(client.post("/api/report-tasks", headers=admin, json={
            "source": "dashboard", "projectId": maonan["id"], "period": maonan_cycle["name"], "townNames": ["金塘镇"], "outputs": ["summary"]
        }), "summary only report")
        summary_only_result = assert_ok(client.get(f"/api/report-tasks/{summary_only['id']}"), "summary only report result")
        assert summary_only_result["status"] == "completed", summary_only_result.get("error")
        assert len(summary_only_result["reports"]) == 1
        assert "汇总报告" in summary_only_result["reports"][0]["name"]
        assert {item["town"] for item in summary_only_result["dataSnapshot"]["towns"]} == {"金塘镇", "镇盛镇", "茂南区"}

        dynamic_cycle_payload = complete_payload(
            project=maonan,
            cycle=maonan_cycle,
            town="山阁镇",
            village="",
            facility_type="town_plant",
            indicators=standards[("茂南项目", "town_plant")][1],
        )
        dynamic_cycle_payload.pop("cycleId", None)
        dynamic_cycle_payload["period"] = "2030年第4季度"
        dynamic_created = assert_ok(
            client.post("/api/mobile/assessment-records", json=dynamic_cycle_payload, headers=inspector),
            "dynamic cycle assessment",
        )
        dynamic_record_id = dynamic_created["recordIds"][0]
        assert_ok(client.post(f"/api/mobile/assessment-records/{dynamic_record_id}/submit", headers=inspector), "submit dynamic cycle")
        dynamic_detail = assert_ok(client.get(f"/api/records/{dynamic_record_id}"), "dynamic cycle detail")
        assert dynamic_detail["cycleName"] == "2030年第4季度"
        dynamic_cycles = assert_ok(client.get("/api/mobile/assessment-cycles", params={"city_id": maonan["id"]}), "dynamic cycles")["items"]
        assert any(item["name"] == "2030年第4季度" for item in dynamic_cycles)
        dynamic_cleared = assert_ok(client.delete(
            "/api/mobile/assessment-records",
            headers=inspector,
            params={"city_id": maonan["id"], "period": "2030年第4季度"},
        ), "clear dynamic cycle")
        assert dynamic_cleared["recordCount"] == 1

        dashboard = assert_ok(client.get("/api/dashboard/towns", params={"city_id": yunan["id"]}), "dashboard")
        assert any(item["name"] == "桂圩镇" and item["recordCount"] >= 1 for item in dashboard["items"])
        maonan_dashboard = assert_ok(client.get("/api/dashboard/towns", params={"city_id": maonan["id"]}), "maonan dashboard")
        assert maonan_dashboard["overview"]["villageCount"] == 12
        print("PASS: two-project mobile -> dashboard -> review -> DOCX pipeline")


if __name__ == "__main__":
    main()
